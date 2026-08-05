import streamlit as st
import pandas as pd
import requests
import base64
from datetime import datetime, timedelta, timezone
import os
import tempfile
import textwrap
import time
from fpdf import FPDF
import plotly.express as px
import re
import io
import unicodedata
import zipfile

# --- ARRANQUE BLINDADO: IMPORTACIÓN OPCIONAL DE WORD ---
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    HAS_DOCX = True
except (ImportError, KeyError, Exception):
    # Evita caídas de arranque en el servidor de Streamlit Cloud
    HAS_DOCX = False

# --- IMPORTACIÓN DE HERRAMIENTAS GCS ---
try:
    from tools import leer_espejo_gcs, sobrescribir_archivo_gcs
except ImportError:
    pass

try:
    from tools import leer_espejo_gcs, sobrescribir_archivo_gcs, generar_docx_reporte_faltas_individual, NOMBRE_BUCKET_SISTEMA
except ImportError:
    NOMBRE_BUCKET_SISTEMA = "jovial-trilogy-306216.appspot.com"  # respaldo si tools.py no está disponible


# Usuarios autorizados para el repositorio documental. OJO: en login.py el nombre
# de usuario y el rol son valores DISTINTOS (st.secrets["credenciales"][usuario]["rol"]),
# así que aquí van NOMBRES DE USUARIO, no roles: tener rol de admin no da acceso.
# La comparación es exacta (en minúsculas) a propósito, para que nadie entre por
# parecido de nombre a documentos laborales sensibles.
USUARIOS_REPOSITORIO = {"oscar", "afajardo", "jaison"}

# ==============================================================================
# ALMACENAMIENTO DEL REPOSITORIO: CATBOX
# ==============================================================================
# Los archivos se suben a Catbox y el índice (quién es dueño de qué) vive en una
# hoja de Google Sheets, porque sin GCS no hay otro lugar donde guardarlo.
#
# ADVERTENCIA DE PRIVACIDAD: los enlaces de Catbox son PÚBLICOS. Cualquiera que
# tenga la URL puede abrir el documento sin contraseña. La app protege la LISTA
# (solo la ven los usuarios autorizados), pero no los archivos en sí.
CARPETA_REPOSITORIO = "expedientes_documentos"
HOJA_REPOSITORIO = "RepositorioDocs"

# Catbox rechaza estas extensiones. Como Word está entre ellas y sí se necesita,
# el archivo se sube renombrado a .bin y el nombre real queda en el índice; al
# descargarlo se restaura, así que para el usuario es transparente.
EXTENSIONES_BLOQUEADAS_CATBOX = {"doc", "docx", "exe", "scr", "cpl", "jar"}

MIMES_REPOSITORIO = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "doc": "application/msword",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "xls": "application/vnd.ms-excel",
    "csv": "text/csv",
    "zip": "application/zip",
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
}


def _subir_documento_repositorio(datos_bytes, nombre_original):
    """
    Sube un archivo a Catbox. Devuelve (url, None) o (None, error).
    Si la extensión está bloqueada por Catbox, se sube con extensión .bin.
    """
    ext = nombre_original.rsplit('.', 1)[-1].lower() if '.' in nombre_original else ''
    nombre_envio = nombre_original
    if ext in EXTENSIONES_BLOQUEADAS_CATBOX:
        base = nombre_original.rsplit('.', 1)[0]
        nombre_envio = f"{base}__{ext}.bin"

    try:
        url = subir_archivo_catbox(datos_bytes, nombre_envio)
        if url and str(url).startswith("http"):
            return str(url).strip(), None
        return None, f"Catbox devolvió una respuesta inesperada: {url}"
    except Exception as e:
        return None, str(e)


def _borrar_documento_repositorio(url):
    """
    Elimina un archivo de Catbox usando el userhash de la cuenta.
    Devuelve (True, None) o (False, error). Si el archivo ya no existe se
    considera éxito, para poder limpiar el índice igual.
    """
    try:
        nombre_en_catbox = str(url).rstrip('/').split('/')[-1]
        resp = requests.post(
            "https://catbox.moe/user/api.php",
            data={
                "reqtype": "deletefiles",
                "userhash": CATBOX_USERHASH,
                "files": nombre_en_catbox,
            },
            timeout=30
        )
        if resp.status_code == 200:
            return True, None
        return False, f"Catbox respondió {resp.status_code}: {resp.text[:120]}"
    except Exception as e:
        return False, str(e)


def _cargar_indice_repositorio(conn):
    """Lee el índice de documentos desde Google Sheets."""
    columnas = ["COLABORADOR", "CARPETA", "NOMBRE_ARCHIVO", "ENLACE", "TIPO",
                "TAMANO_KB", "DESCRIPCION", "SUBIDO_POR", "FECHA_SUBIDA"]
    try:
        df = conn.read(spreadsheet=st.secrets["url_base_datos"], worksheet=HOJA_REPOSITORIO, ttl=0)
        df = df.dropna(how="all")
    except Exception:
        df = None

    if df is None or df.empty:
        return pd.DataFrame(columns=columnas)

    for c in columnas:
        if c not in df.columns:
            df[c] = ""
    return df


def _guardar_indice_repositorio(conn, df):
    """
    Guarda el índice en Google Sheets. Si la hoja todavía no existe, la crea:
    conn.update() falla con WorksheetNotFound cuando la worksheet no está creada,
    y esa era la causa del error 'No se pudo guardar el índice: RepositorioDocs'.
    """
    try:
        conn.update(spreadsheet=st.secrets["url_base_datos"], worksheet=HOJA_REPOSITORIO, data=df)
        return True
    except Exception:
        pass

    # Segundo intento: crear la hoja y escribir en ella.
    try:
        conn.create(spreadsheet=st.secrets["url_base_datos"], worksheet=HOJA_REPOSITORIO, data=df)
        return True
    except Exception:
        pass

    # Tercer intento: crear vacía y luego actualizar (algunas versiones de
    # st-gsheets-connection no aceptan 'data' en create()).
    try:
        conn.create(spreadsheet=st.secrets["url_base_datos"], worksheet=HOJA_REPOSITORIO)
        conn.update(spreadsheet=st.secrets["url_base_datos"], worksheet=HOJA_REPOSITORIO, data=df)
        return True
    except Exception as e:
        st.error(f"No se pudo guardar el índice en la hoja '{HOJA_REPOSITORIO}': {e}")
        return False

# ==============================================================================
# CONFIGURACIÓN Y CARGA DE PERSONAL
# ==============================================================================
API_KEY_FREEIMAGE = st.secrets.get("api_freeimage", "6d207e02198a847aa98d0a2a901485a5")
CATBOX_USERHASH = st.secrets.get("catbox_userhash", "327c87ffe7f915a6d1ec367ee") # Tu userhash integrado de forma nativa [3]

def get_honduras_time():
    return datetime.now(timezone.utc) - timedelta(hours=6)

@st.cache_data(show_spinner=False)
def cargar_personal(filepath="personal_tecnico.txt"):
    try:
        if not os.path.exists(filepath): return []
        with open(filepath, 'r', encoding='utf-8') as f:
            lineas = f.readlines()
        nombres = []
        for linea in lineas:
            linea = linea.strip()
            if linea:
                nombre_crudo = linea.split(',')[0]
                nombre_limpio = " ".join(nombre_crudo.replace('\t', ' ').split()).upper()
                if nombre_limpio: nombres.append(nombre_limpio)
        return sorted(list(set(nombres)))
    except: return []

@st.cache_data(show_spinner=False)
def cargar_personal_admin(filepath="personal_sac.txt"):
    personal = {}
    try:
        if not os.path.exists(filepath): return personal
        with open(filepath, 'r', encoding='utf-8') as f:
            contenido = f.read().replace('\n', ' ')
        
        bloques_departamentos = contenido.split('.')
        
        for bloque in bloques_departamentos:
            bloque = bloque.strip()
            if not bloque: continue
            
            partes = bloque.split(',', 1)
            
            if len(partes) > 0:
                departamento = partes[0].strip().upper()
                
                if len(partes) > 1:
                    empleados = [e.strip().upper() for e in partes[1].split(';') if e.strip()]
                else:
                    empleados = []
                
                if departamento not in personal:
                    personal[departamento] = []
                personal[departamento].extend(empleados)
                
        return personal
    except:
        return {}

# ==============================================================================
# MOTOR AUXILIAR DE SUBIDA A NUBE (CATBOX PARA PDF / FREEIMAGE PARA IMÁGENES)
# ==============================================================================
def subir_archivo_catbox(file_bytes, file_name):
    """
    Sube cualquier tipo de archivo (especialmente PDFs) a Catbox.moe
    utilizando el userhash del usuario para guardarlo en su cuenta [3].
    """
    url = "https://catbox.moe/user/api.php"
    payload = {
        "reqtype": "fileupload",
        "userhash": CATBOX_USERHASH
    }
    files = {
        "fileToUpload": (file_name, file_bytes)
    }
    try:
        response = requests.post(url, data=payload, files=files, timeout=35)
        if response.status_code == 200:
            return response.text.strip()
        else:
            st.error(f"Error de Catbox ({response.status_code}): {response.text}")
            return None
    except Exception as e:
        st.error(f"Error al conectar con Catbox: {e}")
        return None

def subir_evidencias_inteligente(file_uploader_obj):
    """
    Detecta el formato de cada archivo. Si es PDF lo sube a Catbox [3],
    si es imagen lo hospeda en Freeimage de forma transparente.
    """
    name_lower = file_uploader_obj.name.lower()
    if name_lower.endswith('.pdf'):
        return subir_archivo_catbox(file_uploader_obj.getvalue(), file_uploader_obj.name)
    else:
        # Enrutamiento de imágenes a Freeimage
        try:
            res = requests.post(
                "https://freeimage.host/api/1/upload",
                data={
                    "key": API_KEY_FREEIMAGE,
                    "action": "upload",
                    "source": base64.b64encode(file_uploader_obj.getvalue()).decode('utf-8'),
                    "format": "json"
                },
                timeout=25
            )
            if res.status_code == 200:
                return res.json()["image"]["url"]
            return None
        except:
            # Fallback a Catbox si falla Freeimage
            return subir_archivo_catbox(file_uploader_obj.getvalue(), file_uploader_obj.name)

# ==============================================================================
# MOTOR DE CLASIFICACIÓN INTELIGENTE DE INCIDENCIAS (EXCLUSIÓN DE ÓRDENES)
# ==============================================================================
def es_llegada_tarde(motivo, comentario):
    motivo_u = str(motivo).upper().strip()
    com_u = str(comentario).upper().strip()
    
    EXCL = [
        'ALMUERZO', 'BREAK', 'DESCANSO', 'ORDEN', 'ÓRDEN', 'CIERRE', 
        'CERRADA', 'APERTURA', 'APERTURADA', 'LIQUIDADA', 'LIQUIDACION', 
        'LIQUIDACIÓN', 'CEQUI', 'SOP'
    ]
    KWS  = ['LLEGADA TARDE', 'LLEGADA TARDIA', 'LLEGADA TARDÍA', 'TARDANZA', 'TARDE', 'RETRASO', 'LLEGADA TARDÍA / AUSENCIA']
    
    texto_completo = motivo_u + " " + com_u
    if any(ex in texto_completo for ex in EXCL):
        return False
        
    for kw in KWS:
        if kw in motivo_u:
            return True
    for kw in KWS:
        if kw in com_u:
            return True
    return False

def clasificar_grave_o_leve(motivo, comentario, n_tardes=0):
    motivo_u = str(motivo).upper().strip()
    com_u = str(comentario).upper().strip()
    texto = motivo_u + " " + com_u

    # 1. VEHÍCULO / MAL CUIDADO O DESCUIDO - SIEMPRE GRAVE
    palabras_vehiculo_neglect = [
        "MAL CUIDADO", "CUIDADO VEHICULO", "CUIDADO DEL VEHICULO", "CUIDADO DE VEHICULO",
        "CUIDADO VEHÍCULO", "CUIDADO DEL VEHÍCULO", "CUIDADO DE VEHÍCULO",
        "MALTRATO VEHICULO", "MALTRATO DE VEHICULO", "MALTRATO DEL VEHICULO",
        "DESCUIDO VEHICULO", "DESCUIDO DE VEHICULO", "DESCUIDO DEL VEHICULO",
        "DESCUIDO VEHÍCULO", "DESCUIDO DE VEHÍCULO", "DESCUIDO DEL VEHÍCULO",
        "DAÑO VEHICULO", "DAÑO AL VEHICULO", "DAÑO DE VEHICULO",
        "DAÑO VEHÍCULO", "DAÑO AL VEHÍCULO", "DAÑO DE VEHÍCULO",
        "GOLPE VEHICULO", "GOLPE AL VEHICULO", "GOLPE DE VEHICULO",
        "ABUSO VEHICULO", "ABUSO DEL VEHICULO", "ABUSO DE VEHICULO",
        "DAÑO CARRO", "DAÑO AL CARRO", "DAÑO UNIDAD", "DAÑO A UNIDAD"
    ]
    if any(x in texto for x in palabras_vehiculo_neglect):
        return 'GRAVE'

    # 2. EVALUACIÓN DE SLA & MANEJO DE ÓRDENES (REGLA ASOCIATIVA INTELIGENTE)
    roots_objeto = ["ORDEN", "ÓRDEN", "RUTA"]
    roots_accion = ["APERTUR", "CERR", "CIERR", "INIC", "LIQUID", "FINALIZ"]
    roots_anomalia = ["TARDE", "TARDÍ", "TARDI", "DESFAS", "DESFAC", "RETRAS", "INCUMPLI"]

    has_objeto = any(obj in texto for obj in roots_objeto)
    has_accion = any(acc in texto for acc in roots_accion)
    has_anomalia = any(anom in texto for anom in roots_anomalia)

    if (has_objeto and has_accion and has_anomalia):
        return 'GRAVE'

    # 3. ACCIONES GRAVES DIRECTAS (Seguridad y Faltas de Respeto)
    palabras_graves_directas = [
        "FALTA DE RESPETO", "FALTA DE RESPET", "FALTAS DE RESPETO",
        "IRRESPETO", "INSULTO", "INSULTOS", "GOLPE", "PELEA",
        "AGRESION", "AGRESIÓN", "AMENAZA", "HOSTIGAMIENTO",
        "FALTA AL RESPETO", "OFENSA", "OFENSAS", "FALTA RESPETO",
        "APERTURADA TARDE", "CERRADA TARDE", "APERTURADAS TARDES", "CERRADAS TARDES",
        "MAL USO", "MAL MANEJO", "MALA MANIPULACIÓN", "MALA MANIPULACION",
        "CHOQUE", "ACCIDENTE", "ALCOHOL", "EBRIEDAD", "EBRIO", "DROGA", 
        "ROBO", "HURTO", "ABANDONO DE RUTA", "ABANDONO RUTA", "ABANDONO", 
        "INJUSTIFICADA", "DAÑO", "PÉRDIDA", "PERDIDA", "FRAUDE", "NEGLIGENCIA", 
        "REINCIDENCIA", "DORMIDO", "GRAVE", "IRRESPONSABILIDAD"
    ]
    if any(kw in texto for kw in palabras_graves_directas):
        return 'GRAVE'

    graves_combos = [
        ('ABANDONO', 'RUTA'),
        ('DAÑO', 'EQUIPO'), ('DAÑO', 'VEHICULO'),
        ('DAÑO', 'VEHÍCULO'), ('DAÑO', 'CARRO'), ('DAÑO', 'UNIDAD'),
        ('FALTA', 'RESPETO'),
        ('RESPETO', 'COMPAÑERO'), ('RESPETO', 'SUPERVISOR'),
        ('ORDENES', 'PENDIENTES'), ('ÓRDENES', 'PENDIENTES'),
        ('AUSENCIA', 'AVISO'), ('AUSENCIA', 'JUSTIF'),
        ('INASISTENCIA', 'AVISO'), ('INASISTENCIA', 'JUSTIF'),
        ('IRRESPETO', 'COMPAÑERO'), ('IRRESPETO', 'SUPERVISOR'),
    ]
    for w1, w2 in graves_combos:
        if w1 in texto and w2 in texto:
            return 'GRAVE'

    # 4. PROMOCIÓN AUTOMÁTICA POR REINCIDENCIA (3 o más llegadas tarde = GRAVE)
    if es_llegada_tarde(motivo, comentario):
        if n_tardes >= 3:
            return 'GRAVE'
        return 'LEVE'

    # 5. ACCIONES LEVES
    palabras_leves = [
        'ALMUERZO EXCEDIDO', 'HORA ALMUERZO', 'HORA DE ALMUERZO',
        'EXCEDIÓ ALMUERZO', 'EXCEDIO ALMUERZO',
        'TIEMPO DE ALMUERZO', 'TIEMPO ALMUERZO', 'EXCESO ALMUERZO',
        'BREAK EXCEDIDO', 'HORA BREAK', 'HORA DE BREAK',
        'EXCEDIÓ BREAK', 'EXCEDIO BREAK',
        'TIEMPO DE BREAK', 'TIEMPO BREAK', 'EXCESO BREAK',
        'DESCANSO EXCEDIDO',
        'NO MARCÓ', 'NO MARCO', 'SIN MARCAJE', 'MARCAJE FALTANTE',
        'NO MARCÓ ENTRADA', 'NO MARCÓ SALIDA',
        'NO MARCO ENTRADA', 'NO MARCO SALIDA',
        'NO REGISTRO ENTRADA', 'NO REGISTRO SALIDA',
        'OLVIDO MARCAJE', 'FALTA DE MARCAJE', 'OLVIDO DE MARCAJE',
        'MALA DOCUMENTACION', 'MALA DOCUMENTACIÓN',
        "TARDE", "RETRASO", "TRÁFICO", "TRAFICO", "LLANTA", "UNIFORME",
        "GAFETE", "SUCIO", "DESORDEN", "OLVIDO", "MINUTOS", "LEVE"
    ]
    if any(p in texto for p in palabras_leves):
        return 'LEVE'

    leves_combos = [
        ('HORA', 'ALMUERZO'), ('HORA', 'BREAK'),
        ('EXCEDIÓ', 'ALMUERZO'), ('EXCEDIÓ', 'BREAK'),
        ('NO', 'MARCÓ'), ('NO', 'MARCO'),
        ('SIN', 'MARCAJE'), ('SIN', 'MARCA'),
        ('MALA', 'DOCUMENTACION'), ('MALA', 'DOCUMENTACIÓN'),
    ]
    for w1, w2 in leves_combos:
        if w1 in texto and w2 in texto:
            return 'LEVE'

    return 'OTRO'

def asignar_rubro_automatico(motivo, comentario, n_tardes=0):
    motivo_str = str(motivo).upper().strip()
    clasificacion = clasificar_grave_o_leve(motivo, comentario, n_tardes)
    etiqueta = f"[{clasificacion}]"
    motivo_limpio = motivo_str.replace("[GRAVE]", "").replace("[LEVE]", "").replace("[OTRO]", "").strip()
    return f"{etiqueta} {motivo_limpio}"

# ==============================================================================
# 1. LÓGICA DE PDF (Clase Base)
# ==============================================================================
class MemoPDF(FPDF):
    def header(self):
        logo_path = 'logo.png'
        if os.path.exists(logo_path):
            try: self.image(logo_path, 10, 6, 35)
            except: pass
        self.set_y(10); self.set_x(50); self.set_text_color(0, 0, 0)
        self.set_font("Helvetica", "B", 10)
        self.cell(0, 5, "MAXCOM - DEPARTAMENTO DE CONTROL OPERATIVO", ln=True, align="R")
        self.set_font("Helvetica", "", 8); self.set_x(50)
        self.cell(0, 5, "Reporte Oficial de Gestion de Personal", ln=True, align="R")
        self.set_draw_color(200, 200, 200); self.line(10, 22, 200, 22); self.ln(10)
        
    def footer(self):
        self.set_y(-15); self.set_text_color(150, 150, 150); self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Pagina {self.page_no()}", align="C")

def sanitizar(texto):
    import unicodedata
    if pd.isna(texto) or texto is None: return "N/D"
    return unicodedata.normalize('NFKD', str(texto)).encode('ascii', 'ignore').decode('ascii')

# ==============================================================================
# 2. GENERADORES DE DOCUMENTOS DE REPORTES (PDF Y WORD)
# ==============================================================================
def generar_pdf_consolidado(df):
    df_work = pd.DataFrame()
    df_leves = df_graves = df_otros = pd.DataFrame()
    conteo_tardes = {}

    if not df.empty:
        df_work = df.copy()

        # Contar llegadas tarde filtradas para la promoción por reincidencia
        mask_tarde = df_work.apply(
            lambda r: es_llegada_tarde(
                str(r.get('TIPO_FALTA', '')).upper(),
                str(r.get('COMENTARIO', '')).upper()
            ), axis=1
        )
        conteo_tardes = (
            df_work[mask_tarde]['TECNICO']
            .astype(str).str.upper().str.strip()
            .value_counts().to_dict()
        )

        df_work['_CLASIF'] = df_work.apply(
            lambda r: clasificar_grave_o_leve(
                str(r.get('TIPO_FALTA', '')),
                str(r.get('COMENTARIO', '')),
                conteo_tardes.get(str(r.get('TECNICO', '')).upper().strip(), 0)
            ), axis=1
        )
        df_leves  = df_work[df_work['_CLASIF'] == 'LEVE'].copy()
        df_graves = df_work[df_work['_CLASIF'] == 'GRAVE'].copy()
        df_otros  = df_work[df_work['_CLASIF'] == 'OTRO'].copy()

    n_leves  = len(df_leves)
    n_graves = len(df_graves)
    n_otros  = len(df_otros)

    def _dibujar_tabla_clasif(pdf_obj, df_t, etiqueta, desc_corta,
                               hr, hg, hb,
                               rr, rg, rb,
                               thr=255, thg=255, thb=255):
        if df_t.empty:
            return

        n = len(df_t)

        pdf_obj.set_font("Helvetica", "B", 11)
        pdf_obj.set_fill_color(hr, hg, hb)
        pdf_obj.set_text_color(thr, thg, thb)
        pdf_obj.cell(0, 8,
                     f"  {etiqueta}  ({n} incidencia{'s' if n != 1 else ''})",
                     border=1, fill=True, ln=True)
        pdf_obj.set_font("Helvetica", "I", 7)
        pdf_obj.set_text_color(100, 100, 100)
        pdf_obj.cell(0, 5, f"  {desc_corta}", ln=True)
        pdf_obj.ln(2)

        W    = [25, 48, 42, 55, 20]
        HDRS = ["FECHA", "COLABORADOR", "TIPO DE FALTA", "DESCRIPCION", "SUPERVISOR"]

        pdf_obj.set_fill_color(hr, hg, hb)
        pdf_obj.set_text_color(thr, thg, thb)
        pdf_obj.set_font("Helvetica", "B", 7)
        for i, h in enumerate(HDRS):
            pdf_obj.cell(W[i], 6, h, border=1, fill=True, align="C")
        pdf_obj.ln()

        pdf_obj.set_text_color(40, 40, 40)
        for idx_fila, (_, row) in enumerate(df_t.iterrows()):
            f_inc = sanitizar(str(row.get('FECHA_INCIDENCIA', ''))[:16])
            tec   = sanitizar(str(row.get('TECNICO', ''))[:40])
            mot   = sanitizar(str(row.get('TIPO_FALTA', ''))[:35])
            com   = sanitizar(str(row.get('COMENTARIO', '')))
            sup   = sanitizar(str(row.get('SUPERVISOR', ''))[:18])

            lineas = textwrap.wrap(com, width=40)
            if not lineas:
                lineas = [""]

            if idx_fila % 2 == 0:
                pdf_obj.set_fill_color(rr, rg, rb)
            else:
                pdf_obj.set_fill_color(255, 255, 255)

            pdf_obj.set_font("Helvetica", "", 7)
            for i_l, linea in enumerate(lineas):
                b_t = 'T' if i_l == 0 else ''
                b_b = 'B' if i_l == len(lineas) - 1 else ''
                bs  = 'LR' + b_t + b_b

                pdf_obj.cell(W[0], 5, f_inc if i_l == 0 else "", border=bs, fill=True)
                pdf_obj.cell(W[1], 5, tec   if i_l == 0 else "", border=bs, fill=True)
                pdf_obj.cell(W[2], 5, mot   if i_l == 0 else "", border=bs, fill=True)
                pdf_obj.cell(W[3], 5, f" {linea}",               border=bs, fill=True)
                pdf_obj.cell(W[4], 5, sup   if i_l == 0 else "", border=bs, fill=True, ln=True)

        pdf_obj.ln(6)

    pdf = MemoPDF()
    pdf.alias_nb_pages()
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(40, 50, 100)
    pdf.cell(0, 10, "REPORTE CONSOLIDADO DE EXPEDIENTES", ln=True, align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 6, f"Generado el: {get_honduras_time().strftime('%d/%m/%Y a las %H:%M:%S')}", ln=True, align="C")
    pdf.ln(8)

    if df.empty:
        pdf.set_font("Helvetica", "I", 12)
        pdf.cell(0, 10, "No hay registros disponibles.", ln=True, align="C")
    else:
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(40, 50, 100)
        pdf.cell(0, 6, "RESUMEN DE INCIDENCIAS POR NIVEL DE GRAVEDAD", ln=True)
        pdf.ln(2)

        W_KPI = 63
        KPI_DATA = [
            ("  FALTAS LEVES",        n_leves,  200, 140,  20, 255, 248, 220, 100, 50,  0),
            ("  FALTAS GRAVES",       n_graves, 180,  30,  30, 255, 235, 235, 255, 255, 255),
            ("  OTRAS INCIDENCIAS",   n_otros,   80,  95, 115, 235, 238, 245, 255, 255, 255),
        ]
        for (label, _, hr, hg, hb, rr, rg, rb, thr, thg, thb) in KPI_DATA:
            pdf.set_fill_color(hr, hg, hb)
            pdf.set_text_color(thr, thg, thb)
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(W_KPI, 7, label, border=1, fill=True)
        pdf.ln()
        for (_, count, hr, hg, hb, rr, rg, rb, thr, thg, thb) in KPI_DATA:
            pdf.set_fill_color(rr, rg, rb)
            pdf.set_text_color(hr, hg, hb)
            pdf.set_font("Helvetica", "B", 18)
            pdf.cell(W_KPI, 12, f"  {count}", border=1, fill=True)
        pdf.ln()
        pdf.ln(10)

        pdf.set_font("Helvetica", "I", 7)
        pdf.set_text_color(120, 120, 120)
        pdf.cell(0, 5,
                 "Nota: Las llegadas tardes se clasifican como LEVE hasta 2 ocurrencias por colaborador. "
                 "A partir de la 3a se promueven automaticamente a GRAVE.",
                 ln=True)
        pdf.ln(5)

        if n_leves > 0:
            _dibujar_tabla_clasif(
                pdf, df_leves,
                etiqueta   = "FALTAS LEVES",
                desc_corta = "Llegadas tardes (<3), almuerzo/break excedido, no marco entrada/salida, mala documentacion",
                hr=200, hg=140, hb=20,
                rr=255, rg=248, rb=220,
                thr=255, thg=255, thb=255,
            )

        if n_graves > 0:
            if pdf.get_y() > 200:
                pdf.add_page()
            _dibujar_tabla_clasif(
                pdf, df_graves,
                etiqueta   = "FALTAS GRAVES",
                desc_corta = "Mal cuidado de vehiculos, no apertura/cierre, ordenes pendientes, insultos/irrespeto, o >=3 llegadas tarde",
                hr=180, hg=30, hb=30,
                rr=255, rg=235, rb=235,
                thr=255, thg=255, thb=255,
            )

        if n_otros > 0:
            if pdf.get_y() > 200:
                pdf.add_page()
            _dibujar_tabla_clasif(
                pdf, df_otros,
                etiqueta   = "OTRAS INCIDENCIAS",
                desc_corta = "Incidencias medicas, meritos, documentos administrativos y registros sin clasificacion definida",
                hr=80, hg=95, hb=115,
                rr=235, rg=238, rb=245,
                thr=255, thg=255, thb=255,
            )

    # BLINDAJE DE SEGURIDAD: Evita que el generador intente insertar PDFs como fotos (Imagen 1) [3]
    tiene_anexos = False
    for _, row in df.iterrows():
        urls = str(row.get('URL_FOTO', '')).split(',')
        if any(u.strip().startswith('http') and not u.strip().lower().endswith('.pdf') for u in urls):
            tiene_anexos = True
            break

    if tiene_anexos:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(40, 50, 100)
        pdf.cell(0, 10, "ANEXOS - EVIDENCIA FOTOGRÁFICA", ln=True, align="C")
        pdf.ln(5)
        for _, row in df.iterrows():
            urls    = str(row.get('URL_FOTO', '')).split(',')
            validas = [u.strip() for u in urls if u.strip().startswith('http') and not u.strip().lower().endswith('.pdf')]
            if validas:
                tec_name     = sanitizar(str(row.get('TECNICO', '')))
                f_inc        = sanitizar(str(row.get('FECHA_INCIDENCIA', '')))
                motivo_falta = sanitizar(str(row.get('TIPO_FALTA', '')))
                for url in validas:
                    try:
                        r = requests.get(url, timeout=5)
                        if r.status_code == 200:
                            fd, tp = tempfile.mkstemp(suffix=".png")
                            os.close(fd)
                            try:
                                with open(tp, 'wb') as f: f.write(r.content)
                                if pdf.get_y() > 60:
                                    pdf.add_page()
                                pdf.set_font("Helvetica", "B", 9)
                                pdf.set_text_color(0, 0, 0)
                                pdf.set_fill_color(240, 240, 240)
                                pdf.cell(0, 8, f" Evidencia: {tec_name} | {motivo_falta} | {f_inc}",
                                         ln=True, fill=True, border=1)
                                pdf.ln(3)
                                pdf.image(tp, x=20, w=150)
                                pdf.ln(10)
                            finally:
                                if os.path.exists(tp):
                                    os.remove(tp)
                    except:
                        pass

    fd, path = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    pdf.output(path)
    with open(path, "rb") as f:
        data = f.read()
    os.remove(path)
    return data


def generar_docx_consolidado(df):
    """
    Genera un documento oficial de Word (.docx) estructurado con tablas,
    resúmenes de KPIs y evidencia fotográfica adjunta en el mismo formato.
    Si se detecta un solo colaborador, genera automáticamente la plantilla idéntica
    al reporte de faltas de la imagen.
    """
    if not HAS_DOCX:
        return b""
        
    # --- REDIRECCIÓN AL FORMATO EXACTO DE REPORTE DE FALTAS DE LA IMAGEN ---
    if not df.empty and df['TECNICO'].nunique() == 1:
        try:
            return generar_docx_reporte_faltas_individual(df)
        except Exception as e:
            # Fallback en caso de que ocurra algún error inesperado
            pass
        
    doc = Document()
    
    # --- Estilo de Página Compacto (Márgenes de 0.5 pulg para asegurar que quepa en una página) ---
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    # --- Configurar tipografía estándar ---
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Arial'
    font_normal.size = Pt(8.5)

    def set_cell_background(cell, hex_color):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def aplicar_espaciado_celda(cell):
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0

    # Extraemos información base del primer elemento para las cabeceras principales
    primer_registro = df.iloc[0]
    nombre_completo = str(primer_registro.get('TECNICO', ''))
    nombre_limpio_emp = re.sub(r'\s*\(.*\)$', '', nombre_completo).strip()
    
    fecha_falta_val = str(primer_registro.get('FECHA_INCIDENCIA', ''))
    hora_falta_val = extraer_hora_falta(str(primer_registro.get('COMENTARIO', '')), str(primer_registro.get('FECHA_REGISTRO', '')))

    # ==========================================================================
    # 1. TABLA ENCABEZADO (Logo | Título | Metadatos) - Imagen 1
    # ==========================================================================
    header_table = doc.add_table(rows=1, cols=3)
    header_table.style = 'Table Grid'
    header_table.autofit = False
    
    header_table.columns[0].width = Inches(1.3)
    header_table.columns[1].width = Inches(3.7)
    header_table.columns[2].width = Inches(2.5)

    # Columna 1: Logotipo (Utiliza logo.png)
    cell_logo = header_table.cell(0, 0)
    p_logo = cell_logo.paragraphs[0]
    logo_path = 'logo.png'
    if os.path.exists(logo_path):
        try: p_logo.add_run().add_picture(logo_path, width=Inches(1.1))
        except: p_logo.text = "MAXCOM"
    else:
        p_logo.text = "MAXCOM"
        p_logo.runs[0].font.bold = True
        p_logo.runs[0].font.size = Pt(12)
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Columna 2: Título Central de la tabla de cabecera
    cell_title = header_table.cell(0, 1)
    p_title = cell_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("\nREPORTE DE FALTAS")
    run_title.font.size = Pt(10)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(128, 128, 128) # Gris opaco

    # Columna 3: Cuadro de Metadatos
    cell_meta = header_table.cell(0, 2)
    meta_table = cell_meta.add_table(rows=4, cols=2)
    meta_table.style = 'Table Grid'
    meta_table.autofit = True
    
    meta_fields = [
        ("CÓDIGO", "HN-GG-RH-FR-09"),
        ("VERSIÓN", "1.0"),
        ("FECHA", "12/05/2026"),
        ("CLASIFICACIÓN", "INTERNO")
    ]
    for r_idx, (k, v) in enumerate(meta_fields):
        m_cells = meta_table.rows[r_idx].cells
        m_cells[0].text = k
        m_cells[0].paragraphs[0].runs[0].font.bold = True
        m_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        m_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(100, 116, 139)
        
        m_cells[1].text = v
        m_cells[1].paragraphs[0].runs[0].font.size = Pt(7)
        m_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(100, 116, 139)
        
        aplicar_espaciado_celda(m_cells[0])
        aplicar_espaciado_celda(m_cells[1])

    for cell in header_table.rows[0].cells:
        aplicar_espaciado_celda(cell)

    # --- TÍTULO INDEPENDIENTE DEBAJO DE LA CABECERA (Imagen 1) ---
    p_title_below = doc.add_paragraph()
    p_title_below.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title_below.paragraph_format.space_before = Pt(12)
    p_title_below.paragraph_format.space_after = Pt(12)
    run_title_below = p_title_below.add_run("REPORTE DE FALTAS")
    run_title_below.font.name = 'Arial'
    run_title_below.font.size = Pt(11)
    run_title_below.font.bold = True
    run_title_below.font.color.rgb = RGBColor(0, 0, 0)

    # ==========================================================================
    # 2. TABLA: DATOS DEL JEFE SOLICITANTE
    # ==========================================================================
    table_jefe = doc.add_table(rows=6, cols=2)
    table_jefe.style = 'Table Grid'
    table_jefe.columns[0].width = Inches(2.5)
    table_jefe.columns[1].width = Inches(5.0)

    # Cabecera azul
    hdr_cell_j = table_jefe.rows[0].cells[0].merge(table_jefe.rows[0].cells[1])
    hdr_cell_j.text = "Datos del Jefe del Departamento Solicitante"
    set_cell_background(hdr_cell_j, "1E1B4B") # Navy oscuro institucional
    p_hdr_j = hdr_cell_j.paragraphs[0]
    p_hdr_j.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_hdr_j = p_hdr_j.runs[0]
    run_hdr_j.font.bold = True
    run_hdr_j.font.size = Pt(9)
    run_hdr_j.font.color.rgb = RGBColor(255, 255, 255)

    jefe_labels = [
        "Nombre del Jefe del Departamento",
        "Puesto",
        "Departamento",
        "Ciudad",
        "Fecha de Ingresos"
    ]
    for f_idx, label in enumerate(jefe_labels):
        cells = table_jefe.rows[f_idx + 1].cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].font.bold = True
        cells[0].paragraphs[0].runs[0].font.size = Pt(8.5)
        cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(30, 41, 59)
        
        cells[1].text = "" # Editable manual
        
        aplicar_espaciado_celda(cells[0])
        aplicar_espaciado_celda(cells[1])

    aplicar_espaciado_celda(hdr_cell_j)
    doc.add_paragraph()

    # ==========================================================================
    # 3. TABLA: DATOS DEL EMPLEADO REPORTADO (AUTO-COMPLETADOS)
    # ==========================================================================
    table_emp = doc.add_table(rows=7, cols=2)
    table_emp.style = 'Table Grid'
    table_emp.columns[0].width = Inches(2.5)
    table_emp.columns[1].width = Inches(5.0)

    # Cabecera azul
    hdr_cell_e = table_emp.rows[0].cells[0].merge(table_emp.rows[0].cells[1])
    hdr_cell_e.text = "Datos de Empleado Reportado"
    set_cell_background(hdr_cell_e, "1E1B4B")
    p_hdr_e = hdr_cell_e.paragraphs[0]
    p_hdr_e.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_hdr_e = p_hdr_e.runs[0]
    run_hdr_e.font.bold = True
    run_hdr_e.font.size = Pt(9)
    run_hdr_e.font.color.rgb = RGBColor(255, 255, 255)

    emp_fields = [
        ("Nombre del Empleado", nombre_limpio_emp), # Autocompletado
        ("Código de Empleado", ""),
        ("Puesto", ""),
        ("Horario de Trabajo", ""),
        ("Fecha de la Falta Ocurrida", fecha_falta_val), # Autocompletado
        ("Hora de la Falta Ocurrida", hora_falta_val)   # Autocompletado
    ]

    for f_idx, (label, val) in enumerate(emp_fields):
        cells = table_emp.rows[f_idx + 1].cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].font.bold = True
        cells[0].paragraphs[0].runs[0].font.size = Pt(8.5)
        cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(30, 41, 59)
        
        cells[1].text = val
        if val:
            cells[1].paragraphs[0].runs[0].font.size = Pt(8.5)
            
        aplicar_espaciado_celda(cells[0])
        aplicar_espaciado_celda(cells[1])

    aplicar_espaciado_celda(hdr_cell_e)
    doc.add_paragraph()

    # ==========================================================================
    # 4. TABLA: DETALLE DE LA FALTA (CON CONSOLIDACIÓN EN RENGLONES DISTINTOS)
    # ==========================================================================
    table_det = doc.add_table(rows=12, cols=1)
    table_det.style = 'Table Grid'

    # Cabecera azul
    hdr_cell_d = table_det.rows[0].cells[0]
    hdr_cell_d.text = "Detalle de la Falta"
    set_cell_background(hdr_cell_d, "1E1B4B")
    p_hdr_d = hdr_cell_d.paragraphs[0]
    p_hdr_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_hdr_d = p_hdr_d.runs[0]
    run_hdr_d.font.bold = True
    run_hdr_d.font.size = Pt(9)
    run_hdr_d.font.color.rgb = RGBColor(255, 255, 255)
    aplicar_espaciado_celda(hdr_cell_d)

    # --- AUTOCOMPLETADO DE CADA INCIDENCIA EN RENGLONES DISTINTOS ---
    idx_renglon = 1
    for _, row_inc in df.iterrows():
        if idx_renglon >= 12: # Límite para evitar que salte de página
            break
        fecha_p = str(row_inc.get('FECHA_INCIDENCIA', ''))
        tipo_p = str(row_inc.get('TIPO_FALTA', ''))
        desc_p = str(row_inc.get('COMENTARIO', ''))
        
        cell_renglon = table_det.rows[idx_renglon].cells[0]
        cell_renglon.text = f"• [{fecha_p}] {tipo_p}: {desc_p}"
        cell_renglon.paragraphs[0].runs[0].font.size = Pt(8.5)
        aplicar_espaciado_celda(cell_renglon)
        idx_renglon += 1

    # Renglones restantes en blanco con altura para simular diseño impreso original
    for r_idx in range(idx_renglon, 12):
        cell_vacia = table_det.rows[r_idx].cells[0]
        cell_vacia.text = ""
        table_det.rows[r_idx].height = Inches(0.18)
        aplicar_espaciado_celda(cell_vacia)

    doc.add_paragraph()

    # Texto Legal intermedio
    p_legal = doc.add_paragraph()
    run_legal = p_legal.add_run("Firmo en señal de solicitud/autorización/aprobación lo de arriba detallado.")
    run_legal.font.size = Pt(8.5)
    run_legal.italic = True
    p_legal.paragraph_format.space_before = Pt(4)
    p_legal.paragraph_format.space_after = Pt(4)

    # ==========================================================================
    # 5. TABLA INFERIOR: FIRMAS Y ELABORACIÓN (CON HORA DEL SISTEMA ACTUAL)
    # ==========================================================================
    table_bottom = doc.add_table(rows=2, cols=2)
    table_bottom.style = 'Table Grid'
    table_bottom.columns[0].width = Inches(4.5)
    table_bottom.columns[1].width = Inches(3.0)

    # --- FILA 0: ENCABEZADOS AZULES ---
    # Celda izquierda: Totalmente azul oscuro sin texto (Imagen 2)
    cell_elab_hdr = table_bottom.cell(0, 0)
    cell_elab_hdr.text = "" 
    set_cell_background(cell_elab_hdr, "1E1B4B")

    # Celda derecha: Azul oscuro con texto de Firma
    cell_firma_hdr = table_bottom.cell(0, 1)
    cell_firma_hdr.text = "" 
    set_cell_background(cell_firma_hdr, "1E1B4B")
    p_firma_hdr = cell_firma_hdr.paragraphs[0]
    p_firma_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_firma_hdr = p_firma_hdr.add_run("Firma de Jefe de\nDepartamento Solicitante")
    run_firma_hdr.font.bold = True
    run_firma_hdr.font.size = Pt(8.5)
    run_firma_hdr.font.color.rgb = RGBColor(255, 255, 255)

    # --- FILA 1: CONTENIDO ---
    cell_elab_cont = table_bottom.cell(1, 0)
    left_subtable = cell_elab_cont.add_table(rows=2, cols=2)
    left_subtable.style = 'Table Grid'

    # Fecha y hora actual del sistema en Honduras
    ahora_hn = get_honduras_time()
    fecha_elab = ahora_hn.strftime("%d/%m/%Y")
    hora_elab = ahora_hn.strftime("%I:%M%p").lower() # Formato compacto como 10:10am

    elab_fields = [
        ("Fecha de elaboración de reporte", fecha_elab),
        ("Hora de elaboración de reporte", hora_elab)
    ]
    for f_idx, (k, v) in enumerate(elab_fields):
        l_cells = left_subtable.rows[f_idx].cells
        l_cells[0].text = k
        l_cells[0].paragraphs[0].runs[0].font.bold = True
        l_cells[0].paragraphs[0].runs[0].font.size = Pt(8)
        
        l_cells[1].text = v
        if v:
            l_cells[1].paragraphs[0].runs[0].font.size = Pt(8)
            
        aplicar_espaciado_celda(l_cells[0])
        aplicar_espaciado_celda(l_cells[1])

    # Columna Derecha de Firma: SE MANTIENE TOTALMENTE COLOR BLANCO PARA LA PLUMA
    cell_firma_cont = table_bottom.cell(1, 1)
    cell_firma_cont.text = "" # Fondo blanco liso
    cell_firma_cont.add_paragraph("\n\n\n")

    # Aplicar sangría final a las celdas de la tabla inferior
    for r_idx in range(2):
        for c_idx in range(2):
            aplicar_espaciado_celda(table_bottom.cell(r_idx, c_idx))

    # Retornar como Bytes seguros en memoria
    b_io = io.BytesIO()
    doc.save(b_io)
    return b_io.getvalue()


# ==============================================================================
# MOTOR DE MEMORIA
# ==============================================================================
def forzar_actualizacion_memoria(conn):
    df = leer_espejo_gcs(NOMBRE_BUCKET_SISTEMA, "expedientes_maestro.csv")
    if df is None or df.empty:
        df = conn.read(spreadsheet=st.secrets["url_base_datos"], worksheet="Expedientes", ttl=0)
    st.session_state['df_exp_memoria'] = df

def obtener_datos_memoria(conn):
    if 'df_exp_memoria' not in st.session_state:
        forzar_actualizacion_memoria(conn)
    return st.session_state['df_exp_memoria']

# ==============================================================================
# 3. INTERFAZ DE EXPEDIENTES Y VISTAS AISLADAS
# ==============================================================================
# ==============================================================================
# REPOSITORIO DOCUMENTAL DE EXPEDIENTES (archivos en Catbox, agrupados por colaborador)
# ==============================================================================
def _sanitizar_nombre_carpeta(nombre):
    """
    Convierte el nombre de un colaborador en un identificador de carpeta seguro:
    sin acentos, sin caracteres raros y con guiones bajos en vez de espacios.
    El nombre legible original se conserva en el índice CSV.
    """
    txt = unicodedata.normalize('NFKD', str(nombre)).encode('ascii', 'ignore').decode('ascii')
    txt = txt.upper().strip()
    txt = re.sub(r'[^A-Z0-9\s_-]', '', txt)
    txt = re.sub(r'\s+', '_', txt)
    return txt or "SIN_NOMBRE"


def _finalizar_subida(conn, df_indice, nuevos_registros, fallidos, colaborador=None):
    """
    Guarda el índice y reporta el resultado. La comparte la subida de archivos
    sueltos y la de ZIP para que ambas se comporten igual ante un fallo.
    """
    if nuevos_registros:
        df_actualizado = pd.concat([df_indice, pd.DataFrame(nuevos_registros)], ignore_index=True)
        if _guardar_indice_repositorio(conn, df_actualizado):
            destino = f" en la carpeta de {colaborador}" if colaborador else ""
            personas = sorted({r["COLABORADOR"] for r in nuevos_registros})
            detalle = f" ({len(personas)} colaborador(es))" if len(personas) > 1 else ""
            st.success(f"✅ {len(nuevos_registros)} archivo(s) guardado(s){destino}{detalle}.")
            time.sleep(1.5)
            st.rerun()
        else:
            # CRÍTICO: el archivo ya está en Catbox, pero si el índice no se
            # guarda, el enlace se pierde para siempre (Catbox no permite listar
            # los archivos de una cuenta). Por eso se muestran aquí para poder
            # rescatarlos a mano antes de recargar la página.
            st.error(
                "⚠️ Los archivos SÍ se subieron, pero no se pudo guardar el índice. "
                "**Copia estos enlaces ahora**: no hay forma de recuperarlos después."
            )
            df_rescate = pd.DataFrame(nuevos_registros)[["COLABORADOR", "NOMBRE_ARCHIVO", "ENLACE"]]
            st.dataframe(df_rescate, use_container_width=True, hide_index=True)
            st.download_button(
                "📥 Descargar estos enlaces (CSV)",
                data=df_rescate.to_csv(index=False).encode('utf-8'),
                file_name=f"enlaces_rescate_{get_honduras_time().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
                key=f"rescate_{get_honduras_time().strftime('%H%M%S')}"
            )

    if fallidos:
        st.error("❌ No se pudieron subir:\n\n- " + "\n- ".join(fallidos[:15]))
        if len(fallidos) > 15:
            st.caption(f"...y {len(fallidos) - 15} más.")


def mostrar_repositorio_documentos(conn):
    """
    Repositorio de documentos (PDF / Word / Excel) guardados en Google Cloud
    Storage, organizados en una carpeta por colaborador.

    A diferencia de las evidencias de incidencias (que van a Catbox y quedan con
    URL pública), estos archivos viven en el bucket privado de la empresa: solo
    se pueden abrir a través de la aplicación y por usuarios autorizados. Esto
    importa porque aquí se guardan documentos laborales sensibles.
    """
    st.subheader("🗄️ Repositorio de Documentos")

    usuario_actual = str(st.session_state.get('usuario_actual', '')).strip().lower()
    if usuario_actual not in USUARIOS_REPOSITORIO:
        st.warning("🔒 No tienes acceso a este repositorio. Solicítalo al administrador.")
        st.caption(f"Usuario conectado: `{usuario_actual or '(no detectado)'}`")
        return

    st.caption(
        "Los archivos se guardan en Catbox y se organizan por colaborador. "
        "⚠️ Los enlaces son públicos: quien tenga la URL puede abrir el documento."
    )

    df_indice = _cargar_indice_repositorio(conn)

    # ----------------------------------------------------------------------
    # LISTA DE COLABORADORES (técnicos + administrativos + carpetas existentes)
    # ----------------------------------------------------------------------
    nombres_tecnicos = cargar_personal("personal_tecnico.txt")
    nombres_admin = []
    try:
        for _dep, _emps in cargar_personal_admin("personal_sac.txt").items():
            nombres_admin.extend(_emps)
    except Exception:
        pass

    nombres_existentes = []
    if not df_indice.empty:
        nombres_existentes = df_indice['COLABORADOR'].dropna().astype(str).unique().tolist()

    lista_colaboradores = sorted(set(
        [n for n in nombres_tecnicos if n] +
        [n for n in nombres_admin if n] +
        [n for n in nombres_existentes if n]
    ))

    tab_subir, tab_explorar = st.tabs(["📤 Subir Documento", "📂 Explorar Carpetas"])

    # ======================================================================
    # SUBIR
    # ======================================================================
    with tab_subir:
        col_a, col_b = st.columns(2)
        with col_a:
            modo_nombre = st.radio(
                "Colaborador:",
                ["Seleccionar de la lista", "Escribir uno nuevo"],
                horizontal=True,
                key="repo_modo_nombre"
            )
            if modo_nombre == "Seleccionar de la lista" and lista_colaboradores:
                colaborador_sel = st.selectbox("👤 Nombre:", options=lista_colaboradores, key="repo_colab_sel")
            else:
                colaborador_sel = st.text_input("👤 Nombre del colaborador:", key="repo_colab_txt", placeholder="Ej: JUAN PEREZ").strip().upper()

        with col_b:
            descripcion_doc = st.text_input("📝 Descripción (opcional):", key="repo_desc", placeholder="Ej: Contrato firmado 2026")

        st.divider()
        modo_carga = st.radio(
            "¿Qué vas a subir?",
            ["📄 Archivos sueltos", "🗂️ Una carpeta completa (ZIP)"],
            horizontal=True,
            key="repo_modo_carga"
        )

        # ------------------------------------------------------------------
        # MODO A: ARCHIVOS SUELTOS
        # ------------------------------------------------------------------
        if modo_carga == "📄 Archivos sueltos":
            archivos_subir = st.file_uploader(
                "📎 Archivos (PDF, Word, Excel, imágenes):",
                type=list(MIMES_REPOSITORIO.keys()),
                accept_multiple_files=True,
                key="repo_uploader"
            )
            st.caption("💡 Puedes abrir la carpeta, seleccionar todo con **Ctrl+A** y arrastrar los archivos aquí.")

            if colaborador_sel:
                st.caption(f"📁 Se guardará en: `{CARPETA_REPOSITORIO}/{_sanitizar_nombre_carpeta(colaborador_sel)}/`")

            if st.button("☁️ Subir al Repositorio", type="primary", use_container_width=True, key="repo_btn_subir"):
                if not colaborador_sel:
                    st.error("⚠️ Debes indicar el colaborador.")
                elif not archivos_subir:
                    st.error("⚠️ Selecciona al menos un archivo.")
                else:
                    carpeta = _sanitizar_nombre_carpeta(colaborador_sel)
                    nuevos_registros = []
                    fallidos = []

                    barra = st.progress(0.0, text="Subiendo archivos...")
                    for i, arch in enumerate(archivos_subir):
                        extension = arch.name.rsplit('.', 1)[-1].lower() if '.' in arch.name else ''
                        try:
                            arch.seek(0)
                        except Exception:
                            pass
                        enlace, err = _subir_documento_repositorio(arch.read(), arch.name)
                        if enlace:
                            nuevos_registros.append({
                                "COLABORADOR": colaborador_sel,
                                "CARPETA": carpeta,
                                "NOMBRE_ARCHIVO": arch.name,
                                "ENLACE": enlace,
                                "TIPO": extension.upper(),
                                "TAMANO_KB": round(arch.size / 1024, 1) if getattr(arch, 'size', None) else "",
                                "DESCRIPCION": descripcion_doc,
                                "SUBIDO_POR": st.session_state.get('usuario_actual', ''),
                                "FECHA_SUBIDA": get_honduras_time().strftime('%Y-%m-%d %H:%M:%S'),
                            })
                        else:
                            fallidos.append(f"{arch.name} ({err})")

                        barra.progress((i + 1) / len(archivos_subir), text=f"Subiendo {arch.name}...")

                    barra.empty()
                    _finalizar_subida(conn, df_indice, nuevos_registros, fallidos, colaborador_sel)

        # ------------------------------------------------------------------
        # MODO B: CARPETA COMPLETA EN ZIP
        # ------------------------------------------------------------------
        else:
            st.info(
                "Comprime la carpeta en tu computadora (clic derecho → **Enviar a → Carpeta comprimida** en Windows, "
                "o **Comprimir** en Mac) y sube el ZIP aquí. Se conservan las subcarpetas."
            )

            modo_zip = st.radio(
                "¿Cómo se organiza el ZIP?",
                [
                    "Todo pertenece a UN colaborador",
                    "Cada carpeta dentro del ZIP es un colaborador distinto",
                ],
                key="repo_modo_zip"
            )

            if modo_zip == "Cada carpeta dentro del ZIP es un colaborador distinto":
                st.caption(
                    "Ejemplo: un ZIP con `JUAN PEREZ/contrato.pdf` y `MARIA LOPEZ/constancia.pdf` "
                    "creará una carpeta por cada persona automáticamente."
                )

            archivo_zip = st.file_uploader(
                "🗂️ Archivo ZIP:",
                type=["zip"],
                accept_multiple_files=False,
                key="repo_uploader_zip"
            )

            # Vista previa del contenido antes de subir nada, para que no haya
            # sorpresas con archivos basura o carpetas mal nombradas.
            contenido_zip = []
            if archivo_zip is not None:
                try:
                    archivo_zip.seek(0)
                    with zipfile.ZipFile(io.BytesIO(archivo_zip.read())) as zf:
                        for info in zf.infolist():
                            if info.is_dir():
                                continue
                            nombre_interno = info.filename
                            base = nombre_interno.split('/')[-1]
                            # Basura típica de Windows/Mac que no debe subirse
                            if base.startswith('.') or base.startswith('~$') or '__MACOSX' in nombre_interno:
                                continue
                            ext = base.rsplit('.', 1)[-1].lower() if '.' in base else ''
                            contenido_zip.append({
                                "ruta_interna": nombre_interno,
                                "nombre": base,
                                "ext": ext,
                                "permitido": ext in MIMES_REPOSITORIO,
                                "tam_kb": round(info.file_size / 1024, 1),
                            })
                except zipfile.BadZipFile:
                    st.error("❌ El archivo no es un ZIP válido o está dañado.")
                except Exception as e_zip:
                    st.error(f"❌ No se pudo leer el ZIP: {e_zip}")

            if contenido_zip:
                validos = [c for c in contenido_zip if c["permitido"]]
                omitidos = [c for c in contenido_zip if not c["permitido"]]

                st.success(f"📦 {len(validos)} archivo(s) listos para subir.")
                if omitidos:
                    with st.expander(f"⚠️ {len(omitidos)} archivo(s) se omitirán (tipo no permitido)"):
                        st.write(", ".join(sorted({c['nombre'] for c in omitidos})))

                if modo_zip == "Cada carpeta dentro del ZIP es un colaborador distinto":
                    carpetas_detectadas = sorted({
                        c["ruta_interna"].split('/')[0].upper()
                        for c in validos if '/' in c["ruta_interna"]
                    })
                    sueltos = [c for c in validos if '/' not in c["ruta_interna"]]
                    if carpetas_detectadas:
                        st.caption(f"👥 Colaboradores detectados: {', '.join(carpetas_detectadas)}")
                    if sueltos:
                        st.warning(
                            f"⚠️ {len(sueltos)} archivo(s) están en la raíz del ZIP, sin carpeta. "
                            "Se asignarán al colaborador escrito arriba."
                        )

                if st.button("☁️ Subir Carpeta al Repositorio", type="primary", use_container_width=True, key="repo_btn_subir_zip"):
                    necesita_nombre = (
                        modo_zip == "Todo pertenece a UN colaborador"
                        or any('/' not in c["ruta_interna"] for c in validos)
                    )
                    if necesita_nombre and not colaborador_sel:
                        st.error("⚠️ Debes indicar el colaborador.")
                    else:
                        nuevos_registros = []
                        fallidos = []
                        sello = get_honduras_time().strftime('%Y%m%d_%H%M%S')

                        archivo_zip.seek(0)
                        datos_zip = io.BytesIO(archivo_zip.read())

                        barra = st.progress(0.0, text="Procesando carpeta...")
                        with zipfile.ZipFile(datos_zip) as zf:
                            for i, item in enumerate(validos):
                                ruta_interna = item["ruta_interna"]
                                partes = ruta_interna.split('/')

                                # Determinar a qué colaborador pertenece el archivo
                                if modo_zip == "Cada carpeta dentro del ZIP es un colaborador distinto" and len(partes) > 1:
                                    dueno = partes[0].strip().upper()
                                    subruta = partes[1:]
                                else:
                                    dueno = colaborador_sel
                                    subruta = partes

                                carpeta_dueno = _sanitizar_nombre_carpeta(dueno)

                                # En Catbox no existen carpetas reales: la jerarquía
                                # se conserva en el índice (COLABORADOR + DESCRIPCION),
                                # que es lo que la app usa para agrupar y filtrar.
                                subcarpetas = [p for p in subruta[:-1]]

                                try:
                                    contenido = zf.read(ruta_interna)
                                    enlace, err = _subir_documento_repositorio(contenido, item["nombre"])
                                except Exception as e_item:
                                    enlace, err = None, str(e_item)

                                if enlace:
                                    nuevos_registros.append({
                                        "COLABORADOR": dueno,
                                        "CARPETA": carpeta_dueno,
                                        "NOMBRE_ARCHIVO": item["nombre"],
                                        "ENLACE": enlace,
                                        "TIPO": item["ext"].upper(),
                                        "TAMANO_KB": item["tam_kb"],
                                        "DESCRIPCION": descripcion_doc or ("/".join(subcarpetas) if subcarpetas else ""),
                                        "SUBIDO_POR": st.session_state.get('usuario_actual', ''),
                                        "FECHA_SUBIDA": get_honduras_time().strftime('%Y-%m-%d %H:%M:%S'),
                                    })
                                else:
                                    fallidos.append(f"{ruta_interna} ({err})")

                                barra.progress((i + 1) / len(validos), text=f"Subiendo {item['nombre']}...")

                        barra.empty()
                        _finalizar_subida(conn, df_indice, nuevos_registros, fallidos, None)

    # ======================================================================
    # EXPLORAR
    # ======================================================================
    with tab_explorar:
        col_f1, col_f2, col_f3 = st.columns([2, 2, 1])
        with col_f1:
            opciones_filtro = ["VER TODOS"] + (sorted(df_indice['COLABORADOR'].dropna().astype(str).unique().tolist()) if not df_indice.empty else [])
            filtro_colab = st.selectbox("🔍 Carpeta:", options=opciones_filtro, key="repo_filtro_colab")
        with col_f2:
            buscar_archivo = st.text_input("🔎 Buscar por nombre:", key="repo_buscar")
        with col_f3:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("🔄 Refrescar", use_container_width=True, key="repo_refrescar"):
                st.rerun()

        if df_indice.empty:
            st.info("📭 El repositorio está vacío. Sube el primer documento en la pestaña anterior.")
        else:
            df_ver = df_indice.copy()
            if filtro_colab != "VER TODOS":
                df_ver = df_ver[df_ver['COLABORADOR'].astype(str) == filtro_colab]
            if buscar_archivo.strip():
                df_ver = df_ver[df_ver['NOMBRE_ARCHIVO'].astype(str).str.contains(buscar_archivo.strip(), case=False, na=False)]

            if df_ver.empty:
                st.warning("No hay documentos que coincidan con el filtro.")
            else:
                df_ver = df_ver.sort_values(by='FECHA_SUBIDA', ascending=False)
                st.caption(f"📄 {len(df_ver)} documento(s)")

                for idx, fila in df_ver.iterrows():
                    with st.container(border=True):
                        c1, c2, c3 = st.columns([5, 1.2, 1.2])
                        with c1:
                            st.markdown(f"**{fila['NOMBRE_ARCHIVO']}**")
                            detalle = f"👤 {fila['COLABORADOR']} · 📅 {fila['FECHA_SUBIDA']} · ⬆️ {fila['SUBIDO_POR']}"
                            if str(fila.get('TAMANO_KB', '')).strip():
                                detalle += f" · {fila['TAMANO_KB']} KB"
                            st.caption(detalle)
                            if str(fila.get('DESCRIPCION', '')).strip():
                                st.caption(f"📝 {fila['DESCRIPCION']}")

                        with c2:
                            # Se descarga desde Catbox solo cuando se pide, y se
                            # restaura el nombre real: los archivos con extensión
                            # bloqueada (Word) están guardados allá como .bin.
                            if st.button("⬇️ Preparar", key=f"repo_prep_{idx}", use_container_width=True):
                                with st.spinner("Descargando..."):
                                    try:
                                        r_doc = requests.get(str(fila['ENLACE']), timeout=60)
                                        datos = r_doc.content if r_doc.status_code == 200 else None
                                    except Exception:
                                        datos = None
                                if datos:
                                    st.session_state[f"repo_datos_{idx}"] = datos
                                else:
                                    st.error("No se encontró el archivo en la nube.")

                            if st.session_state.get(f"repo_datos_{idx}"):
                                ext_f = str(fila.get('TIPO', '')).lower()
                                st.download_button(
                                    "💾 Guardar",
                                    data=st.session_state[f"repo_datos_{idx}"],
                                    file_name=fila['NOMBRE_ARCHIVO'],
                                    mime=MIMES_REPOSITORIO.get(ext_f, "application/octet-stream"),
                                    key=f"repo_dl_{idx}",
                                    use_container_width=True
                                )

                        with c3:
                            if st.session_state.get(f"repo_confirmar_{idx}"):
                                st.caption("¿Seguro?")
                                if st.button("✅ Sí", key=f"repo_si_{idx}", use_container_width=True):
                                    ok_del, err_del = _borrar_documento_repositorio(fila['ENLACE'])
                                    if ok_del:
                                        df_nuevo = df_indice[df_indice['ENLACE'] != fila['ENLACE']]
                                        _guardar_indice_repositorio(conn, df_nuevo)
                                        st.session_state[f"repo_confirmar_{idx}"] = False
                                        st.success("Eliminado.")
                                        time.sleep(1)
                                        st.rerun()
                                    else:
                                        st.error(f"No se pudo eliminar: {err_del}")
                                if st.button("✖️ No", key=f"repo_no_{idx}", use_container_width=True):
                                    st.session_state[f"repo_confirmar_{idx}"] = False
                                    st.rerun()
                            else:
                                if st.button("🗑️ Borrar", key=f"repo_del_{idx}", use_container_width=True):
                                    st.session_state[f"repo_confirmar_{idx}"] = True
                                    st.rerun()

        # ------------------------------------------------------------------
        # MANTENIMIENTO
        # ------------------------------------------------------------------
        with st.expander("🛠️ Mantenimiento del repositorio"):
            st.markdown("**Almacenamiento:** Catbox (enlaces públicos) · **Índice:** Google Sheets")
            st.caption(
                "⚠️ Los enlaces de Catbox no tienen contraseña: cualquiera que tenga la URL "
                "puede abrir el documento. La app protege la lista, no los archivos."
            )

            st.divider()
            st.caption("Verifica que todos los enlaces del índice sigan vivos en Catbox.")
            if st.button("🩺 Verificar enlaces", key="repo_verificar_enlaces"):
                if df_indice.empty:
                    st.info("No hay documentos registrados.")
                else:
                    rotos = []
                    barra_v = st.progress(0.0, text="Verificando...")
                    total_v = len(df_indice)
                    for j, (_, f_v) in enumerate(df_indice.iterrows()):
                        try:
                            r_v = requests.head(str(f_v['ENLACE']), timeout=15, allow_redirects=True)
                            if r_v.status_code >= 400:
                                rotos.append(f"{f_v['COLABORADOR']} — {f_v['NOMBRE_ARCHIVO']}")
                        except Exception:
                            rotos.append(f"{f_v['COLABORADOR']} — {f_v['NOMBRE_ARCHIVO']}")
                        barra_v.progress((j + 1) / total_v, text=f"Verificando {j + 1}/{total_v}...")
                    barra_v.empty()

                    if rotos:
                        st.error(f"❌ {len(rotos)} enlace(s) no responden:\n\n- " + "\n- ".join(rotos[:20]))
                    else:
                        st.success(f"✅ Los {total_v} enlaces están activos.")

            st.divider()
            st.caption("Si un archivo se subió pero no quedó registrado, pega aquí su enlace de Catbox para agregarlo al índice.")
            with st.form("form_rescate_enlace", clear_on_submit=True):
                col_r1, col_r2 = st.columns(2)
                with col_r1:
                    rescate_colab = st.text_input("Colaborador:", key="rescate_colab").strip().upper()
                with col_r2:
                    rescate_nombre = st.text_input("Nombre del archivo:", key="rescate_nombre", placeholder="ej: CONSTANCIA.pdf").strip()
                rescate_url = st.text_input("Enlace de Catbox:", key="rescate_url", placeholder="https://files.catbox.moe/...").strip()

                if st.form_submit_button("➕ Agregar al índice", use_container_width=True):
                    if not (rescate_colab and rescate_nombre and rescate_url.startswith("http")):
                        st.error("Completa los tres campos. El enlace debe empezar con http.")
                    else:
                        ext_r = rescate_nombre.rsplit('.', 1)[-1].upper() if '.' in rescate_nombre else ""
                        registro_r = [{
                            "COLABORADOR": rescate_colab,
                            "CARPETA": _sanitizar_nombre_carpeta(rescate_colab),
                            "NOMBRE_ARCHIVO": rescate_nombre,
                            "ENLACE": rescate_url,
                            "TIPO": ext_r,
                            "TAMANO_KB": "",
                            "DESCRIPCION": "(agregado manualmente)",
                            "SUBIDO_POR": st.session_state.get('usuario_actual', ''),
                            "FECHA_SUBIDA": get_honduras_time().strftime('%Y-%m-%d %H:%M:%S'),
                        }]
                        df_r = pd.concat([df_indice, pd.DataFrame(registro_r)], ignore_index=True)
                        if _guardar_indice_repositorio(conn, df_r):
                            st.success("✅ Registrado en el índice.")
                            time.sleep(1.2)
                            st.rerun()

            st.divider()
            st.caption("Descarga el índice completo como respaldo (los enlaces quedan incluidos).")
            if not df_indice.empty:
                st.download_button(
                    "📥 Descargar índice (CSV)",
                    data=df_indice.to_csv(index=False).encode('utf-8'),
                    file_name=f"repositorio_expedientes_{get_honduras_time().strftime('%Y%m%d')}.csv",
                    mime="text/csv",
                    key="repo_backup_indice"
                )


def mostrar_modulo_expedientes(conn, df_base):
    supervisor_actual = st.session_state.get('usuario_actual', st.session_state.get('username', 'Supervisor'))
    rol_usuario = st.session_state.get('rol_actual', 'monitoreo')
    es_admin = (str(rol_usuario).strip().lower() == 'admin')

    st.title("📁 Gestión de Expedientes y Reportes")
    
    def generar_vista_historial(df_mostrar, titulo_seccion, tab_id):
        try:
            col_tit, col_ref = st.columns([4, 1])
            with col_tit:
                st.subheader(f"📜 Historial de Expedientes ({titulo_seccion})")
            with col_ref:
                if st.button("🔄 Refrescar Datos", key=f"btn_refrescar_{tab_id}", use_container_width=True):
                    forzar_actualizacion_memoria(conn)
                    st.rerun()
            
            if not df_mostrar.empty:
                with st.container():
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        filtro_nombre = st.selectbox("🔍 Colaborador:", options=["VER TODOS"] + sorted(df_mostrar['TECNICO'].unique().tolist()), key=f"filtro_nom_{tab_id}")
                    with col2:
                        hoy = get_honduras_time().date()
                        rango_fechas = st.date_input("📅 Rango de Fechas:", value=(hoy - timedelta(days=60), hoy), key=f"filtro_fec_{tab_id}")
                    with col3:
                        filtro_tipo = st.selectbox("📋 Tipo de Registro:", options=["Todos los Tipos", "Llamado de Atención", "Incidencia Médica"], key=f"filtro_tip_{tab_id}")

                if filtro_nombre != "VER TODOS":
                    df_mostrar = df_mostrar[df_mostrar['TECNICO'] == filtro_nombre]
                
                if isinstance(rango_fechas, (list, tuple)) and len(rango_fechas) == 2:
                    fecha_inicio, fecha_fin = rango_fechas
                    def parsear_fecha(f_str):
                        for fmt in ('%d/%m/%Y', '%Y-%m-%d', '%d-%m-%Y'):
                            try: return datetime.strptime(str(f_str).strip(), fmt).date()
                            except: continue
                        return None

                    df_mostrar['FECHA_INCIDENCIA_DT'] = df_mostrar['FECHA_INCIDENCIA'].apply(parsear_fecha)
                    df_mostrar = df_mostrar[
                        df_mostrar['FECHA_INCIDENCIA_DT'].notna() &
                        (df_mostrar['FECHA_INCIDENCIA_DT'] >= fecha_inicio) & 
                        (df_mostrar['FECHA_INCIDENCIA_DT'] <= fecha_fin)
                    ]
                
                if filtro_tipo == "Incidencia Médica":
                    df_mostrar = df_mostrar[df_mostrar['TIPO_FALTA'].str.upper().isin(["INCIDENCIA MÉDICA", "INCIDENCIA MEDICA"])]
                elif filtro_tipo == "Llamado de Atención":
                    df_mostrar = df_mostrar[~df_mostrar['TIPO_FALTA'].str.upper().isin(["INCIDENCIA MÉDICA", "INCIDENCIA MEDICA"])]

                # --- PANEL DE KPIs ---
                with st.expander("📊 PANEL DE KPIs: ANALÍTICA DE PERSONAL", expanded=False):
                    if not df_mostrar.empty:
                        df_kpi = df_mostrar.copy()
                        
                        # Contar llegadas tarde para aplicar promoción por reincidencia
                        mask_tarde_kpi = df_kpi.apply(lambda r: es_llegada_tarde(r['TIPO_FALTA'], r['COMENTARIO']), axis=1)
                        conteo_tardes_kpi = (
                            df_kpi[mask_tarde_kpi]['TECNICO']
                            .astype(str).str.upper().str.strip()
                            .value_counts().to_dict()
                        )
                        
                        df_kpi['RUBRO'] = df_kpi.apply(
                            lambda r: asignar_rubro_automatico(
                                r['TIPO_FALTA'], 
                                r['COMENTARIO'], 
                                n_tardes=conteo_tardes_kpi.get(str(r['TECNICO']).upper().strip(), 0)
                            ), axis=1
                        )
                        
                        df_kpi['FECHA_DT'] = pd.to_datetime(df_kpi['FECHA_INCIDENCIA'], format='%d/%m/%Y', errors='coerce')
                        df_kpi['Día Semana'] = df_kpi['FECHA_DT'].dt.day_name()
                        dias_es = {'Monday': 'Lunes', 'Tuesday': 'Martes', 'Wednesday': 'Miércoles', 'Thursday': 'Jueves', 'Friday': 'Viernes', 'Saturday': 'Sábado', 'Sunday': 'Domingo'}
                        df_kpi['Día Semana'] = df_kpi['Día Semana'].map(dias_es)

                        tot_registros = len(df_kpi)
                        colab_unicos = df_kpi['TECNICO'].nunique() 
                        rubro_comun = df_kpi['RUBRO'].value_counts().index[0] if not df_kpi['RUBRO'].empty else "N/D"
                        
                        st.markdown(f"""
                        <div style="display: flex; gap: 10px; margin-bottom: 15px;">
                            <div style="flex: 1; background-color: #1A1D24; padding: 10px; border-radius: 6px; border: 1px solid #2D2F39; text-align: center;">
                                <span style="font-size: 10px; color: #94A3B8; text-transform: uppercase; font-weight: bold;">📦 Casos Totales</span>
                                <h3 style="margin: 2px 0 0 0; color: #FFF; font-size: 18px;">{tot_registros}</h3>
                            </div>
                            <div style="flex: 1; background-color: #1A1D24; padding: 10px; border-radius: 6px; border: 1px solid #2D2F39; text-align: center;">
                                <span style="font-size: 10px; color: #94A3B8; text-transform: uppercase; font-weight: bold;">👤 Personas Implicadas</span>
                                <h3 style="margin: 2px 0 0 0; color: #10B981; font-size: 18px;">{colab_unicos}</h3>
                            </div>
                            <div style="flex: 1; background-color: #1A1D24; padding: 10px; border-radius: 6px; border: 1px solid #2D2F39; text-align: center;">
                                <span style="font-size: 10px; color: #94A3B8; text-transform: uppercase; font-weight: bold;">🎯 Rubro Dominante</span>
                                <h3 style="margin: 2px 0 0 0; color: #F59E0B; font-size: 18px;">{rubro_comun}</h3>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)

                        c_top1, c_top2 = st.columns(2)
                        with c_top1:
                            st.markdown("<h5 style='color:#EF4444; font-size:13px; font-weight:bold;'>🚨 Reincidentes Críticos</h5>", unsafe_allow_html=True)
                            df_reinc = df_kpi.groupby(['TECNICO', 'TIPO_FALTA']).size().reset_index(name='Veces')
                            df_reinc = df_reinc[df_reinc['Veces'] > 1].sort_values(by='Veces', ascending=False)
                            if not df_reinc.empty:
                                st.dataframe(df_reinc.head(4), hide_index=True, use_container_width=True, height=180)
                            else:
                                st.success("✅ Sin reincidentes.")
                                
                        with c_top2:
                            st.markdown("<h5 style='color:#F59E0B; font-size:13px; font-weight:bold;'>📅 Faltas por Día</h5>", unsafe_allow_html=True)
                            orden_dias = ['Lunes', 'Martes', 'Miércoles', 'Jueves', 'Viernes', 'Sábado', 'Domingo']
                            df_dias = df_kpi['Día Semana'].value_counts().reindex(orden_dias, fill_value=0).reset_index()
                            df_dias.columns = ['Día', 'Cantidad']
                            fig_barras_dias = px.bar(df_dias, x='Día', y='Cantidad', template="plotly_dark", height=180, color='Cantidad', color_continuous_scale='Reds')
                            fig_barras_dias.update_layout(margin=dict(t=5, b=5, l=5, r=5), showlegend=False, coloraxis_showscale=False, xaxis_title="", yaxis_title="")
                            st.plotly_chart(fig_barras_dias, use_container_width=True)

                        c_bot1, c_bot2 = st.columns(2)
                        with c_bot1:
                            st.markdown("<h5 style='color:#3B82F6; font-size:13px; font-weight:bold;'>🎛️ Origen por Rubros</h5>", unsafe_allow_html=True)
                            df_rubros_chart = df_kpi['RUBRO'].value_counts().reset_index()
                            df_rubros_chart.columns = ['Rubro', 'Cantidad']
                            colores_rubros = {'GPS': '#EF4444', 'BIOMÉTRICO': '#3B82F6', 'CEPHEUS': '#F59E0B', 'RECO': '#10B981', 'OTROS': '#64748B'}
                            fig_rubros = px.bar(df_rubros_chart, x='Rubro', y='Cantidad', template="plotly_dark", height=180, color='Rubro', color_discrete_map=colores_rubros)
                            fig_rubros.update_layout(margin=dict(t=5, b=5, l=5, r=5), showlegend=False, xaxis_title="", yaxis_title="")
                            st.plotly_chart(fig_rubros, use_container_width=True)
                            
                        with c_bot2:
                            st.markdown("<h5 style='color:#10B981; font-size:13px; font-weight:bold;'>🚫 Tipos de Falta Específica</h5>", unsafe_allow_html=True)
                            df_motivos = df_kpi['TIPO_FALTA'].value_counts().reset_index()
                            df_motivos.columns = ['Motivo', 'Cantidad']
                            fig_pie = px.pie(df_motivos, names='Motivo', values='Cantidad', hole=0.4, template="plotly_dark", height=180)
                            fig_pie.update_traces(textposition='inside', textinfo='percent+label', textfont_size=10)
                            fig_pie.update_layout(margin=dict(t=5, b=5, l=5, r=5), showlegend=False)
                            st.plotly_chart(fig_pie, use_container_width=True)

                        st.markdown("<h4 style='color:#10B981; font-size:14px; font-weight:bold; margin-top:20px;'>📊 Resumen General de Incidencias por Rubro</h4>", unsafe_allow_html=True)
                        df_resumen_rubros = df_kpi['RUBRO'].value_counts().reset_index()
                        df_resumen_rubros.columns = ['Rubro / Tipo de Falta', 'Cantidad de Incidencias']
                        st.dataframe(df_resumen_rubros, use_container_width=True, hide_index=True)
                    else:
                        st.info("📊 No hay datos disponibles para los filtros seleccionados.")

                st.markdown("---")

                # --- 2. SECCIÓN DE DESCARGA DUAL (PDF & WORD) ---
                c_v, c_b_pdf, c_b_docx = st.columns([2, 1, 1])
                with c_b_pdf:
                    if not df_mostrar.empty:
                        if filtro_nombre == "VER TODOS":
                            nombre_archivo_pdf = "Reporte_General.pdf"
                        else:
                            nombre_corto = " ".join(filtro_nombre.split()[:2]) 
                            nombre_archivo_pdf = f"Reporte_{nombre_corto.replace(' ', '_')}.pdf"

                        id_estado_pdf = f"pdf_listo_{tab_id}_{filtro_nombre}_{len(df_mostrar)}"

                        if st.session_state.get('estado_pdf_actual') == id_estado_pdf:
                            st.download_button(
                                label="⬇️ Descargar PDF",
                                data=st.session_state['pdf_bytes_listo'],
                                file_name=nombre_archivo_pdf,
                                mime="application/pdf",
                                use_container_width=True,
                                type="primary"
                            )
                        else:
                            if st.button("📄 Preparar PDF", key=f"btn_pdf_{tab_id}", use_container_width=True):
                                with st.spinner("Generando PDF..."):
                                    st.session_state['pdf_bytes_listo'] = generar_pdf_consolidado(df_mostrar)
                                    st.session_state['estado_pdf_actual'] = id_estado_pdf
                                    st.rerun()

                with c_b_docx:
                    if not df_mostrar.empty:
                        if HAS_DOCX:
                            if filtro_nombre == "VER TODOS":
                                nombre_archivo_docx = "Reporte_General.docx"
                            else:
                                nombre_corto = " ".join(filtro_nombre.split()[:2]) 
                                nombre_archivo_docx = f"Reporte_{nombre_corto.replace(' ', '_')}.docx"

                            id_estado_docx = f"docx_listo_{tab_id}_{filtro_nombre}_{len(df_mostrar)}"

                            if st.session_state.get('estado_docx_actual') == id_estado_docx:
                                st.download_button(
                                    label="⬇️ Descargar Word",
                                    data=st.session_state['docx_bytes_listo'],
                                    file_name=nombre_archivo_docx,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True,
                                    type="primary"
                                )
                            else:
                                if st.button("📝 Preparar Word", key=f"btn_docx_{tab_id}", use_container_width=True):
                                    with st.spinner("Generando Word (.docx)..."):
                                        st.session_state['docx_bytes_listo'] = generar_docx_consolidado(df_mostrar)
                                        st.session_state['estado_docx_actual'] = id_estado_docx
                                        st.rerun()
                        else:
                            st.button("📝 Word Desactivado", disabled=True, help="Agrega 'python-docx' a requirements.txt para habilitar descargas en Word", use_container_width=True, key=f"btn_docx_disabled_{tab_id}")

                # --- 3. DATAFRAME DE EXPEDIENTES ---
                if df_mostrar.empty:
                    st.info("💡 No hay registros para los filtros seleccionados.")
                else:
                    df_tabla = df_mostrar[['FECHA_INCIDENCIA', 'TECNICO', 'TIPO_FALTA', 'COMENTARIO', 'SUPERVISOR']].copy()
                    df_tabla.columns = ['Fecha', 'Colaborador', 'Motivo', 'Descripción', 'Registrado por']
                    
                    st.dataframe(
                        df_tabla.iloc[::-1],
                        hide_index=True, 
                        use_container_width=True,
                        height=250
                    )
                    
                    st.markdown("<br>", unsafe_allow_html=True)

                    # --- 4. ACCIONES SOBRE EL REGISTRO SELECCIONADO ---
                    with st.expander("🛠️ ACCIONES: Ver o Eliminar Registro", expanded=False):
                        st.write("Seleccione un registro de la lista para gestionarlo:")
                        
                        dict_acciones = {}
                        lista_opciones = ["--- Seleccione un registro ---"]
                        
                        for idx, row in df_mostrar.iloc[::-1].iterrows():
                            label = f"{row['FECHA_INCIDENCIA']} | {row['TECNICO']} | {row['TIPO_FALTA']}"
                            lista_opciones.append(label)
                            dict_acciones[label] = (idx, row)
                            
                        registro_sel = st.selectbox("Registro a gestionar:", options=lista_opciones, label_visibility="collapsed", key=f"sel_gestion_{tab_id}")
                        
                        if registro_sel != "--- Seleccione un registro ---":
                            idx_sel, row_sel = dict_acciones[registro_sel]
                            
                            st.markdown("---")
                            st.markdown(f"### 📋 Detalles del Expediente: {row_sel['TECNICO']}")
                            st.info(f"**📝 Descripción:**\n\n{row_sel['COMENTARIO']}")
                            st.write(f"**👤 Registrado por:** {row_sel['SUPERVISOR']}  |  **🕒 Fecha de registro:** {row_sel['FECHA_REGISTRO']}")
                            
                            urls = str(row_sel.get('URL_FOTO', '')).split(',')
                            validas = [u.strip() for u in urls if u.strip().startswith('http')]
                            if validas:
                                st.markdown("#### 🖼️ Evidencias Fotográficas Adjuntas:")
                                cols_img = st.columns(len(validas))
                                for i, u in enumerate(validas):
                                    with cols_img[i]:
                                        if u.lower().endswith('.pdf') or 'catbox' in u.lower() and '.pdf' in u.lower():
                                            st.info(f"📄 **Documento PDF adjunto:** [Haga clic aquí para ver o descargar]({u})")
                                        else:
                                            st.image(u, caption="Evidencia Fotográfica", use_container_width=True)
                            else:
                                st.caption("🚫 No se adjuntaron evidencias.")
                                
                            st.markdown("<br>", unsafe_allow_html=True)
                            if es_admin:
                                if st.button("🗑️ Eliminar Registro de Base de Datos", key=f"del_btn_{idx_sel}_{tab_id}", type="primary", use_container_width=True):
                                    with st.spinner("Eliminando..."):
                                        try:
                                            df_borrado = obtener_datos_memoria(conn)
                                            if idx_sel in df_borrado.index:
                                                df_borrado = df_borrado.drop(idx_sel).reset_index(drop=True)
                                                sobrescribir_archivo_gcs(df_borrado, NOMBRE_BUCKET_SISTEMA, "expedientes_maestro.csv")
                                                conn.update(spreadsheet=st.secrets["url_base_datos"], worksheet="Expedientes", data=df_borrado)
                                            st.success("✅ Eliminado.")
                                            forzar_actualizacion_memoria(conn)
                                            time.sleep(1)
                                            st.rerun()
                                        except Exception as e:
                                            st.error(f"Error al eliminar: {e}")
                            else:
                                st.button("🚫 Eliminar (Solo Admin)", disabled=True, use_container_width=True, key=f"del_block_{idx_sel}_{tab_id}")
            else:
                st.info("No hay registros en la base de datos para esta área.")
        except Exception as e:
            st.warning(f"⚠️ Error al cargar el historial: {e}")

    # ==========================================================================
    # DEFINICIÓN DE PESTAÑAS (Ubicado fuera de generar_vista_historial)
    # ==========================================================================
    # La pestaña del repositorio solo se muestra a los usuarios autorizados, para
    # no exponer siquiera su existencia al resto del personal.
    _usuario_repo = str(st.session_state.get('usuario_actual', '')).strip().lower()
    if _usuario_repo in USUARIOS_REPOSITORIO:
        tab_tecnicos, tab_admin, tab_repositorio = st.tabs([
            "⚙️ Operaciones (Técnicos y Auxiliares)",
            "🏢 Administrativo (SAC, Ventas, etc.)",
            "🗄️ Repositorio de Documentos"
        ])
    else:
        tab_tecnicos, tab_admin = st.tabs([
            "⚙️ Operaciones (Técnicos y Auxiliares)",
            "🏢 Administrativo (SAC, Ventas, etc.)"
        ])
        tab_repositorio = None
    
    # ==========================================================================
    # PESTAÑA 1: OPERACIONES (TÉCNICOS)
    # ==========================================================================
    with tab_tecnicos:
        with st.expander("➕ Crear Nuevo Registro - Operaciones", expanded=True):
            st.info(f"✍️ Supervisor registrando: **{supervisor_actual}**")
            
            c1, c2 = st.columns(2)
            with c1:
                lista_nombres = cargar_personal("personal_tecnico.txt")
                
                tipo_falta_base = st.selectbox("🚫 Motivo:", [
                    "Exceso de Velocidad", "Llegada Tarde", "Abandono de Ruta", 
                    "Mala Documentación", "Incidencia Médica", "Reco / Cambio de Postes", "Otro"
                ], key="sel_falta")
                
                if tipo_falta_base == "Reco / Cambio de Postes":
                    tipo_falta = "RECO / CAMBIO DE POSTES"
                    opciones_tecnicos = ["RECO"]
                    idx_defecto = 0
                    disabled_tec = True
                else:
                    tipo_falta = tipo_falta_base
                    if tipo_falta_base == "Otro":
                        motivo_especifico = st.text_input("📝 Especifique el motivo de la falta:", key="txt_motivo_otro")
                        tipo_falta = motivo_especifico.strip().upper()
                    opciones_tecnicos = ["---"] + lista_nombres
                    idx_defecto = 0
                    disabled_tec = False
                    
                colaborador_sel = st.selectbox("👤 Colaborador:", options=opciones_tecnicos, index=idx_defecto, disabled=disabled_tec, key="sel_colab")
                
            with c2:
                fecha_inc = st.date_input("📅 Fecha:", value=get_honduras_time().date(), key="date_inc")
                archivos = st.file_uploader("🖼️ Evidencias:", type=['png', 'jpg', 'jpeg', 'pdf'], accept_multiple_files=True, key="up_archivos")
            
            comentario = st.text_area("📝 Descripción de los hechos:", key="txt_comentario")
            
            if st.button("💾 GUARDAR EN EXPEDIENTE OPERATIVO", type="primary", use_container_width=True):
                if colaborador_sel == "---" or not comentario:
                    st.error("⚠️ Complete el nombre y el comentario.")
                elif tipo_falta_base == "Otro" and not tipo_falta:
                    st.error("⚠️ Por favor, especifique el motivo de la falta en el campo correspondiente.")
                else:
                    try:
                        urls = []
                        if archivos:
                            with st.spinner("Subiendo archivos al servidor..."):
                                for a in archivos:
                                    url_subido = subir_evidencias_inteligente(a)
                                    if url_subido:
                                        urls.append(url_subido)
                        
                        with st.spinner("Guardando en la Nube y en Sheets..."):
                            df_actual = obtener_datos_memoria(conn)
                                
                            cols_exp = ['FECHA_REGISTRO', 'TECNICO', 'TIPO_FALTA', 'FECHA_INCIDENCIA', 'COMENTARIO', 'URL_FOTO', 'SUPERVISOR']
                            
                            nueva_fila = [
                                get_honduras_time().strftime("%d/%m/%Y %H:%M:%S"),
                                colaborador_sel,
                                tipo_falta,
                                fecha_inc.strftime("%d/%m/%Y"),
                                comentario,
                                ", ".join(urls),
                                supervisor_actual
                            ]
                            nuevo_df = pd.DataFrame([nueva_fila], columns=cols_exp)
                            
                            if df_actual is not None and not df_actual.empty:
                                if len(df_actual.columns) > len(cols_exp):
                                    df_actual = df_actual.iloc[:, :len(cols_exp)]
                                elif len(df_actual.columns) < len(cols_exp):
                                    for i in range(len(cols_exp) - len(df_actual.columns)):
                                        df_actual[f"Columna_Recuperada_{i}"] = ""
                                        
                                df_actual.columns = cols_exp
                                df_final = pd.concat([df_actual, nuevo_df], ignore_index=True)
                            else:
                                df_final = nuevo_df
                                
                            sobrescribir_archivo_gcs(df_final, NOMBRE_BUCKET_SISTEMA, "expedientes_maestro.csv")
                            conn.update(spreadsheet=st.secrets["url_base_datos"], worksheet="Expedientes", data=df_final)

                        st.success(f"✅ ¡Guardado exitosamente! {colaborador_sel} registrado en la base de datos.")
                        forzar_actualizacion_memoria(conn)
                        time.sleep(1.5)
                        
                        llaves_a_borrar = ["sel_colab", "sel_falta", "date_inc", "up_archivos", "txt_comentario", "txt_motivo_otro"]
                        for llave in llaves_a_borrar:
                            if llave in st.session_state:
                                del st.session_state[llave]
                        
                        st.rerun()

                    except Exception as e:
                        st.error(f"❌ Error al intentar escribir en la base de datos: {e}")

            st.markdown("---")
            
            df_view = obtener_datos_memoria(conn)
            if df_view is not None and not df_view.empty and 'TECNICO' in df_view.columns:
                df_view['TECNICO'] = df_view['TECNICO'].astype(str).str.upper().str.strip()
                df_view['TECNICO'] = df_view['TECNICO'].replace(r'\s+', ' ', regex=True)
                es_admin_mask = df_view['TECNICO'].str.contains(r'\(.*\)$', regex=True, na=False)
                df_tecnicos_tab = df_view[~es_admin_mask].copy()
                df_tecnicos_tab = df_tecnicos_tab[~df_tecnicos_tab['TECNICO'].isin(['', 'NAN', 'NONE', 'NULL', 'NAT', 'UNDEFINED'])]
                
                generar_vista_historial(df_tecnicos_tab, "Operaciones y Técnicos", "tec")

        # ==========================================================================
        # PESTAÑA 2: ADMINISTRATIVO (MÓDULO SAC Y VENTAS)
        # ==========================================================================
        with tab_admin:
            with st.expander("➕ Crear Nuevo Registro - Administrativo", expanded=True):
                st.info(f"✍️ Supervisor registrando: **{supervisor_actual}**")
                
                c1_admin, c2_admin = st.columns(2)
                with c1_admin:
                    dict_admin = cargar_personal_admin("personal_sac.txt")
                    opciones_dept = ["--- Seleccione ---"] + list(dict_admin.keys()) if dict_admin else ["--- Seleccione ---"]
                    
                    dept_sel = st.selectbox("🏢 Área / Departamento:", opciones_dept, key="sel_dept_admin")
                    
                    opciones_nombres_admin = ["---"]
                    if dept_sel != "--- Seleccione ---":
                        opciones_nombres_admin += dict_admin.get(dept_sel, [])
                        
                    if len(opciones_nombres_admin) <= 1:
                        nombre_admin = st.text_input("👤 Nombre Completo del Colaborador:*", key="txt_nombre_admin").upper()
                    else:
                        nombre_admin = st.selectbox("👤 Colaborador:", opciones_nombres_admin, key="sel_nombre_admin")
                        
                    tipo_falta_admin_base = st.selectbox("📄 Tipo de Registro / Incidencia:", [
                        "Llamado de Atención Verbal", "Amonestación Escrita", 
                        "Llegada Tardía / Ausencia", "Incidencia Médica", 
                        "Felicitación / Mérito", "Curriculum / Contrato", "Otro"
                    ], key="sel_falta_admin")
                    
                    if tipo_falta_admin_base == "Otro":
                        tipo_falta_admin = st.text_input("📝 Especifique la incidencia:", key="txt_otro_admin").upper()
                    else:
                        tipo_falta_admin = tipo_falta_admin_base.upper()
                        
                with c2_admin:
                    fecha_inc_admin = st.date_input("📅 Fecha del Evento:", value=get_honduras_time().date(), key="date_inc_admin")
                    archivos_admin = st.file_uploader("🖼️ Evidencias Fotográficas:", type=['png', 'jpg', 'jpeg', 'pdf'], accept_multiple_files=True, key="up_archivos_admin")
                    
                comentario_admin = st.text_area("📝 Descripción de los hechos o detalles:", key="txt_comentario_admin")
                
                if st.button("💾 REGISTRAR INCIDENCIA ADMINISTRATIVA", type="primary", use_container_width=True):
                    if dept_sel == "--- Seleccione ---":
                        st.error("⚠️ Debes seleccionar el departamento.")
                    elif not nombre_admin or nombre_admin == "---":
                        st.error("⚠️ El nombre del empleado es obligatorio.")
                    elif not comentario_admin:
                        st.error("⚠️ Ingresa una descripción de los hechos.")
                    elif tipo_falta_admin_base == "Otro" and not tipo_falta_admin:
                        st.error("⚠️ Especifique el motivo en el campo 'Otro'.")
                    else:
                        try:
                            urls_admin = []
                            if archivos_admin:
                                with st.spinner("Subiendo archivos al servidor..."):
                                    for a in archivos_admin:
                                        url_subido = subir_evidencias_inteligente(a)
                                        if url_subido:
                                            urls_admin.append(url_subido)
                            
                            with st.spinner("Guardando en la base de datos principal..."):
                                df_actual = obtener_datos_memoria(conn)
                                cols_exp = ['FECHA_REGISTRO', 'TECNICO', 'TIPO_FALTA', 'FECHA_INCIDENCIA', 'COMENTARIO', 'URL_FOTO', 'SUPERVISOR']
                                nombre_etiquetado = f"{nombre_admin} ({dept_sel})"
                                
                                nueva_fila = [
                                    get_honduras_time().strftime("%d/%m/%Y %H:%M:%S"),
                                    nombre_etiquetado,
                                    tipo_falta_admin,
                                    fecha_inc_admin.strftime("%d/%m/%Y"),
                                    comentario_admin,
                                    ", ".join(urls_admin),
                                    supervisor_actual
                                ]
                                nuevo_df = pd.DataFrame([nueva_fila], columns=cols_exp)
                                
                                if df_actual is not None and not df_actual.empty:
                                    if len(df_actual.columns) > len(cols_exp):
                                        df_actual = df_actual.iloc[:, :len(cols_exp)]
                                    elif len(df_actual.columns) < len(cols_exp):
                                        for i in range(len(cols_exp) - len(df_actual.columns)):
                                            df_actual[f"Columna_Recuperada_{i}"] = ""
                                            
                                    df_actual.columns = cols_exp
                                    df_final = pd.concat([df_actual, nuevo_df], ignore_index=True)
                                else:
                                    df_final = nuevo_df
                                    
                                sobrescribir_archivo_gcs(df_final, NOMBRE_BUCKET_SISTEMA, "expedientes_maestro.csv")
                                conn.update(spreadsheet=st.secrets["url_base_datos"], worksheet="Expedientes", data=df_final)

                            st.success(f"✅ ¡Guardado exitosamente! {nombre_admin} registrado en la base de datos.")
                            forzar_actualizacion_memoria(conn)
                            time.sleep(1.5)
                            
                            llaves_a_borrar = ["sel_dept_admin", "sel_nombre_admin", "txt_nombre_admin", "sel_falta_admin", "date_inc_admin", "up_archivos_admin", "txt_comentario_admin", "txt_otro_admin"]
                            for llave in llaves_a_borrar:
                                if llave in st.session_state:
                                    del st.session_state[llave]
                            
                            st.rerun()

                        except Exception as e:
                            st.error(f"❌ Error al intentar escribir en la base de datos: {e}")

            st.markdown("---")
            
            df_view_admin = obtener_datos_memoria(conn)
            if df_view_admin is not None and not df_view_admin.empty and 'TECNICO' in df_view_admin.columns:
                df_view_admin['TECNICO'] = df_view_admin['TECNICO'].astype(str).str.upper().str.strip()
                df_view_admin['TECNICO'] = df_view_admin['TECNICO'].replace(r'\s+', ' ', regex=True)
                es_admin_mask = df_view_admin['TECNICO'].str.contains(r'\(.*\)$', regex=True, na=False)
                df_admin_tab = df_view_admin[es_admin_mask].copy()
                df_admin_tab = df_admin_tab[~df_admin_tab['TECNICO'].isin(['', 'NAN', 'NONE', 'NULL', 'NAT', 'UNDEFINED'])]
                
                generar_vista_historial(df_admin_tab, "Administración, SAC y Ventas", "adm")

    # ==========================================================================
    # PESTAÑA 3: REPOSITORIO DOCUMENTAL (solo usuarios autorizados)
    # ==========================================================================
    if tab_repositorio is not None:
        with tab_repositorio:
            mostrar_repositorio_documentos(conn)
