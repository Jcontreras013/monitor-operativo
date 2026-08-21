import pandas as pd
import re
from fpdf import FPDF
from datetime import datetime, timedelta, time as dt_time
import unicodedata
import tempfile
import os
import numpy as np
import streamlit as st
import io
import json
from typing import Any
from google.cloud import storage
from google.oauth2 import service_account
import requests
import urllib3
from requests.auth import HTTPBasicAuth

# Desactivar advertencias de certificados locales/autofirmados de red
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# ==============================================================================
# MOTOR SEGURO DE FECHAS, ZONA HORARIA Y UTILIDADES BASE
# ==============================================================================
def safestr(texto: Any) -> str:
    """Sanitizador CRÍTICO: Previene corrupción de PDFs eliminando caracteres especiales."""
    if pd.isna(texto):
        return ""
    return unicodedata.normalize('NFKD', str(texto)).encode('ascii', 'ignore').decode('ascii')

def get_honduras_time() -> datetime:
    """Ajusta la hora a UTC-6 internamente."""
    return datetime.utcnow() - timedelta(hours=6)

# Alias para mantener compatibilidad con las funciones de auditoría
get_hn_time = get_honduras_time

ALIAS_TECNICOS_CONOCIDOS = {
    "JOSUE MIGUEL SAUCEDA": "JOSE MIGUEL SAUCEDA",
    "JERMY MODESTO PADILLA": "JERMI MODESTO PADILLA",
    "JEREMY MODESTO PADILLA": "JERMI MODESTO PADILLA",
    "JERMY MODESTO PADILLA CARDONA": "JERMI MODESTO PADILLA CARDONA",
    "JEREMY MODESTO PADILLA CARDONA": "JERMI MODESTO PADILLA CARDONA",
    "JERNY MODESTO PADILLA CARDONA": "JERMI MODESTO PADILLA CARDONA",
    "JERNY MODESTO PADILLA": "JERMI MODESTO PADILLA",
    "ELIAS MIZAEL SABILLON": "ELIAS MISAEL ALONZO SABILLON",
    "ELIAS MISAEL SABILLON": "ELIAS MISAEL ALONZO SABILLON",
    "ELIAS MISAEL ALONZO": "ELIAS MISAEL ALONZO SABILLON",
    "DANIEL EZEQUIEL PONCE GUZMAN": "DANIEL EZEQUIEL GUZMAN PONCE",
    "NELAON RAMON FERRUFINO LEON": "NELSON RAMON FERRUFINO LEON"
}

def normalizar_nombre_cruce(texto):
    """
    Normaliza texto eliminando acentos, caracteres invisibles (como el Zero-Width Non-Joiner)
    y espacios extra para asegurar un cruce de nombres óptimo, y corrige alias conocidos
    de técnicos cuyo nombre varía entre fuentes de datos.
    """
    if pd.isnull(texto):
        return ""
    t = str(texto).upper().strip()
    # Limpieza de caracteres invisibles típicos en archivos txt/copias de portales
    t = t.replace('\u200c', '').replace('\u200b', '')
    t = ''.join(c for c in unicodedata.normalize('NFD', t) if unicodedata.category(c) != 'Mn')
    t = ' '.join(t.split())
    return ALIAS_TECNICOS_CONOCIDOS.get(t, t)

# ==============================================================================
# REEMPLAZAR ESTAS DOS FUNCIONES EN TOOLS.PY
# ==============================================================================

def parse_date_ultra_safe(val: Any, permitir_asumir_hoy: bool = True) -> pd.Timestamp:
    """Motor de conversión de fechas a partir de múltiples formatos entrantes."""
    if pd.isnull(val) or str(val).strip() == "" or str(val).upper() in ["NONE", "NAN", "NAT", "NULL"]:
        return pd.NaT
    
    hoy = pd.Timestamp(get_honduras_time()).normalize()

    if isinstance(val, dt_time):
        return pd.Timestamp.combine(hoy.date(), val) if permitir_asumir_hoy else pd.NaT
    if isinstance(val, datetime):
        if val.year <= 1970:
            return (hoy + pd.Timedelta(hours=val.hour, minutes=val.minute, seconds=val.second)) if permitir_asumir_hoy else pd.NaT
        return pd.Timestamp(val)
    if isinstance(val, (int, float)):
        if val == 0 or val == 0.0: return pd.NaT
        if val > 10000: return pd.to_datetime(val, unit='D', origin='1899-12-30')
        elif 0 < val < 1: return (hoy + pd.to_timedelta(val, unit='D')) if permitir_asumir_hoy else pd.NaT
        else: return pd.NaT

    # === LA CURA: Limpieza de formatos MaxCom (AM/PM) ===
    str_val = str(val).strip()
    str_val = re.sub(r'(?i)a\.?\s*m\.?', 'AM', str_val)
    str_val = re.sub(r'(?i)p\.?\s*m\.?', 'PM', str_val)

    try:
        if re.match(r'^\d{1,2}:\d{2}(:\d{2})?\s*(AM|PM)?$', str_val, re.I):
            if not permitir_asumir_hoy:
                return pd.NaT
            parsed_time = pd.to_datetime(str_val).time()
            return pd.Timestamp.combine(hoy.date(), parsed_time)

        if re.match(r'^\d{4}-\d{2}-\d{2}', str_val):
            parsed = pd.to_datetime(str_val, errors='coerce')
        else:
            parsed = pd.to_datetime(str_val, dayfirst=True, errors='coerce')

        if pd.notnull(parsed):
            if parsed.year <= 1970:
                return (hoy + pd.Timedelta(hours=parsed.hour, minutes=parsed.minute, seconds=parsed.second)) if permitir_asumir_hoy else pd.NaT
            return parsed
        return pd.NaT
    except:
        return pd.NaT

def procesar_fechas_seguro(df_input: pd.DataFrame, columnas: list, columnas_sin_asumir_hoy: list = None) -> pd.DataFrame:
    df = df_input.copy()
    columnas_sin_asumir_hoy = columnas_sin_asumir_hoy or []
    for col in columnas:
        if col in df.columns:
            permitir_hoy = col not in columnas_sin_asumir_hoy
            df[col] = df[col].apply(lambda v: parse_date_ultra_safe(v, permitir_asumir_hoy=permitir_hoy))
            df[col] = pd.to_datetime(df[col], errors='coerce')
    return df

# ==============================================================================
# 1. MAPEO UNIVERSAL DE COLUMNAS (Soporta archivos Cepheus y formatos de la API)
# ==============================================================================
COLUMNS_MAPPING = {
    'HORA_INI': ['HORA ENTRADA', 'HORA INICIO', 'HORAINICIOORDEN', 'HORA INICIO ORDEN', 'FECHA ENTRADA', 'INICIO', 'FECHAENTRADA'],
    'HORA_LIQ': ['HORA LIQUIDADO', 'HORA CIERRE', 'HORACIERREORDEN', 'HORA CIERRE ORDEN', 'FECHA LIQUIDADO', 'LIQUIDADO', 'FECHALIQUIDADO'],
    'TECNICO': ['TÉCNICO', 'TECNICO', 'OPERADOR', 'USER NAME'],
    'ACTIVIDAD': ['NOMBRE ACTIVIDAD', 'TIPO ORDEN', 'ACTIVIDAD', 'NOMACTIVIDAD'],
    'FECHA_APE': ['FECHA APERTURA', 'APERTURA', 'DIASASIGNADA', 'Días', 'FECHAAPERTURA'],
    'ESTADO': ['ESTADO', 'STATUS'],
    'SECTOR': ['SECTOR', 'Sect', 'Sector', 'CIUDAD', 'Ciudad', 'Zona'],
    'COLONIA': ['COLONIA', 'BARRIO', 'DIRECCION', 'LOCALIDAD'],
    'NUM': ['NUM', 'IDORDEN', 'NÚMERO'],
    'CLIENTE': ['CLIENTE', 'CUENTA', 'NO. CLIENTE'], 
    'NOMBRE': ['NOMBRE CLIENTE', 'SUSCRIPTOR', 'NOMBRE', 'NOMBRECLIENTE'], 
    'COMENTARIO': ['COMENTARIO', 'OBSERVACIONES'],
    'MX': ['MX', 'VEHICULO', 'UNIDAD'],
    'GPS': ['GPS', 'UBICACION', 'LINK', 'COORDENADAS'],
    # Infraestructura de red (FTTH). Permite ubicar en qué nodo se concentran
    # los cortes y atenuaciones, en vez de verlos como incidentes sueltos.
    'OLT': ['OLT', 'NOMBRE OLT', 'NOMBREOLT', 'EQUIPO OLT', 'OLT NAME'],
    'PON': ['PON', 'PUERTO PON', 'PUERTOPON', 'PON PORT', 'PUERTO', 'SPLITTER', 'TARJETA/PUERTO']
}

COLUMNAS_VITALES_SISTEMA = [
    'HORA_INI', 'HORA_LIQ', 'TECNICO', 'ACTIVIDAD', 'FECHA_APE',
    'ESTADO', 'SECTOR', 'COLONIA', 'NUM', 'CLIENTE', 'NOMBRE', 'COMENTARIO', 'MX', 'GPS',
    'OLT', 'PON'
]

PATRON_ASIGNADAS_VIVA_STR = 'PENDIENTE|INICIADA|PROCESO|ASIGNADA|DESPACHO|RUTA|SITIO|VIAJANDO|CAMINO|LLEGADA'

# Fuente única de verdad para el nombre del bucket de GCS. Antes estaba escrito
# a mano en 8 lugares distintos repartidos entre app.py, auditorv.py,
# expediente.py, tiempot.py, sync_job.py y 4 funciones internas de este mismo
# archivo -- si el bucket cambiaba algún día, había que acordarse de tocar los
# 8. Ahora todos importan esta constante.
NOMBRE_BUCKET_SISTEMA = "jovial-trilogy-306216.appspot.com"

# ==============================================================================
# IDENTIDAD DE MARCA PARA REPORTES (logo y pie de membrete)
# ==============================================================================
# Fuente ÚNICA de la marca. Antes la ruta del logo estaba repetida en 5 lugares
# distintos (tools.py, expediente.py, auditorv.py), así que cambiar el logo
# obligaba a editar archivo por archivo. Ahora se edita solo aquí.

# Rutas candidatas, en orden de preferencia. Se usa la primera que exista, de
# modo que si el archivo se renombra o falta, el reporte no sale sin logo.
LOGOS_REPORTE = [
    "image/LogoMaxcom-4.png",
    "image/LogoMaxcom-3.png",
    "image/LogoMaxcom-1.png",
    "logo.png",
]

# Membrete oficial. De él se recorta la franja del pie (iconos de redes,
# teléfonos, web, correo y código QR) para reproducirla EXACTAMENTE en los
# reportes, en vez de intentar redibujar los iconos con texto.
RUTA_MEMBRETE = "image/MEMBRETE-MAXCOM.png"

# Datos tomados del membrete oficial (MEMBRETE-MAXCOM).
MARCA_TAGLINE = "connecting your world"
MARCA_REDES = "@maxcomhn"
MARCA_TELEFONOS = "+504 2454-7070 / 8730-2675"
MARCA_WEB = "www.maxcom.hn"
MARCA_CORREO = "info@maxcom.hn"
MARCA_PIE = f"{MARCA_REDES}   |   {MARCA_TELEFONOS}   |   {MARCA_WEB}   |   {MARCA_CORREO}"

_LOGO_RECORTADO_CACHE = {}


def obtener_logo_reporte():
    """
    Devuelve la ruta de un logo listo para insertar en un PDF/DOCX.

    Los PNG de marca vienen sobre un lienzo enorme con mucho margen vacío: en
    LogoMaxcom-1 la figura ocupa apenas un 11% del ancho. Si se insertara tal
    cual con un ancho de 35 mm, la marca visible mediría ~4 mm y el reporte se
    vería roto. Por eso se recorta el margen transparente/blanco una sola vez y
    se guarda el resultado en un archivo temporal reutilizable.

    Si Pillow no está disponible, se devuelve el archivo original sin recortar:
    el logo saldrá pequeño, pero el reporte no falla.
    """
    ruta_origen = next((p for p in LOGOS_REPORTE if os.path.exists(p)), None)
    if not ruta_origen:
        return None

    if ruta_origen in _LOGO_RECORTADO_CACHE:
        cacheado = _LOGO_RECORTADO_CACHE[ruta_origen]
        if cacheado and os.path.exists(cacheado):
            return cacheado

    try:
        from PIL import Image
        import numpy as _np

        img = Image.open(ruta_origen).convert("RGBA")
        arr = _np.array(img)
        rgb, alfa = arr[..., :3], arr[..., 3]
        # Se considera "visible" lo que no es transparente ni casi blanco.
        visible = (alfa > 20) & (rgb.min(axis=-1) < 240)
        ys, xs = _np.where(visible)

        if len(xs) == 0:
            _LOGO_RECORTADO_CACHE[ruta_origen] = ruta_origen
            return ruta_origen

        recorte = img.crop((int(xs.min()), int(ys.min()), int(xs.max()) + 1, int(ys.max()) + 1))

        # Fondo blanco: los PDF no manejan bien la transparencia del PNG.
        fondo = Image.new("RGBA", recorte.size, (255, 255, 255, 255))
        fondo.alpha_composite(recorte)

        destino = os.path.join(tempfile.gettempdir(), "maxcom_logo_reporte.png")
        fondo.convert("RGB").save(destino, "PNG")
        _LOGO_RECORTADO_CACHE[ruta_origen] = destino
        return destino
    except Exception:
        _LOGO_RECORTADO_CACHE[ruta_origen] = ruta_origen
        return ruta_origen


def medidas_logo_en_caja(ruta_logo, ancho_max=34.0, alto_max=9.0):
    """
    Calcula (ancho, alto) en mm para que el logo quepa dentro de una caja, sin
    deformarse. Es necesario porque los logos de marca tienen proporciones muy
    distintas: el isotipo es cuadrado (1:1) y el logotipo completo es muy
    horizontal (~7.8:1). Fijar solo el ancho deformaría el encabezado con uno,
    y fijar solo el alto haría que el otro invadiera el título.
    """
    try:
        from PIL import Image
        with Image.open(ruta_logo) as im:
            w_px, h_px = im.size
        if not w_px or not h_px:
            return ancho_max, 0
        proporcion = w_px / h_px
        alto = alto_max
        ancho = alto * proporcion
        if ancho > ancho_max:
            ancho = ancho_max
            alto = ancho / proporcion
        return round(ancho, 2), round(alto, 2)
    except Exception:
        # Sin Pillow: se fija el alto y se deja que FPDF calcule el ancho.
        return 0, alto_max


_PIE_IMAGEN_CACHE = {}


def obtener_pie_membrete_imagen():
    """
    Recorta del membrete oficial la franja inferior (iconos de redes sociales,
    teléfonos, sitio web, correo y código QR) y la devuelve como archivo listo
    para insertar en un PDF.

    Se recorta de la imagen real en vez de redibujar los iconos con texto,
    porque así el pie queda idéntico al membrete institucional, incluido el QR.
    Devuelve (ruta, proporcion_ancho_alto) o (None, None) si no se puede.
    """
    if "pie" in _PIE_IMAGEN_CACHE:
        ruta, prop = _PIE_IMAGEN_CACHE["pie"]
        if ruta and os.path.exists(ruta):
            return ruta, prop

    if not os.path.exists(RUTA_MEMBRETE):
        return None, None

    try:
        from PIL import Image
        import numpy as _np

        img = Image.open(RUTA_MEMBRETE).convert("RGB")
        arr = _np.array(img)
        alto_total = arr.shape[0]

        # La franja del pie es la última banda con contenido de la hoja. Se
        # localiza dinámicamente (no con coordenadas fijas) para que siga
        # funcionando si el membrete se reemplaza por otra versión.
        filas_con_tinta = _np.where((arr.min(axis=2) < 240).any(axis=1))[0]
        if len(filas_con_tinta) == 0:
            return None, None

        y_fin = int(filas_con_tinta.max())
        y_ini = y_fin
        for y in range(y_fin, 0, -1):
            if not (arr[y].min(axis=1) < 240).any():
                # Se corta al encontrar 40 px seguidos en blanco hacia arriba.
                if all(not (arr[yy].min(axis=1) < 240).any() for yy in range(max(0, y - 40), y)):
                    y_ini = y
                    break
        if y_ini >= y_fin:
            y_ini = max(0, y_fin - int(alto_total * 0.05))

        banda = arr[y_ini:y_fin + 1]
        cols = _np.where((banda.min(axis=2) < 240).any(axis=0))[0]
        x_ini, x_fin = int(cols.min()), int(cols.max())

        margen = 10
        recorte = img.crop((
            max(0, x_ini - margen), max(0, y_ini - margen),
            min(img.width, x_fin + margen), min(img.height, y_fin + margen)
        ))

        destino = os.path.join(tempfile.gettempdir(), "maxcom_pie_membrete.png")
        recorte.save(destino, "PNG")
        proporcion = recorte.width / recorte.height
        _PIE_IMAGEN_CACHE["pie"] = (destino, proporcion)
        return destino, proporcion
    except Exception:
        _PIE_IMAGEN_CACHE["pie"] = (None, None)
        return None, None


def dibujar_pie_membrete(pdf, mostrar_pagina=True):
    """
    Dibuja el pie institucional en cualquier PDF de FPDF. Usa la franja recortada
    del membrete oficial (con iconos y QR); si no está disponible, cae a una
    versión en texto para que el reporte nunca quede sin datos de contacto.
    """
    try:
        ruta_pie, proporcion = obtener_pie_membrete_imagen()
        ancho_util = pdf.w - 20

        if ruta_pie and proporcion:
            # Se ajusta a una altura fija y se centra. La altura manda para que
            # el pie ocupe lo mismo en hojas verticales y horizontales.
            alto_pie = 9.0
            ancho_pie = alto_pie * proporcion
            if ancho_pie > ancho_util:
                ancho_pie = ancho_util
                alto_pie = ancho_pie / proporcion

            y_pie = pdf.h - alto_pie - 8
            x_pie = (pdf.w - ancho_pie) / 2

            pdf.set_draw_color(210, 210, 210)
            pdf.line(10, y_pie - 2.5, pdf.w - 10, y_pie - 2.5)
            pdf.image(ruta_pie, x_pie, y_pie, ancho_pie, alto_pie)

            if mostrar_pagina:
                pdf.set_y(-6)
                pdf.set_font("Helvetica", "", 6.5)
                pdf.set_text_color(150, 150, 150)
                pdf.cell(0, 4, f"{pdf.page_no()} / {{nb}}", align="R")
        else:
            pdf.set_y(-15)
            pdf.set_draw_color(210, 210, 210)
            pdf.line(10, pdf.get_y(), pdf.w - 10, pdf.get_y())
            pdf.set_y(-13)
            pdf.set_font("Helvetica", "", 6.5)
            pdf.set_text_color(90, 90, 90)
            pdf.cell(0, 5, safestr(MARCA_PIE), ln=True, align="C")
            if mostrar_pagina:
                pdf.set_text_color(150, 150, 150)
                pdf.cell(0, 4, f"{pdf.page_no()} / {{nb}}", align="R")

        pdf.set_text_color(0, 0, 0)
    except Exception:
        pass


# Lista de actividades "basura" (tareas internas de mantenimiento de equipo,
# no visitas técnicas reales) que se excluyen de la consolidación de datos.
# Antes duplicada entre app.py y sync_job.py.
ACTIVIDADES_BASURA = ['ACTUALIZACIONDATOS', 'ACTUALIZACIOFW', 'ACTUALIZAINFOTECNICA', 'ACTUALIZARDATOSTECNICOS', 'ACTUALIZARSENSOR']

# ==============================================================================
# MOTOR DE CONEXIÓN CON LA API DE CEPHEUS
# ==============================================================================
def consultar_api_ordenes(fecha_inicio_dt: datetime) -> pd.DataFrame:
    api_config = st.secrets.get("cepheus_api", {})
    base_url = api_config.get("url")
    auth_user = api_config.get("usuario", "monitorISCA")
    auth_pass = api_config.get("contrasena", "SV#8xgE4U#")
    
    usuarios_autorizados = api_config.get("usuarios_consulta", [])
    if isinstance(usuarios_autorizados, str):
        usuarios_autorizados = [usuarios_autorizados]
    
    fecha_str = fecha_inicio_dt.strftime("%d/%m/%Y")
    hora_str = "07:00" 
    auth = HTTPBasicAuth(auth_user, auth_pass)
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)", 
        "Accept": "*/*"
    }
    
    try:
        session = requests.Session()
        session.trust_env = False  
        
        for query_user in usuarios_autorizados:
            query_user = query_user.strip()
            url_exacta = f"{base_url}?usuario={query_user}&medios=FTTH,C&fechaInicio={fecha_str}&horaInicio={hora_str}"
            
            try:
                # verify=False porque es un servidor interno
                response = session.get(url_exacta, headers=headers, auth=auth, verify=False, timeout=240)
                
                if response.status_code == 200:
                    tipo_archivo = response.headers.get('Content-Type', 'Desconocido')
                    texto_resp = response.text.strip()
                    print(f"  -> [INFO] Cepheus respondió con un archivo tipo: {tipo_archivo}")
                    print(f"  -> [INFO] Tamaño de la respuesta: {len(response.content)} bytes")

                    # === WEBSERVICE NUEVO (producción): la respuesta SIEMPRE es JSON con
                    # la forma {"codigo":200, "mensaje":"...", "ordenes":[ {...}, ... ]}.
                    # Un JSON no es sinónimo de error: hay que parsearlo y revisar "codigo".
                    if 'json' in tipo_archivo.lower() or texto_resp.startswith('{'):
                        try:
                            data_json = response.json()
                        except Exception as e_json:
                            print(f"  -> [-] La respuesta decía ser JSON pero no se pudo parsear para {query_user}: {e_json}")
                            continue

                        codigo_resp = data_json.get('codigo')
                        ordenes = data_json.get('ordenes', [])
                        mensaje_resp = data_json.get('mensaje', 'Sin mensaje')

                        # Límite de 5 consultas/hora impuesto por IT (código 102). Insistir con
                        # las demás cuentas de la lista solo consumiría más cupo del límite
                        # compartido, así que se corta de inmediato y se espera al próximo ciclo.
                        if codigo_resp == 102:
                            print(f"  -> [!] Límite de consultas de Cepheus alcanzado (código 102): {mensaje_resp}. Se reintentará en el próximo ciclo, sin insistir con otras cuentas.")
                            return pd.DataFrame()

                        if codigo_resp != 200 or not ordenes:
                            print(f"  -> [!] Cepheus respondió sin órdenes para {query_user} (código {codigo_resp}): {mensaje_resp}. Probando siguiente usuario...")
                            continue

                        df_temp = pd.DataFrame(ordenes)
                        df_temp.columns = df_temp.columns.astype(str).str.upper().str.strip()
                        print(f"  -> [+] ¡{len(df_temp)} órdenes descargadas con éxito desde la cuenta {query_user}!")
                        return df_temp

                    # === Compatibilidad retro: por si algún día vuelve a responder un archivo binario ===
                    try:
                        df_temp = pd.read_excel(io.BytesIO(response.content))
                    except Exception as e_excel:
                        print(f"  -> [!] No es un Excel válido: {e_excel}. Intentando como CSV...")
                        try:
                            df_temp = pd.read_csv(io.StringIO(texto_resp), sep=None, engine='python')
                        except Exception as e_csv:
                            print(f"  -> [-] Falla total al leer el archivo. Error: {e_csv}")
                            df_temp = pd.DataFrame() 
                    
                    if not df_temp.empty:
                        df_temp.columns = df_temp.columns.str.upper().str.strip()
                        print(f"  -> [+] ¡Archivo rep_actividades descargado con éxito desde la cuenta {query_user}!")
                        return df_temp 
                    else:
                        print(f"  -> [!] Archivo procesado, pero la tabla quedó completamente vacía para {query_user}.")
                else:
                    print(f"  -> [-] Error de acceso con {query_user} (Código {response.status_code})")
            except Exception as e_user:
                print(f"  -> [-] Error descargando el archivo de {query_user}: {e_user}")
                
        return pd.DataFrame()
            
    except Exception as e:
        print(f"❌ Error crítico en motor de descarga: {e}")
        return pd.DataFrame()
        
def depurar_api_con_dispositivos(df_api: pd.DataFrame, file_dispositivos_o_df: Any) -> pd.DataFrame:
    """
    Cruza el DataFrame de la API con los dispositivos/vehículos (FTTX) subidos de forma manual
    o descargados desde la nube.
    """
    try:
        if isinstance(file_dispositivos_o_df, pd.DataFrame):
            dfdispfull = file_dispositivos_o_df
        elif isinstance(file_dispositivos_o_df, bytes):
            dfdispfull = pd.read_excel(io.BytesIO(file_dispositivos_o_df), engine='openpyxl')
        elif hasattr(file_dispositivos_o_df, 'read'):
            file_dispositivos_o_df.seek(0)
            if getattr(file_dispositivos_o_df, 'name', '').lower().endswith('.csv'):
                dfdispfull = pd.read_csv(file_dispositivos_o_df, sep=None, engine='python')
            else:
                dfdispfull = pd.read_excel(file_dispositivos_o_df, engine='openpyxl')
        else:
            dfdispfull = pd.DataFrame()

        dfdispref = pd.DataFrame()
        if not dfdispfull.empty:
            coltec = [c for c in dfdispfull.columns if any(x in str(c).upper() for x in ['TECNICO', 'USER', 'OPERADOR'])]
            colmx = [c for c in dfdispfull.columns if any(x in str(c).upper() for x in ['MX', 'VEHICULO', 'PLACA'])]
            dfdispref['TECREF'] = dfdispfull[coltec[0]].astype(str).str.strip().str.upper() if coltec else "N/D"
            dfdispref['MXREF'] = dfdispfull[colmx[0]].astype(str).str.strip() if colmx else "N/D"

        dfp = procesar_dataframe_base(df_api)
        dfp['TECKEY'] = dfp['TECNICO'].astype(str).str.strip().str.upper()
        
        if not dfdispref.empty:
            dffinal = dfp.merge(dfdispref.drop_duplicates('TECREF'), left_on='TECKEY', right_on='TECREF', how='left')
            if 'MXREF' in dffinal.columns:
                dffinal['MX'] = dffinal['MXREF'].combine_first(dffinal.get('MX', pd.Series(dtype=str)))
            return dffinal.drop(columns=['TECKEY', 'TECREF', 'MXREF'], errors='ignore')
        
        return dfp.drop(columns=['TECKEY'], errors='ignore')
    except Exception as e:
        raise Exception(f"Error en cruce con API: {str(e)}")

# ==============================================================================
# 2. CLASE PARA PDF (REPORTING AVANZADO Y TABLAS COMPLEJAS)
# ==============================================================================
class ReporteGenerencialPDF(FPDF):
    def header(self):
        _logo = obtener_logo_reporte()
        if _logo:
            # El logo se ajusta a una caja fija: así entra igual de bien el
            # isotipo cuadrado que el logotipo horizontal, sin deformarse ni
            # invadir el título ni la línea del encabezado.
            _lw, _lh = medidas_logo_en_caja(_logo)
            self.image(_logo, 10, 7, _lw, _lh)

        self.set_x(50) 
        self.set_text_color(0, 0, 0)
        self.set_font("Helvetica", "", 7)
        self.cell(80, 5, safestr("Reporte Operativo Consolidado"), ln=False, align="L")
        self.set_font("Helvetica", "B", 7)
        self.cell(0, 5, safestr(MARCA_TAGLINE), ln=True, align="R")

        self.set_draw_color(200, 200, 200)
        y_line = max(self.get_y(), 18) 
        self.line(10, y_line, 200, y_line)
        self.set_y(y_line + 5)

    def footer(self):
        dibujar_pie_membrete(self)

    def seccion_titulo(self, titulo):
        self.set_text_color(84, 98, 143)
        self.set_font("Helvetica", "B", 9)
        self.cell(0, 6, safestr(titulo), ln=True, align="L")
        self.ln(1)

    def dibujar_tabla_rendimiento(self, df, anchos=None, alineaciones=None):
        if df.empty: return
        self.set_fill_color(225, 225, 225)
        self.set_text_color(50, 50, 50)
        self.set_draw_color(230, 230, 230)
        self.set_font("Helvetica", "B", 7)
        numcols = len(df.columns)
        w = anchos if anchos else 190 / numcols
        aligns = alineaciones if (alineaciones and len(alineaciones) == numcols) else ["C"] * numcols
        for i, col in enumerate(df.columns):
            widthcell = w if isinstance(w, (int, float)) else w[i]
            self.cell(widthcell, 6, safestr(str(col).upper()), border=1, align="C", fill=True)
        self.ln()
        self.set_font("Helvetica", "", 7)
        for _, fila in df.iterrows():
            for i, item in enumerate(fila):
                widthcell = w if isinstance(w, (int, float)) else w[i]
                valstr = str(item)[:40]
                valclean = safestr(valstr)
                fillr, fillg, fillb = 255, 255, 255
                textr, textg, textb = 0, 0, 0
                
                if df.columns[i] in ['% LOGRO FINAL', '% LOGRO SEMANAL', '% LOGRO META']:
                    try:
                        pct = float(valstr.replace('%', ''))
                        if pct >= 100: fillr, fillg, fillb = 146, 208, 80 
                        elif pct >= 80: fillr, fillg, fillb = 169, 208, 142 
                        elif pct >= 50: fillr, fillg, fillb = 255, 230, 153 
                        elif pct >= 0: fillr, fillg, fillb = 244, 176, 132 
                    except: pass
                    
                if df.columns[i] == 'BONO MIXTO':
                    if valstr != '+0.0%':
                        fillr, fillg, fillb = 220, 235, 255 

                self.set_fill_color(fillr, fillg, fillb)
                self.set_text_color(textr, textg, textb)
                self.cell(widthcell, 5, valclean, border=1, align=aligns[i], fill=True)
            self.ln()
        self.ln(4)

    def dibujar_tabla(self, df, anchos=None, alineaciones=None):
        if df.empty: return
        self.set_fill_color(225, 225, 225)
        self.set_text_color(50, 50, 50)
        self.set_draw_color(230, 230, 230)
        self.set_font("Helvetica", "B", 7)
        numcols = len(df.columns)
        w = anchos if anchos else 190 / numcols
        aligns = alineaciones if (alineaciones and len(alineaciones) == numcols) else ["C"] * numcols
        for i, col in enumerate(df.columns):
            widthcell = w if isinstance(w, (int, float)) else w[i]
            self.cell(widthcell, 6, safestr(str(col).upper()), border=1, align="C", fill=True)
        self.ln()
        self.set_font("Helvetica", "", 7)
        for _, fila in df.iterrows():
            for i, item in enumerate(fila):
                widthcell = w if isinstance(w, (int, float)) else w[i]
                valstr = str(item)[:40]
                self.cell(widthcell, 5, safestr(valstr), border=1, align=aligns[i], fill=False)
            self.ln()
        self.ln(4)

    def dibujar_tabla_tiempos_rangos(self, titulo, headercolname, dfsubset, pivotcol, showtotalcol=False):
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(84, 98, 143)
        self.cell(0, 6, safestr(titulo), ln=True, align="L")
        if dfsubset.empty:
            self.set_text_color(0, 0, 0); self.set_font("Helvetica", "", 7)
            self.cell(0, 6, "Sin datos disponibles.", ln=True); self.ln(2)
            return
        rangosorden = ['0. Anulada', '1. Menos de 1 Día', '2. De 1 a 3 Días', '3. De 3 a 6 Días', '4. Más de 6 Días', '6. Pendiente']
        pivotvals = dfsubset[pivotcol].value_counts().index.tolist()
        if showtotalcol: pivotvals.append('Total')
        wcol1 = 35
        wsub = 18
        self.set_fill_color(210, 210, 215)
        self.set_text_color(50, 50, 50)
        self.set_font("Helvetica", "B", 7)
        self.set_draw_color(220, 220, 220)
        self.cell(wcol1, 6, safestr(headercolname), border=1, align="C", fill=True)
        for pval in pivotvals:
            self.cell(wsub * 2, 6, safestr(pval), border=1, align="C", fill=True)
        self.ln()
        self.cell(wcol1, 6, "Rango Dias a Visita", border=1, align="C", fill=True)
        for pval in pivotvals:
            self.cell(wsub, 6, "Cantidad", border=1, align="C", fill=True)
            self.cell(wsub, 6, "%", border=1, align="C", fill=True)
        self.ln()
        datos = {}
        for pval in pivotvals:
            dfp = dfsubset if pval == 'Total' else dfsubset[dfsubset[pivotcol] == pval]
            datos[pval] = {'total': len(dfp), 'counts': dfp['RANGOTIEMPO'].value_counts()}
        self.set_font("Helvetica", "", 7)
        self.set_text_color(0, 0, 0)
        for rango in rangosorden:
            self.set_fill_color(255, 255, 255)
            self.cell(wcol1, 5, safestr(rango), border=1, align="L", fill=True)
            for pval in pivotvals:
                count = datos[pval]['counts'].get(rango, 0)
                tot = datos[pval]['total']
                pct = (count / tot * 100) if tot > 0 else 0
                cntstr = str(count) if count > 0 else ""
                pctstr = f"{pct:.0f}%" if count > 0 else ""
                fr, fg, fb = 255, 255, 255
                if count > 0 and 'Menos' in rango: 
                    if pct >= 75: fr, fg, fb = 146, 208, 80 
                    elif pct >= 40: fr, fg, fb = 255, 230, 153 
                    elif pct >= 25: fr, fg, fb = 244, 176, 132 
                    else: fr, fg, fb = 234, 153, 153 
                elif count > 0: 
                    if pct >= 75: fr, fg, fb = 146, 208, 80
                    elif pct >= 40: fr, fg, fb = 255, 230, 153
                    elif pct >= 25: fr, fg, fb = 244, 176, 132
                    else: fr, fg, fb = 234, 153, 153 
                self.set_fill_color(255, 255, 255)
                self.cell(wsub, 5, cntstr, border=1, align="C", fill=True)
                self.set_fill_color(fr, fg, fb)
                self.cell(wsub, 5, pctstr, border=1, align="C", fill=True)
            self.ln()
        self.set_font("Helvetica", "B", 7)
        self.set_fill_color(240, 240, 240)
        self.cell(wcol1, 5, "Total", border=1, align="L", fill=True)
        for pval in pivotvals:
            tot = datos[pval]['total']
            self.cell(wsub, 5, str(tot) if tot>0 else "0", border=1, align="C", fill=True)
            self.cell(wsub, 5, "100%" if tot>0 else "0%", border=1, align="C", fill=True)
        self.ln(6)

    def dibujar_tabla_cerradas_ciudad(self, dfbase):
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(84, 98, 143)
        self.cell(0, 6, safestr("Ordenes Cerradas y Tiempo Promedio de Atencion por Ciudad"), ln=True, align="L")
        dfcerradas = dfbase[dfbase['ESTADO'].astype(str).str.upper() == 'CERRADA'].copy()
        if dfcerradas.empty:
            self.set_text_color(0, 0, 0); self.set_font("Helvetica", "", 7)
            self.cell(0, 6, "Sin datos de ordenes cerradas.", ln=True); self.ln(2)
            return
        dfgrp = dfcerradas.groupby(['SECTOR', 'TIPOACTDETALLE']).agg(
            CANTIDAD=('NUM', 'count'), MINUTOSPROMEDIO=('MINUTOS_CALC', 'mean')
        ).reset_index()
        dfgrp['MINUTOSPROMEDIO'] = dfgrp['MINUTOSPROMEDIO'].round(0).fillna(0).astype(int)
        wcity, wact, wcant, wmin = 40, 60, 30, 40
        self.set_fill_color(210, 210, 215); self.set_text_color(50, 50, 50); self.set_font("Helvetica", "B", 7)
        self.cell(wcity, 6, "Ciudad", border=1, align="C", fill=True)
        self.cell(wact, 6, "Tipo Actividad", border=1, align="C", fill=True)
        self.cell(wcant, 6, "Cantidad", border=1, align="C", fill=True)
        self.cell(wmin, 6, "Minutos Promedio", border=1, align="C", fill=True)
        self.ln()
        self.set_font("Helvetica", "", 7); self.set_text_color(0, 0, 0)
        sectores = sorted(dfgrp['SECTOR'].unique())
        grandtotcant = grandtotminsum = 0
        for sec in sectores:
            dfsec = dfgrp[dfgrp['SECTOR'] == sec].sort_values(by='CANTIDAD', ascending=False)
            first = True; sectotcant = sectotminsum = 0
            for _, row in dfsec.iterrows():
                self.set_fill_color(255, 255, 255)
                bordercity = "LTR" if first else "LR"
                self.cell(wcity, 5, safestr(sec) if first else "", border=bordercity, align="L", fill=True)
                self.cell(wact, 5, safestr(row['TIPOACTDETALLE']), border=1, align="L", fill=True)
                self.cell(wcant, 5, str(row['CANTIDAD']), border=1, align="C", fill=True)
                self.cell(wmin, 5, str(row['MINUTOSPROMEDIO']), border=1, align="C", fill=True)
                self.ln()
                first = False
                sectotcant += row['CANTIDAD']
                sectotminsum += row['MINUTOSPROMEDIO'] * row['CANTIDAD']
            secprom = int(sectotminsum / sectotcant) if sectotcant > 0 else 0
            self.set_font("Helvetica", "B", 7); self.set_fill_color(248, 248, 248)
            self.cell(wcity, 5, "", border="LRB", align="L", fill=True) 
            self.cell(wact, 5, "Total", border=1, align="L", fill=True)
            self.cell(wcant, 5, str(sectotcant), border=1, align="C", fill=True)
            self.cell(wmin, 5, str(secprom), border=1, align="C", fill=True)
            self.ln()
            self.set_font("Helvetica", "", 7)
            grandtotcant += sectotcant
            grandtotminsum += sectotminsum
        grandprom = int(grandtotminsum / grandtotcant) if grandtotcant > 0 else 0
        self.set_font("Helvetica", "B", 7); self.set_fill_color(240, 240, 240)
        self.cell(wcity + wact, 6, "Total General", border=1, align="L", fill=True)
        self.cell(wcant, 6, str(grandtotcant), border=1, align="C", fill=True)
        self.cell(wmin, 6, str(grandprom), border=1, align="C", fill=True)
        self.ln(6)

    def dibujar_tabla_tiempos_actividad(self, dfbase):
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(84, 98, 143)
        self.cell(0, 6, safestr("Tiempos de Atencion Promedio por Colaborador y Actividad"), ln=True, align="L")
        
        if dfbase.empty:
            self.set_text_color(0, 0, 0); self.set_font("Helvetica", "", 7)
            self.cell(0, 6, "Sin datos disponibles.", ln=True); self.ln(2)
            return
            
        dfgrp = dfbase.groupby(['TECNICO', 'ACTIVIDAD']).agg(
            CANTIDAD=('NUM', 'count'), MINUTOSPROMEDIO=('MINUTOS_CALC', 'mean')
        ).reset_index()
        dfgrp['MINUTOSPROMEDIO'] = dfgrp['MINUTOSPROMEDIO'].round(1)
        
        wtec, wact, wcant, wmin = 55, 65, 30, 40
        
        self.set_fill_color(210, 210, 215); self.set_text_color(50, 50, 50); self.set_font("Helvetica", "B", 7)
        self.cell(wtec, 6, "Colaborador", border=1, align="C", fill=True)
        self.cell(wact, 6, "Actividad", border=1, align="C", fill=True)
        self.cell(wcant, 6, "Ordenes Atendidas", border=1, align="C", fill=True)
        self.cell(wmin, 6, "Prom. Duracion (Min)", border=1, align="C", fill=True)
        self.ln()
        
        self.set_font("Helvetica", "", 7); self.set_text_color(0, 0, 0)
        tecnicos = sorted(dfgrp['TECNICO'].unique())
        
        for tec in tecnicos:
            dftec = dfgrp[dfgrp['TECNICO'] == tec].sort_values(by='CANTIDAD', ascending=False)
            first = True
            tectotcant = 0
            tectotminsum = 0
            
            for _, row in dftec.iterrows():
                self.set_fill_color(255, 255, 255)
                bordertec = "LTR" if first else "LR"
                self.cell(wtec, 5, safestr(tec)[:32] if first else "", border=bordertec, align="L", fill=True)
                self.cell(wact, 5, safestr(row['ACTIVIDAD'])[:35], border=1, align="L", fill=True)
                self.cell(wcant, 5, str(row['CANTIDAD']), border=1, align="C", fill=True)
                self.cell(wmin, 5, str(row['MINUTOSPROMEDIO']), border=1, align="C", fill=True)
                self.ln()
                first = False
                tectotcant += row['CANTIDAD']
                tectotminsum += row['MINUTOSPROMEDIO'] * row['CANTIDAD']
                
            tecprom = round((tectotminsum / tectotcant), 1) if tectotcant > 0 else 0
            self.set_font("Helvetica", "B", 7); self.set_fill_color(248, 248, 248)
            self.cell(wtec, 5, "", border="LRB", align="L", fill=True) 
            self.cell(wact, 5, "Total", border=1, align="R", fill=True)
            self.cell(wcant, 5, str(tectotcant), border=1, align="C", fill=True)
            self.cell(wmin, 5, str(tecprom), border=1, align="C", fill=True)
            self.ln()
            self.set_font("Helvetica", "", 7)
        self.ln(6)

def finalizar_pdf(pdfobj):
    """Guarda y retorna el archivo PDF de manera segura."""
    fd, tmppath = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    try:
        pdfobj.output(tmppath)
        with open(tmppath, "rb") as f: return f.read()
    finally:
        try: os.remove(tmppath)
        except: pass

def generar_graficos_temporales(dfbase):
    paths = {}
    try:
        import matplotlib
        matplotlib.use('Agg') 
        import matplotlib.pyplot as plt
        actstr = dfbase['ACTIVIDAD'].astype(str).str.upper()
        maskins = actstr.str.contains('INS|NUEVA|ADIC|CAMBIO|PLEX')
        masksop = actstr.str.contains('SOP|FALLA|MANT')
        totins = len(dfbase[maskins])
        totsop = len(dfbase[masksop])
        tototros = len(dfbase[~(maskins | masksop)])
        labels, sizes, colors = [], [], []
        if totins > 0: labels.append('Instalaciones'); sizes.append(totins); colors.append('#5C82A6')
        if totsop > 0: labels.append('Mantenimientos'); sizes.append(totsop); colors.append('#A5B1C2')
        if tototros > 0: labels.append('Otros'); sizes.append(tototros); colors.append('#D1D8E0')
        if sizes:
            fig1, ax1 = plt.subplots(figsize=(4, 3))
            ax1.pie(sizes, labels=labels, autopct='%1.0f%%', startangle=90, colors=colors,
                    textprops={'fontsize': 8, 'color': '#333333'}, wedgeprops={'edgecolor': 'white'})
            ax1.axis('equal')
            plt.title('Instalaciones vs Mantenimientos', fontsize=9, color='#4A628A', fontweight='bold', pad=10)
            fdpie, pathpie = tempfile.mkstemp(suffix=".png")
            os.close(fdpie)
            plt.savefig(pathpie, bbox_inches='tight', dpi=150, transparent=True)
            plt.close(fig1)
            paths['pie'] = pathpie
            
        dffechas = dfbase.copy()
        dffechas['FECHAAPEDT'] = pd.to_datetime(dffechas['FECHA_APE'], errors='coerce')
        dffechas = dffechas.dropna(subset=['FECHAAPEDT'])
        if not dffechas.empty:
            conteofechas = dffechas.groupby(dffechas['FECHAAPEDT'].dt.date).size().tail(7)
            if not conteofechas.empty:
                fig2, ax2 = plt.subplots(figsize=(5, 3))
                etiquetasx = [d.strftime('%d/%m') for d in conteofechas.index]
                bars = ax2.bar(etiquetasx, conteofechas.values, color='#8FA1B3')
                ax2.set_title('Creacion de Ordenes por Fecha (Ultimos 7 dias)', fontsize=9, color='#4A628A', fontweight='bold', pad=10)
                ax2.tick_params(axis='x', rotation=30, labelsize=7, colors='#555555')
                ax2.tick_params(axis='y', labelsize=7, colors='#555555')
                ax2.spines['top'].set_visible(False)
                ax2.spines['right'].set_visible(False)
                ax2.spines['left'].set_color('#DDDDDD')
                ax2.spines['bottom'].set_color('#DDDDDD')
                for bar in bars:
                    yval = bar.get_height()
                    ax2.text(bar.get_x() + bar.get_width()/2, yval + (yval*0.02), int(yval),
                             ha='center', va='bottom', fontsize=7, color='#333333')
                fdbar, pathbar = tempfile.mkstemp(suffix=".png")
                os.close(fdbar)
                plt.tight_layout()
                plt.savefig(pathbar, bbox_inches='tight', dpi=150, transparent=True)
                plt.close(fig2)
                paths['bar'] = pathbar
        return paths
    except ImportError:
        return {}
    except Exception as e:
        return {}

def _generar_dona_png(pct, titulo):
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        color = "#EF4444" if pct < 50 else ("#F59E0B" if pct < 80 else "#10B981")
        fig, ax = plt.subplots(figsize=(2.5, 2.5))
        ax.pie([pct, max(0, 100-pct)], colors=[color, '#E5E7EB'], startangle=90, counterclock=False, wedgeprops=dict(width=0.3, edgecolor='w'))
        ax.text(0, 0, f"{pct:.0f}%", ha='center', va='center', fontsize=20, fontweight='bold', color=color)
        plt.title(titulo, fontsize=10, color='#333333', fontweight='bold', pad=5)
        fd, path = tempfile.mkstemp(suffix=".png")
        os.close(fd)
        plt.savefig(path, bbox_inches='tight', dpi=120, transparent=True)
        plt.close(fig)
        return path
    except:
        return None

def calcular_aporte_meta(row):
    act = str(row.get('ACTIVIDAD', '')).upper()
    com = str(row.get('COMENTARIO', '')).upper()
    txt = act + " " + com
    if 'PEXTERNO' in act: return 100.0  
    elif re.search('ADIC|CAMBIO|MIGRACI|RECUP', txt): return 12.5   
    elif re.search('INS|NUEVA|PLEX|SPLITTEROPT', act): return 25.0   
    elif re.search('SOP|FALLA|MANT|RECON|TRASLADO', act): return 12.5   
    else: return 12.5   

def generar_pdf_cierre_diario(dfbase, fechatarget):
    dfc = dfbase[
        (dfbase['HORA_LIQ'].dt.date == fechatarget) & 
        (dfbase['ESTADO'].astype(str).str.contains('CERRADA', na=False, case=False))
    ].copy()
    
    def get_tipo_detalle(row):
        txt = (str(row.get('ACTIVIDAD', '')) + " " + str(row.get('COMENTARIO', ''))).upper()
        if 'RECON' in txt: return 'RECONEXIONES'
        if 'TRASLADO' in txt: return 'TRASLADOS'
        if re.search('INS|NUEVA|ADIC|CAMBIO|PLEX|MIGRACI|RECUP', txt): return 'INSTALACION'
        if re.search('SOP|FALLA|MANT', txt): return 'MANTENIMIENTO'
        return 'OTROS'
        
    def get_tipo_orden(row):
        txt = (str(row.get('ACTIVIDAD', '')) + " " + str(row.get('COMENTARIO', ''))).upper()
        if re.search('INS|NUEVA|ADIC|CAMBIO|PLEX|MIGRACI|RECUP', txt): return 'INSTALACION'
        if re.search('SOP|FALLA|MANT', txt): return 'MANTENIMIENTO'
        return 'OTROS'

    def get_rango(row):
        est = str(row.get('ESTADO', '')).upper()
        dias = row.get('DIAS_RETRASO', 0)
        if 'ANULADA' in est: return '0. Anulada'
        if 'CERRADA' not in est: return '6. Pendiente'
        if dias < 1: return '1. Menos de 1 Día'
        if 1 <= dias <= 3: return '2. De 1 a 3 Días'
        if 4 <= dias <= 6: return '3. De 3 a 6 Días'
        return '4. Más de 6 Días'

    if not dfc.empty:
        dfc['TIPOACTDETALLE'] = dfc.apply(get_tipo_detalle, axis=1)
        dfc['TIPOORDEN'] = dfc.apply(get_tipo_orden, axis=1)
        if 'DIAS_RETRASO' not in dfc.columns:
            ahora = pd.Timestamp(datetime.now())
            dfc['DIAS_RETRASO'] = (ahora.normalize() - pd.to_datetime(dfc['FECHA_APE'], errors='coerce').dt.normalize()).dt.days.fillna(0).clip(lower=0).astype(int)
        dfc['RANGOTIEMPO'] = dfc.apply(get_rango, axis=1)

    # 1. Inicialización de la clase PDF al nivel principal de la función
    pdf = ReporteGenerencialPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(84, 98, 143)
    pdf.set_draw_color(220, 220, 220)
    pdf.set_fill_color(252, 252, 252)
    pdf.cell(0, 10, safestr(f" Reporte Analitico de Cierre Diario: {fechatarget}"), border=1, ln=True, fill=True)
    pdf.ln(5)
    
    pdf.seccion_titulo("Analisis de Eficiencia (Puntos por Meta + 10% Bono por Ruta Mixta)")
    if not dfc.empty:
        dfc['CANT_INS'] = (dfc['TIPOORDEN'] == 'INSTALACION').astype(int)
        dfc['CANT_SOP'] = (dfc['TIPOORDEN'] == 'MANTENIMIENTO').astype(int)
        dfc['CANT_OTR'] = (dfc['TIPOORDEN'] == 'OTROS').astype(int)
        dfc['%_APORTE'] = dfc.apply(calcular_aporte_meta, axis=1)
        
        df_tec = dfc.groupby('TECNICO').agg(
            CANT_INS=('CANT_INS', 'sum'),
            CANT_SOP=('CANT_SOP', 'sum'),
            CANT_OTR=('CANT_OTR', 'sum'),
            PUNTOS_BASE=('%_APORTE', 'sum')
        ).reset_index()
        
        def calcular_bono(row):
            types_count = sum([1 for x in [row['CANT_INS'], row['CANT_SOP'], row['CANT_OTR']] if x > 0])
            if types_count > 1: return 10.0 
            return 0.0
            
        df_tec['BONO_MIXTO'] = df_tec.apply(calcular_bono, axis=1)
        df_tec['LOGRO_FINAL'] = df_tec['PUNTOS_BASE'] + df_tec['BONO_MIXTO']
        df_tec = df_tec.sort_values(by='LOGRO_FINAL', ascending=False)
        
        df_tec_table = df_tec[['TECNICO', 'CANT_INS', 'CANT_SOP', 'CANT_OTR', 'PUNTOS_BASE', 'BONO_MIXTO', 'LOGRO_FINAL']].copy()
        
        # 1. Renombramos las columnas (ahora tienen espacios)
        df_tec_table.columns = ['TECNICO', 'INS', 'SOP', 'OTR', 'PUNTOS BASE', 'BONO MIXTO', '% LOGRO FINAL']
        
        # 2. Aplicamos el formato leyendo exactamente el nombre NUEVO con espacio
        df_tec_table['PUNTOS BASE'] = df_tec_table['PUNTOS BASE'].round(1).astype(str) + '%'
        df_tec_table['BONO MIXTO'] = '+' + df_tec_table['BONO MIXTO'].round(1).astype(str) + '%' # <-- CORREGIDO AQUÍ
        df_tec_table['% LOGRO FINAL'] = df_tec_table['% LOGRO FINAL'].round(1).astype(str) + '%'
        
        pdf.dibujar_tabla_rendimiento(df_tec_table, anchos=[55, 15, 15, 15, 30, 30, 30], alineaciones=["L", "C", "C", "C", "C", "C", "C"])
    else:
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 6, "Sin datos de productividad para hoy.", ln=True)

    if not dfc.empty:
        pdf.add_page()
        pdf.seccion_titulo("Indicadores de Avance Operativo (Completado vs Pendiente)")
        
        mask_tec = (dfbase['TECNICO'].notna() & (dfbase['TECNICO'].astype(str).str.strip() != '') & (~dfbase['TECNICO'].astype(str).str.upper().isin(['NONE', 'NAN', 'N/D', 'NULL'])))
        dfv = dfbase[mask_tec].copy()
        df_vivas = dfv[dfv['ESTADO'].astype(str).str.contains(PATRON_ASIGNADAS_VIVA_STR, na=False, case=False)]
        
        resi_pend = len(df_vivas[df_vivas['SEGMENTO'] == 'RESIDENCIAL'])
        resi_cerr = len(dfc[dfc['SEGMENTO'] == 'RESIDENCIAL'])
        t_resi = resi_pend + resi_cerr
        pct_resi = (resi_cerr / t_resi * 100) if t_resi > 0 else 0
        
        plex_pend = len(df_vivas[df_vivas['SEGMENTO'] == 'PLEX'])
        plex_cerr = len(dfc[dfc['SEGMENTO'] == 'PLEX'])
        t_plex = plex_pend + plex_cerr
        pct_plex = (plex_cerr / t_plex * 100) if t_plex > 0 else 0
        
        t_global = len(df_vivas) + len(dfc)
        pct_global = (len(dfc) / t_global * 100) if t_global > 0 else 0

        path_resi = _generar_dona_png(pct_resi, "Residencial")
        path_plex = _generar_dona_png(pct_plex, "PLEX")
        path_global = _generar_dona_png(pct_global, "Global")

        current_y = pdf.get_y()
        if path_resi: pdf.image(path_resi, x=20, y=current_y, w=50)
        if path_plex: pdf.image(path_plex, x=80, y=current_y, w=50)
        if path_global: pdf.image(path_global, x=140, y=current_y, w=50)
        
        pdf.ln(60) 
        
        for path in [path_resi, path_plex, path_global]:
            if path:
                try: os.remove(path)
                except: pass
        
        pdf.add_page()
        pdf.seccion_titulo("Tiempos de Atencion (Antiguedad de Ordenes Liquidadas)")
        pdf.ln(2)
        dfins = dfc[dfc['TIPOORDEN'] == 'INSTALACION']
        pdf.dibujar_tabla_tiempos_rangos("Instalaciones Liquidadas por Rango", "Ciudad", dfins, 'SECTOR', showtotalcol=False)
        dfmant = dfc[dfc['TIPOORDEN'] == 'MANTENIMIENTO']
        pdf.dibujar_tabla_tiempos_rangos("Mantenimientos Liquidados por Rango", "Ciudad", dfmant, 'SECTOR', showtotalcol=False)
        
        pdf.add_page()
        pdf.dibujar_tabla_cerradas_ciudad(dfc)

        # ==============================================================================
        # RESUMEN CONSOLIDADO POR TIPO DE ACTIVIDAD (CON SUMATORIA TOTAL)
        # ==============================================================================
        pdf.add_page()
        pdf.seccion_titulo("Resumen Consolidado por Tipo de Actividad")
        df_act_summary = dfc['ACTIVIDAD'].value_counts().reset_index()
        df_act_summary.columns = ['Actividad Realizada', 'Total de Ordenes']
        
        # Dibujamos la tabla estándar
        pdf.dibujar_tabla(df_act_summary, anchos=[120, 40], alineaciones=["L", "C"])
        
        # Ajustamos sutilmente el espaciado para pegar el totalizador a la tabla
        pdf.set_y(pdf.get_y() - 4)
        
        # Configuramos el estilo de la fila de Sumatoria Total
        pdf.set_fill_color(240, 240, 240)  # Fondo gris claro institucional
        pdf.set_text_color(0, 0, 0)        # Texto negro
        pdf.set_font("Helvetica", "B", 7)  # Fuente en negrita
        
        # Calcular la suma total de las órdenes
        total_ords = df_act_summary['Total de Ordenes'].sum()
        
        # Dibujamos las celdas alineadas con las columnas superiores (120 y 40 de ancho)
        pdf.cell(120, 6, safestr("TOTAL GENERAL"), border=1, fill=True, align="R")
        pdf.cell(40, 6, safestr(str(int(total_ords))), border=1, fill=True, align="C")
        pdf.ln(10)

        pdf.add_page()
        pdf.dibujar_tabla_tiempos_actividad(dfc)

    pdf.add_page()
    pdf.seccion_titulo("Consolidado General de Ordenes Liquidadas")
    if not dfc.empty:
        pdf.dibujar_tabla(dfc[['NUM', 'TECNICO', 'ACTIVIDAD', 'TIEMPO_REAL']], anchos=[30, 60, 60, 40], alineaciones=["C", "L", "L", "C"])

    if not dfc.empty:
        imagenes = generar_graficos_temporales(dfc)
        if imagenes and 'pie' in imagenes:
            pdf.add_page()
            pdf.seccion_titulo("Distribucion Grafica de la Jornada")
            pdf.image(imagenes['pie'], x=60, y=pdf.get_y() + 5, w=90)
            for path in imagenes.values():
                try: os.remove(path)
                except: pass
                
    return finalizar_pdf(pdf)

def generar_pdf_trimestral_detallado(tabla_produccion, tabla_eficiencia, resumen_jornada):
    pdf = ReporteGenerencialPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(40, 50, 100)
    pdf.cell(0, 10, safestr("REPORTE GERENCIAL: RENDIMIENTO Y JORNADA DE TECNICOS"), border=0, ln=True, align="C")
    
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(100, 100, 100)
    ahorastr = datetime.now().strftime('%d/%m/%Y %I:%M %p')
    pdf.cell(0, 6, safestr(f"Generado el: {ahorastr}"), ln=True, align="C")
    pdf.ln(5)
    
    if resumen_jornada.empty:
        pdf.cell(0, 10, "No hay datos suficientes para generar el reporte.", ln=True)
        return finalizar_pdf(pdf)

    lista_tecnicos = resumen_jornada['TECNICO'].dropna().unique()
    
    for tecnico in lista_tecnicos:
        if pdf.get_y() > 220:
            pdf.add_page()
            
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_fill_color(230, 240, 255)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 8, safestr(f"   TECNICO: {tecnico}"), border=1, ln=True, fill=True)
        
        df_jor = resumen_jornada[resumen_jornada['TECNICO'] == tecnico]
        df_prod = tabla_produccion[tabla_produccion['TECNICO'] == tecnico]
        df_efi = tabla_eficiencia[tabla_eficiencia['TECNICO'] == tecnico]
        
        pdf.set_font("Helvetica", "B", 8)
        pdf.cell(0, 6, "   RESUMEN DE JORNADA LABORAL", ln=True)
        
        pdf.set_font("Helvetica", "", 8)
        prom_horas = df_jor['Promedio_Horas_Dia'].values[0] if not df_jor.empty else 0
        dias_lab = df_jor['Dias_Laborados'].values[0] if not df_jor.empty else 0
        max_horas = df_jor['Max_Horas_Dia'].values[0] if not df_jor.empty else 0
        
        pdf.cell(10, 5, "", border=0)
        pdf.cell(50, 5, safestr(f"Dias Trabajados: {dias_lab}"), border=0)
        pdf.cell(60, 5, safestr(f"Promedio en Calle: {prom_horas:.2f} hrs/dia"), border=0)
        pdf.cell(50, 5, safestr(f"Dia mas largo: {max_horas:.2f} hrs"), border=0, ln=True)
        pdf.ln(2)
        
        pdf.set_font("Helvetica", "B", 8)
        pdf.cell(0, 6, "   DESGLOSE DE ACTIVIDAD Y TIEMPOS", ln=True)
        
        pdf.set_fill_color(245, 245, 245)
        pdf.cell(10, 5, "", border=0)
        pdf.cell(60, 5, "Tipo de Actividad", border=1, align="C", fill=True)
        pdf.cell(25, 5, "Volumen", border=1, align="C", fill=True)
        pdf.cell(25, 5, "% del Total", border=1, align="C", fill=True)
        pdf.cell(40, 5, "Promedio de Resolucion", border=1, align="C", fill=True)
        pdf.ln()
        
        pdf.set_font("Helvetica", "", 8)
        total_ordenes_tec = 0
        
        df_prod = df_prod.sort_values(by='Cantidad', ascending=False)
        
        for _, fila_p in df_prod.iterrows():
            actividad = str(fila_p['ACTIVIDAD'])
            cantidad = fila_p['Cantidad']
            porcentaje = fila_p['Participacion_%']
            total_ordenes_tec += cantidad
            
            fila_efi = df_efi[df_efi['ACTIVIDAD'] == actividad]
            minutos_prom = fila_efi['Promedio_Minutos'].values[0] if not fila_efi.empty else 0
            
            pdf.cell(10, 5, "", border=0)
            pdf.cell(60, 5, safestr(actividad[:35]), border=1)
            pdf.cell(25, 5, safestr(str(cantidad)), border=1, align="C")
            pdf.cell(25, 5, safestr(f"{porcentaje}%"), border=1, align="C")
            
            if pd.notnull(minutos_prom) and minutos_prom > 120:
                pdf.set_text_color(200, 0, 0)
                pdf.cell(40, 5, safestr(f"{minutos_prom:.0f} min [!]"), border=1, align="C")
                pdf.set_text_color(0, 0, 0)
            elif pd.notnull(minutos_prom):
                pdf.cell(40, 5, safestr(f"{minutos_prom:.0f} min"), border=1, align="C")
            else:
                pdf.cell(40, 5, "---", border=1, align="C")
            
            pdf.ln()
            
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_fill_color(240, 240, 240)
        pdf.cell(10, 5, "", border=0)
        pdf.cell(60, 5, "TOTAL ORDENES", border=1, align="R", fill=True)
        pdf.cell(25, 5, safestr(str(total_ordenes_tec)), border=1, align="C", fill=True)
        pdf.cell(65, 5, "", border=0, ln=True)
        
        pdf.ln(8)
        
    return finalizar_pdf(pdf)

def generar_pdf_primera_orden(df_base, fecha_cierre):
    try:
        mask_vivas = df_base['ESTADO'].astype(str).str.contains(PATRON_ASIGNADAS_VIVA_STR, na=False, case=False)
        mask_cerradas = (pd.to_datetime(df_base['HORA_LIQ'], errors='coerce').dt.date == fecha_cierre) & (df_base['ESTADO'].astype(str).str.contains('CERRADA', na=False, case=False))
        
        df_universo = pd.concat([df_base[mask_vivas], df_base[mask_cerradas]]).drop_duplicates(subset=['NUM'])
        
        if 'HORA_INI' in df_universo.columns:
            df_universo['HORA_INI_DT'] = pd.to_datetime(df_universo['HORA_INI'], errors='coerce')
            df_universo = df_universo.dropna(subset=['HORA_INI_DT'])
            
            mask_fecha_ini = df_universo['HORA_INI_DT'].dt.date == pd.to_datetime(fecha_cierre).date()
            df_primera = df_universo[mask_fecha_ini].sort_values(by='HORA_INI_DT').drop_duplicates(subset=['TECNICO'], keep='first')
            df_primera = df_primera.sort_values(by='HORA_INI_DT')
        else:
            return None 

        pdf = ReporteGenerencialPDF()
        pdf.alias_nb_pages()
        pdf.add_page()
        
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(40, 50, 100)
        pdf.cell(0, 10, safestr(f"REPORTE: PRIMERA ORDEN DEL DIA ({fecha_cierre})"), border=0, ln=True, align="C")
        pdf.ln(5)

        if not df_primera.empty:
            df_mostrar = df_primera[['TECNICO', 'HORA_INI_DT', 'COLONIA', 'NUM']].copy()
            df_mostrar['HORA_INI'] = df_mostrar['HORA_INI_DT'].dt.strftime('%H:%M:%S')
            df_mostrar = df_mostrar.drop(columns=['HORA_INI_DT'])
            df_mostrar = df_mostrar[['TECNICO', 'HORA_INI', 'COLONIA', 'NUM']]
            
            pdf.dibujar_tabla(df_mostrar, anchos=[70, 30, 60, 30], alineaciones=["L", "C", "L", "C"])
        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(0, 10, "No hay registros de primera orden para esta fecha.", ln=True, align="C")

        return finalizar_pdf(pdf)
    except Exception as e:
        print(f"Error al generar PDF de Primera Orden: {e}")
        return None

def _valor_flexible(row, candidatos, default=''):
    """Busca un valor en la fila probando varios nombres de columna posibles
    (soporta encabezados crudos del Excel que aun no fueron mapeados)."""
    for nombre in candidatos:
        if nombre in row.index:
            val = row.get(nombre)
            if pd.notna(val) and str(val).strip().upper() not in ('', 'NAN', 'NONE', 'N/D', 'NULL'):
                return val
    return default


def _color_por_dias(pdf, dias_val):
    if dias_val >= 7:
        pdf.set_fill_color(211, 47, 47)
        pdf.set_text_color(255, 255, 255)
    elif dias_val >= 4:
        pdf.set_fill_color(245, 124, 0)
        pdf.set_text_color(255, 255, 255)
    elif dias_val >= 1:
        pdf.set_fill_color(251, 192, 45)
        pdf.set_text_color(0, 0, 0)
    else:
        pdf.set_fill_color(56, 142, 60)
        pdf.set_text_color(255, 255, 255)


def _ajustar_texto_a_celda(pdf, texto, ancho_celda, margen=1.6):
    """Recorta el texto (con puntos suspensivos) para que quepa dentro del ancho real
    de la celda según la fuente activa, evitando que se monte sobre celdas vecinas."""
    texto = safestr(str(texto)) if texto is not None else ""
    ancho_max = ancho_celda - margen
    if ancho_max <= 0:
        return ""
    if pdf.get_string_width(texto) <= ancho_max:
        return texto
    recortado = texto
    while recortado and pdf.get_string_width(recortado + "...") > ancho_max:
        recortado = recortado[:-1]
    return (recortado + "...") if recortado else ""


def _clasificar_subtipo_instalacion(row):
    txt = (str(row.get('ACTIVIDAD', '')) + " " + str(row.get('COMENTARIO', ''))).upper()
    if re.search('ADIC', txt): return 'Instalación Adición'
    if re.search('CAMBIO|MIGRACI', txt): return 'Instalación Cambio de Medio'
    if re.search('RECUP', txt): return 'Recuperado'
    return 'Instalación Nueva'


def _clasificar_subtipo_sopfibra(row):
    if bool(row.get('ES_OFFLINE', False)):
        return 'ONT/ONU Offline'
    com = str(row.get('COMENTARIO', '')).upper()
    if re.search('TV|ANALOGIC|DIGITAL', com):
        return 'Problemas de TV (Analógica o Digital)'
    if re.search('FIBRA|OPTIC|OPTICO|PON|SEÑAL', com):
        return 'Problemas de Fibra / Óptica'
    return 'Otro Diagnóstico'


def _dibujar_mini_tabla(pdf, x, y, w, titulo, filas, total_label="Total"):
    """Dibuja una mini-tabla tipo resumen (Categoría / Pendiente / Total) como en el
    panel izquierdo del reporte de referencia."""
    pdf.set_xy(x, y)
    pdf.set_font("Helvetica", "B", 8)
    pdf.set_text_color(40, 50, 100)
    pdf.set_fill_color(235, 238, 245)
    pdf.cell(w, 6, safestr(titulo), border=1, align="L", fill=True)
    pdf.set_xy(x, pdf.get_y() + 6)

    pdf.set_x(x)
    pdf.set_font("Helvetica", "B", 7)
    pdf.set_fill_color(245, 245, 245)
    pdf.set_text_color(0, 0, 0)
    w_label = w - 30
    pdf.cell(w_label, 5, "Categoria", border=1, fill=True)
    pdf.cell(30, 5, "Pendiente", border=1, align="C", fill=True)
    pdf.set_xy(x, pdf.get_y() + 5)

    total = 0
    pdf.set_font("Helvetica", "", 7)
    for etiqueta, cantidad in filas:
        pdf.set_x(x)
        pdf.set_fill_color(255, 255, 255)
        pdf.cell(w_label, 5, safestr(str(etiqueta))[:38], border=1)
        pdf.cell(30, 5, str(int(cantidad)), border=1, align="C")
        pdf.set_xy(x, pdf.get_y() + 5)
        total += cantidad

    pdf.set_x(x)
    pdf.set_font("Helvetica", "B", 7)
    pdf.set_fill_color(220, 230, 245)
    pdf.cell(w_label, 5, total_label, border=1, fill=True)
    pdf.cell(30, 5, str(int(total)), border=1, align="C", fill=True)
    pdf.set_y(pdf.get_y() + 7)
    return pdf.get_y()


def generar_pdf_pendientes_dispatch(df_totales, df_detalle, hoy_str):
    pdf = ReporteGenerencialPDF()
    pdf.alias_nb_pages()
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(40, 50, 100)
    pdf.cell(0, 10, safestr("REPORTE DE PENDIENTES GENERALES (DISPATCH)"), border=0, ln=True, align="C")

    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 6, safestr(f"Corte Operativo del Día: {hoy_str}"), ln=True, align="C")
    pdf.ln(8)

    # --- Preparación de la base completa (asignadas + sin asignar) con columna de días ---
    df_detalle = df_detalle.copy()
    if 'DIAS_RETRASO' not in df_detalle.columns:
        df_detalle['DIAS_RETRASO'] = 0
    df_detalle['DIAS_RETRASO'] = pd.to_numeric(df_detalle['DIAS_RETRASO'], errors='coerce').fillna(0).astype(int)
    if 'ES_OFFLINE' not in df_detalle.columns:
        df_detalle['ES_OFFLINE'] = False

    # --- 1) Caja de Días (>=7 / >=4 / >=1 / =0) ---
    y_top = pdf.get_y()
    bucket_7 = int((df_detalle['DIAS_RETRASO'] >= 7).sum())
    bucket_4 = int(((df_detalle['DIAS_RETRASO'] >= 4) & (df_detalle['DIAS_RETRASO'] < 7)).sum())
    bucket_1 = int(((df_detalle['DIAS_RETRASO'] >= 1) & (df_detalle['DIAS_RETRASO'] < 4)).sum())
    bucket_0 = int((df_detalle['DIAS_RETRASO'] == 0).sum())

    pdf.set_font("Helvetica", "B", 8)
    pdf.set_fill_color(235, 238, 245)
    pdf.set_text_color(40, 50, 100)
    pdf.set_xy(10, y_top)
    pdf.cell(30, 5, "Días", border=1, align="C", fill=True)
    pdf.cell(20, 5, "Cant", border=1, align="C", fill=True)
    pdf.set_xy(10, y_top + 5)

    buckets = [
        (">= 7 Dias", bucket_7, (211, 47, 47), (255, 255, 255)),
        (">= 4 Dias", bucket_4, (245, 124, 0), (255, 255, 255)),
        (">= 1 Dias", bucket_1, (251, 192, 45), (0, 0, 0)),
        (">= 0 Dias", bucket_0, (56, 142, 60), (255, 255, 255)),
    ]
    pdf.set_font("Helvetica", "B", 7)
    for etiqueta, cant, fill, txt in buckets:
        pdf.set_x(10)
        pdf.set_fill_color(*fill)
        pdf.set_text_color(*txt)
        pdf.cell(30, 5, etiqueta, border=1, fill=True)
        pdf.cell(20, 5, str(cant), border=1, align="C", fill=True)
        pdf.set_xy(10, pdf.get_y() + 5)

    pdf.set_text_color(0, 0, 0)
    y_after_bucket = max(pdf.get_y() + 3, y_top)

    # --- 2) RESUMEN DE CARGA PARA EL SIGUIENTE TURNO ---
    pdf.set_xy(65, y_top)
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(84, 98, 143)
    pdf.cell(0, 6, safestr("RESUMEN DE CARGA PARA EL SIGUIENTE TURNO"), ln=True, align="L")
    pdf.set_x(65)

    pdf.set_fill_color(240, 240, 240)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "B", 8)
    pdf.cell(50, 7, "Clasificacion", border=1, fill=True)
    pdf.cell(30, 7, "Asignadas", border=1, align="C", fill=True)
    pdf.cell(30, 7, "Sin Asignar", border=1, align="C", fill=True)
    pdf.cell(30, 7, "Total", border=1, align="C", fill=True)
    pdf.set_xy(65, pdf.get_y() + 7)

    for _, row in df_totales.iterrows():
        if row['Categoría'] == 'TOTAL PENDIENTES':
            pdf.set_font("Helvetica", "B", 8)
            pdf.set_fill_color(220, 230, 245)
            fill = True
        else:
            pdf.set_font("Helvetica", "", 8)
            fill = False
        pdf.set_x(65)
        pdf.cell(50, 6, safestr(row['Categoría'])[:28], border=1, fill=fill)
        pdf.cell(30, 6, str(row['Asignadas (En Ruta)']), border=1, align="C", fill=fill)
        pdf.cell(30, 6, str(row['Nuevas (Sin Asignar)']), border=1, align="C", fill=fill)
        pdf.cell(30, 6, str(row['TOTAL GENERAL']), border=1, align="C", fill=fill)
        pdf.set_xy(65, pdf.get_y() + 6)

    y_after_resumen = pdf.get_y() + 4
    pdf.set_y(max(y_after_bucket, y_after_resumen))
    pdf.ln(4)

    # --- 3) Mini-tablas de desglose (FTTH/FIBRA, SOPFIBRA, Resto de Actividades) ---
    act_upper = df_detalle['ACTIVIDAD'].astype(str).str.upper()
    com_upper = df_detalle.get('COMENTARIO', pd.Series('', index=df_detalle.index)).astype(str).str.upper()
    txt_completo = act_upper + " " + com_upper

    mask_ins = txt_completo.str.contains('INS|NUEVA|ADIC|CAMBIO|MIGRACI|RECUP', na=False) & ~act_upper.str.contains('SOP|FALLA|MANT', na=False)
    df_ins = df_detalle[mask_ins].copy()

    mask_sop = act_upper.str.contains('SOP|FALLA|MANT', na=False)
    df_sop = df_detalle[mask_sop].copy()

    mask_resto = ~mask_ins & ~mask_sop
    df_resto = df_detalle[mask_resto].copy()

    y_tablas = pdf.get_y()

    if not df_ins.empty:
        df_ins['SUBTIPO'] = df_ins.apply(_clasificar_subtipo_instalacion, axis=1)
        filas_ins = list(df_ins['SUBTIPO'].value_counts().items())
    else:
        filas_ins = []
    y1 = _dibujar_mini_tabla(pdf, 10, y_tablas, 62, "FTTH / FIBRA", filas_ins)

    if not df_sop.empty:
        df_sop['SUBTIPO'] = df_sop.apply(_clasificar_subtipo_sopfibra, axis=1)
        filas_sop = list(df_sop['SUBTIPO'].value_counts().items())
    else:
        filas_sop = []
    y2 = _dibujar_mini_tabla(pdf, 76, y_tablas, 62, "SOPFIBRA", filas_sop)

    if not df_resto.empty:
        filas_resto = list(df_resto['ACTIVIDAD'].astype(str).value_counts().items())
    else:
        filas_resto = []
    y3 = _dibujar_mini_tabla(pdf, 142, y_tablas, 58, "Resto de las Actividades", filas_resto)

    pdf.set_y(max(y1, y2, y3) + 4)

    # --- 4) Listado prioritario: órdenes nuevas sin asignar ---
    mask_sin_tec = (df_detalle['TECNICO'].isna()) | (df_detalle['TECNICO'].astype(str).str.strip() == '') | (df_detalle['TECNICO'].astype(str).str.upper().isin(['NONE', 'NAN', 'N/D', 'NULL']))
    df_no_asig = df_detalle[mask_sin_tec].copy()

    if pdf.get_y() > 230:
        pdf.add_page()

    if not df_no_asig.empty:
        pdf.seccion_titulo("LISTADO PRIORITARIO: ORDENES NUEVAS (SIN ASIGNAR)")
        pdf.set_fill_color(255, 235, 235)
        pdf.set_text_color(50, 50, 50)
        pdf.set_font("Helvetica", "B", 8)
        pdf.cell(20, 6, "Orden", border=1, align="C", fill=True)
        pdf.cell(30, 6, "Cliente", border=1, align="C", fill=True)
        pdf.cell(60, 6, "Actividad", border=1, align="C", fill=True)
        pdf.cell(70, 6, "Colonia", border=1, align="C", fill=True)
        pdf.ln()

        pdf.set_font("Helvetica", "", 7)
        pdf.set_text_color(0, 0, 0)
        for _, row in df_no_asig.iterrows():
            if pdf.get_y() > 275:
                pdf.add_page()
            pdf.cell(20, 5, safestr(str(row['NUM'])), border=1, align="C")
            pdf.cell(30, 5, safestr(str(row['CLIENTE'])), border=1, align="C")
            pdf.cell(60, 5, safestr(str(row['ACTIVIDAD']))[:35], border=1, align="L")
            pdf.cell(70, 5, safestr(str(row.get('COLONIA', '')))[:40], border=1, align="L")
            pdf.ln()
    else:
        pdf.seccion_titulo("LISTADO PRIORITARIO: ORDENES NUEVAS (SIN ASIGNAR)")
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(0, 100, 0)
        pdf.cell(0, 6, "Excelente. Todas las ordenes se encuentran asignadas a tecnicos.", ln=True)

    # --- 5) Listado general detallado (formato ancho, tipo hoja de cálculo) ---
    if not df_detalle.empty:
        pdf.add_page(orientation='L')
        pdf.seccion_titulo("LISTADO GENERAL DETALLADO: TODAS LAS ORDENES PENDIENTES")

        columnas = [
            ("Días", 9), ("NUM", 17), ("ACTIVIDAD", 30), ("MOTIVO UNIFICADO", 32),
            ("SECTOR", 20), ("COLONIA", 30), ("CLIENTE", 17), ("CONTRATO", 19),
            ("NOMBRE DEL CLIENTE", 43), ("TECNICO", 34), ("SUBESTADO", 23),
        ]
        # Suma de anchos = 274mm; el área útil en A4 horizontal con márgenes
        # de 10mm por lado es de ~277mm, así que queda margen sin desbordar.

        def _dibujar_encabezado():
            pdf.set_font("Helvetica", "B", 6.5)
            pdf.set_text_color(50, 50, 50)
            pdf.set_fill_color(240, 240, 240)
            for titulo, ancho in columnas:
                pdf.cell(ancho, 6, _ajustar_texto_a_celda(pdf, titulo, ancho), border=1, align="C", fill=True)
            pdf.ln()

        _dibujar_encabezado()

        df_ordenado = df_detalle.sort_values(by='DIAS_RETRASO', ascending=False)

        for _, row in df_ordenado.iterrows():
            if pdf.get_y() > 190:
                pdf.add_page(orientation='L')
                _dibujar_encabezado()

            dias_val = int(row['DIAS_RETRASO'])
            estado_asignado = not bool(mask_sin_tec.loc[row.name]) if row.name in mask_sin_tec.index else True
            subestado = _valor_flexible(row, ['SUBESTADO', 'SUB ESTADO', 'SUB-ESTADO'], default=("ASIGNADA TECNICO" if estado_asignado else "EN PROCESO"))
            contrato = _valor_flexible(row, ['CONTRATO', 'CONTRATO FÍSICO', 'CONTRATO FISICO', 'CONTRATO_FISICO', 'NO. CONTRATO'])
            motivo = _valor_flexible(row, ['MOTIVO UNIFICADO', 'MOTIVO'])
            sector = _valor_flexible(row, ['SECTOR'])
            nombre_cliente = _valor_flexible(row, ['NOMBRE', 'NOMBRE CLIENTE', 'SUSCRIPTOR'])

            valores_crudos = [
                (str(dias_val), "C", True),
                (row.get('NUM', ''), "C", False),
                (row.get('ACTIVIDAD', ''), "L", False),
                (motivo, "L", False),
                (sector, "L", False),
                (row.get('COLONIA', ''), "L", False),
                (row.get('CLIENTE', ''), "C", False),
                (contrato, "L", False),
                (nombre_cliente, "L", False),
                (row.get('TECNICO', ''), "L", False),
                (subestado, "L", False),
            ]

            pdf.set_font("Helvetica", "", 6.3)
            for (valor, alineacion, es_dias), (titulo, ancho) in zip(valores_crudos, columnas):
                texto = _ajustar_texto_a_celda(pdf, valor, ancho)
                if es_dias:
                    _color_por_dias(pdf, dias_val)
                    pdf.cell(ancho, 5, texto, border=1, align=alineacion, fill=True)
                    pdf.set_text_color(0, 0, 0)
                else:
                    pdf.cell(ancho, 5, texto, border=1, align=alineacion)
            pdf.ln()

    pdf.set_text_color(0, 0, 0)
    return finalizar_pdf(pdf)

def generar_pdf_tiempos_muertos(df_dia, fecha_sel):
    pdf = ReporteGenerencialPDF(orientation='L', unit='mm', format='A4')
    pdf.alias_nb_pages()
    pdf.add_page()
    
    pdf.set_font("Helvetica", 'B', 14)
    pdf.set_text_color(40, 50, 100)
    pdf.cell(0, 10, safestr(f"REPORTE DE EFICIENCIA OPERATIVA Y TIEMPO PERDIDO - {fecha_sel.strftime('%d/%m/%Y')}"), ln=True, align='C')
    pdf.ln(5)
    
    df_valido = df_dia.dropna(subset=['HORA_INI', 'TECNICO']).copy()
    
    if df_valido.empty:
        pdf.set_font("Helvetica", '', 12)
        pdf.cell(0, 10, "No hay datos operativos para calcular tiempos.", ln=True, align='C')
        return finalizar_pdf(pdf)

    df_valido['HORA_INI'] = pd.to_datetime(df_valido['HORA_INI'])
    df_valido['HORA_LIQ'] = pd.to_datetime(df_valido['HORA_LIQ'])
    tecnicos = sorted(df_valido['TECNICO'].astype(str).unique())
    
    ahora_hx = get_honduras_time()
    inicio_jornada = pd.Timestamp.combine(fecha_sel, dt_time(8, 0))
    fin_jornada = pd.Timestamp.combine(fecha_sel, dt_time(17, 0))
    limite_evaluacion = min(ahora_hx, fin_jornada) if fecha_sel == ahora_hx.date() else fin_jornada
    
    for tec in tecnicos:
        df_tec = df_valido[df_valido['TECNICO'] == tec].sort_values(by='HORA_INI')
        
        pdf.set_font("Helvetica", 'B', 10)
        pdf.set_fill_color(220, 230, 250)
        pdf.cell(0, 8, safestr(f" TECNICO: {tec}"), border=1, ln=True, fill=True)
        
        pdf.set_font("Helvetica", 'B', 8)
        pdf.set_fill_color(240, 240, 240)
        pdf.cell(25, 6, "ORDEN", border=1, align='C', fill=True)
        pdf.cell(100, 6, "ACTIVIDAD", border=1, align='C', fill=True)
        pdf.cell(25, 6, "INICIO", border=1, align='C', fill=True)
        pdf.cell(25, 6, "FIN", border=1, align='C', fill=True)
        pdf.cell(25, 6, "DURACION", border=1, align='C', fill=True)
        pdf.ln()
        
        total_minutos_trabajados = 0
        tiempo_muerto_acumulado = 0
        cursor_tiempo = inicio_jornada 
        
        pdf.set_font("Helvetica", '', 8)
        for _, row in df_tec.iterrows():
            num = str(row.get('NUM', 'N/D'))
            act = str(row.get('ACTIVIDAD', ''))[:55]
            
            h_ini_dt = row['HORA_INI']
            h_liq_dt = row['HORA_LIQ']
            
            h_ini_str = h_ini_dt.strftime('%H:%M') if pd.notnull(h_ini_dt) else "N/D"
            h_fin_str = h_liq_dt.strftime('%H:%M') if pd.notnull(h_liq_dt) else "En curso"
            
            duracion_str = "---"
            if pd.notnull(h_ini_dt):
                if pd.notnull(h_liq_dt): fin_real = h_liq_dt
                else: fin_real = ahora_hx if h_ini_dt.date() == ahora_hx.date() else fin_jornada
                    
                if fin_real > h_ini_dt:
                    mins_reales = (fin_real - h_ini_dt).total_seconds() / 60
                    if mins_reales > 0:
                        total_minutos_trabajados += mins_reales
                        hrs_d, mins_d = divmod(mins_reales, 60)
                        duracion_str = f"{int(hrs_d)}h {int(mins_d)}m"

                if h_ini_dt > cursor_tiempo and cursor_tiempo < limite_evaluacion:
                    gap_end = min(h_ini_dt, limite_evaluacion)
                    gap_mins = (gap_end - cursor_tiempo).total_seconds() / 60
                    if gap_mins > 0: tiempo_muerto_acumulado += gap_mins
                
                fin_orden_gap = h_liq_dt if pd.notnull(h_liq_dt) else ahora_hx
                if pd.notnull(fin_orden_gap): cursor_tiempo = max(cursor_tiempo, fin_orden_gap)

            pdf.cell(25, 6, safestr(num), border=1, align='C')
            pdf.cell(100, 6, safestr(act), border=1)
            pdf.cell(25, 6, safestr(h_ini_str), border=1, align='C')
            pdf.cell(25, 6, safestr(h_fin_str), border=1, align='C')
            pdf.cell(25, 6, safestr(duracion_str), border=1, align='C')
            pdf.ln()
            
        if cursor_tiempo < limite_evaluacion:
            gap_mins = (limite_evaluacion - cursor_tiempo).total_seconds() / 60
            if gap_mins > 0: tiempo_muerto_acumulado += gap_mins
        
        tiempo_perdido_mins = max(0, tiempo_muerto_acumulado - 60)
        hrs_t, mins_t = divmod(total_minutos_trabajados, 60)
        hrs_p, mins_p = divmod(tiempo_perdido_mins, 60)
        
        pdf.set_font("Helvetica", 'B', 8)
        pdf.cell(175, 6, "TOTAL TIEMPO TRABAJADO EN ORDENES (Incluye Extras):", border=1, align='R')
        pdf.set_text_color(0, 100, 0)
        pdf.cell(25, 6, safestr(f"{int(hrs_t)}h {int(mins_t)}m"), border=1, align='C')
        pdf.set_text_color(0, 0, 0)
        pdf.ln()
        
        pdf.cell(175, 6, "TIEMPO PERDIDO / MUERTO (Base Brechas 8am-5pm - Almuerzo):", border=1, align='R')
        if tiempo_perdido_mins > 0: pdf.set_text_color(200, 0, 0)
        pdf.cell(25, 6, safestr(f"{int(hrs_p)}h {int(mins_p)}m"), border=1, align='C')
        pdf.set_text_color(0, 0, 0)
        pdf.ln()
        pdf.ln(5)

    return finalizar_pdf(pdf)

def generar_pdf_promedio_arranque(df_promedios, f_inicio, f_fin):
    pdf = ReporteGenerencialPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    
    pdf.set_font("Helvetica", 'B', 14)
    pdf.set_text_color(40, 50, 100)
    pdf.cell(0, 10, safestr("PROMEDIO DE ARRANQUE DE JORNADA"), ln=True, align='C')
    
    pdf.set_font("Helvetica", '', 10)
    pdf.set_text_color(100, 100, 100)
    inicio_str = f_inicio.strftime('%d/%m/%Y') if hasattr(f_inicio, 'strftime') else str(f_inicio)
    fin_str = f_fin.strftime('%d/%m/%Y') if hasattr(f_fin, 'strftime') else str(f_fin)
    pdf.cell(0, 6, safestr(f"Periodo: {inicio_str} al {fin_str}"), ln=True, align='C')
    pdf.ln(10)
    
    if not df_promedios.empty:
        pdf.set_x(15)
        pdf.set_fill_color(220, 230, 250)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", 'B', 9)
        pdf.cell(90, 8, "TECNICO", border=1, align='C', fill=True)
        pdf.cell(40, 8, "DIAS EVALUADOS", border=1, align='C', fill=True)
        pdf.cell(50, 8, "HORA PROMEDIO", border=1, align='C', fill=True)
        pdf.ln()
        
        pdf.set_font("Helvetica", '', 9)
        for _, row in df_promedios.iterrows():
            pdf.set_x(15)
            tec = str(row['TECNICO'])[:45]
            dias = str(row['Dias_Computados'])
            hora = str(row['Hora_Promedio_Inicio'])
            
            pdf.cell(90, 7, safestr(tec), border=1, align='L')
            pdf.cell(40, 7, safestr(dias), border=1, align='C')
            pdf.cell(50, 7, safestr(hora), border=1, align='C')
            pdf.ln()
            
    return finalizar_pdf(pdf)

def generar_pdf_unificado_rrhh(df_ausencias, df_tardanzas):
    pdf = ReporteGenerencialPDF(orientation='L', unit='mm', format='A4')
    pdf.alias_nb_pages()
    pdf.add_page()
    
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(40, 50, 100)
    pdf.cell(0, 12, safestr("REPORTE UNIFICADO RRHH: CONTROL DE ASISTENCIA"), ln=True, align="C")
    
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 6, safestr(f"Fecha de emision: {datetime.now().strftime('%d/%m/%Y %I:%M %p')}"), ln=True, align="C")
    pdf.ln(10)
    
    def dibujar_tabla(pdf_obj, df, titulo):
        pdf_obj.set_font("Helvetica", "B", 12)
        pdf_obj.set_text_color(40, 50, 100)
        pdf_obj.cell(0, 10, safestr(titulo), ln=True)
        
        if df.empty:
            pdf_obj.set_font("Helvetica", "I", 10)
            pdf_obj.set_text_color(150, 0, 0)
            pdf_obj.cell(0, 8, "No se registraron datos en esta categoria.", ln=True)
            pdf_obj.ln(5)
            return
            
        cols_deseadas = ['Nombre completo', 'Departamento', 'Fecha', 'Horario', 'Hora de inicio del trabajo', 'Hora final del trabajo']
        cols_finales = [c for c in cols_deseadas if c in df.columns]
        if not cols_finales: cols_finales = list(df.columns)[:6]
            
        df_sub = df[cols_finales]
        pdf_obj.set_font("Helvetica", "B", 8)
        pdf_obj.set_fill_color(230, 235, 245)
        pdf_obj.set_text_color(0, 0, 0)
        
        ancho_total = 275 
        w = ancho_total / len(cols_finales)
        
        for col in cols_finales:
            pdf_obj.cell(w, 8, safestr(str(col))[:25], border=1, align="C", fill=True)
        pdf_obj.ln()
        
        pdf_obj.set_font("Helvetica", "", 7)
        for _, row in df_sub.iterrows():
            if pdf_obj.get_y() > 185:
                pdf_obj.add_page()
                pdf_obj.set_font("Helvetica", "B", 8)
                pdf_obj.set_fill_color(230, 235, 245)
                for col in cols_finales:
                    pdf_obj.cell(w, 8, safestr(str(col))[:25], border=1, align="C", fill=True)
                pdf_obj.ln()
                pdf_obj.set_font("Helvetica", "", 7)
                
            for col in cols_finales:
                pdf_obj.cell(w, 6, safestr(str(row[col]))[:40], border=1, align="C")
            pdf_obj.ln()
        pdf_obj.ln(12)
        
    dibujar_tabla(pdf, df_ausencias, "1. DETALLE DE AUSENCIAS")
    dibujar_tabla(pdf, df_tardanzas, "2. DETALLE DE LLEGADAS TARDE")
    
    return finalizar_pdf(pdf)

def generar_pdf_infracciones(df_res):
    pdf = ReporteGenerencialPDF(orientation='L', unit='mm', format='A4')
    pdf.alias_nb_pages()
    pdf.add_page()
    
    pdf.set_font("Helvetica", 'B', 16)
    pdf.cell(277, 10, safestr("Resumen Conosolidado de Infracciones Biometricas"), ln=True, align='C')
    pdf.set_font("Helvetica", 'I', 10)
    pdf.cell(277, 6, safestr(f"Generado el: {datetime.now().strftime('%d/%m/%Y')}"), ln=True, align='C')
    pdf.ln(10)

    df_res['Es_Tarde'] = df_res['Tardanza'] == 'Sí'
    df_res['Tiene_Exc_Alm'] = df_res['Almuerzo (min)'] > 60
    df_res['Tiene_Exc_Brk'] = df_res['Break (min)'] > 15

    resumen = df_res.groupby(['ID', 'Empleado']).agg(
        Tardanzas=('Es_Tarde', 'sum'),
        Suma_Tar=('Exceso_Tardanza_min', 'sum'),
        Almuerzos=('Tiene_Exc_Alm', 'sum'),
        Suma_Alm=('Exceso_Alm_min', 'sum'),
        Breaks=('Tiene_Exc_Brk', 'sum'),
        Suma_Brk=('Exceso_Brk_min', 'sum')
    ).reset_index()

    resumen['Total_Faltas'] = resumen['Tardanzas'] + resumen['Almuerzos'] + resumen['Breaks']
    infractores = resumen[resumen['Total_Faltas'] > 0].sort_values(by='Total_Faltas', ascending=False)

    if infractores.empty:
        pdf.set_font("Helvetica", '', 12)
        pdf.cell(277, 10, safestr("Excelente: No se registraron infracciones en este periodo."), ln=True, align='C')
        return finalizar_pdf(pdf)

    pdf.set_font("Helvetica", 'B', 8)
    pdf.set_fill_color(220, 230, 241) 
    
    w_id, w_emp, w_tar, w_ptar, w_alm, w_palm, w_brk, w_pbrk, w_tot = 15, 60, 22, 28, 25, 33, 25, 33, 25
    
    pdf.cell(w_id, 8, "ID", border=1, fill=True, align='C')
    pdf.cell(w_emp, 8, "Empleado", border=1, fill=True, align='C')
    pdf.cell(w_tar, 8, "Tardanzas", border=1, fill=True, align='C')
    pdf.cell(w_ptar, 8, "Prom. Tardanza", border=1, fill=True, align='C')
    pdf.cell(w_alm, 8, "Exc. Almuerzo", border=1, fill=True, align='C')
    pdf.cell(w_palm, 8, "Prom. Exc. Alm.", border=1, fill=True, align='C')
    pdf.cell(w_brk, 8, "Exc. Break", border=1, fill=True, align='C')
    pdf.cell(w_pbrk, 8, "Prom. Exc. Brk.", border=1, fill=True, align='C')
    pdf.cell(w_tot, 8, "TOTAL FALTAS", border=1, fill=True, align='C')
    pdf.ln()

    pdf.set_font("Helvetica", '', 8)
    for _, row in infractores.iterrows():
        nombre_corto = str(row['Empleado'])[:35]
        p_tar = f"{int(row['Suma_Tar'] / row['Tardanzas'])} min" if row['Tardanzas'] > 0 else "---"
        p_alm = f"{int(row['Suma_Alm'] / row['Almuerzos'])} min" if row['Almuerzos'] > 0 else "---"
        p_brk = f"{int(row['Suma_Brk'] / row['Breaks'])} min" if row['Breaks'] > 0 else "---"
        
        pdf.cell(w_id, 8, safestr(row['ID']), border=1, align='C')
        pdf.cell(w_emp, 8, safestr(f" {nombre_corto}"), border=1)
        pdf.cell(w_tar, 8, str(int(row['Tardanzas'])), border=1, align='C')
        pdf.cell(w_ptar, 8, safestr(p_tar), border=1, align='C')
        pdf.cell(w_alm, 8, str(int(row['Almuerzos'])), border=1, align='C')
        pdf.cell(w_palm, 8, safestr(p_alm), border=1, align='C')
        pdf.cell(w_brk, 8, str(int(row['Breaks'])), border=1, align='C')
        pdf.cell(w_pbrk, 8, safestr(p_brk), border=1, align='C')
        
        pdf.set_font("Helvetica", 'B', 9)
        pdf.cell(w_tot, 8, str(int(row['Total_Faltas'])), border=1, align='C')
        pdf.set_font("Helvetica", '', 8)
        pdf.ln()

    return finalizar_pdf(pdf)

# ==============================================================================
# 5. UTILIDADES Y PROCESAMIENTO GENERAL
# ==============================================================================
def calcular_offline_y_alertas(df_act: pd.DataFrame) -> pd.DataFrame:
    """
    Calcula las columnas ES_OFFLINE y ALERTA_TIEMPO sobre un DataFrame de
    órdenes ya normalizado (columnas ACTIVIDAD, ESTADO, TECNICO, COMENTARIO,
    HORA_INI, HORA_LIQ presentes). Se extrajo como función compartida para
    que tanto la carga manual (cargar_y_limpiar_crudos_diamante_monitor) como
    el flujo automático (sync_job.py) apliquen exactamente la misma lógica
    de detección de equipos caídos (offline) y demoras de atención SOP.
    """
    df_act = df_act.copy()
    ahora_momento_ts = pd.Timestamp(get_honduras_time())

    act_upper = df_act['ACTIVIDAD'].fillna('').astype(str).str.upper()
    est_upper = df_act['ESTADO'].fillna('').astype(str).str.upper().str.strip()
    tec_upper = df_act['TECNICO'].fillna('').astype(str).str.upper().str.strip()
    com_upper = df_act['COMENTARIO'].fillna('').astype(str).str.upper()

    mins_diff = (ahora_momento_ts - df_act['HORA_INI']).dt.total_seconds() / 60
    mask_sop = act_upper.str.contains('SOPFIBRA', regex=True)
    mask_falsos = act_upper.str.contains('PLEXISCA|PEXTERNO|SPLITTEROPT|PLEX|INS|NUEVA|ADIC|CAMBIO|RECU|TVADICIONAL|MIGRACI', regex=True)

    df_act['ALERTA_TIEMPO'] = (
        (df_act['HORA_INI'].notnull()) & (df_act['HORA_LIQ'].isnull()) &
        (mins_diff > 120) & (est_upper != 'CERRADA') & mask_sop & ~mask_falsos
    )

    mask_tec_valido = tec_upper != 'JOSUE MIGUEL SAUCEDA'
    mask_est_abierto = est_upper != 'CERRADA'
    mask_com_off = com_upper.str.contains("ONU OFFLINE|OFF LINE|OFFLINE|LOS EN ROJO|PON ROJO", regex=True)
    mask_precisa = com_upper.apply(es_offline_preciso)

    df_act['ES_OFFLINE'] = (mask_tec_valido & mask_est_abierto & mask_sop & ~mask_falsos & (mask_com_off | mask_precisa))
    return df_act


def es_offline_preciso(comentario):
    txt = str(comentario).upper().strip()
    if not txt or txt == 'NAN': return False
    jergasolucion = ['OK', 'LISTO', 'RECUPERADO', 'SOLUCIONADO', 'NAVEGA', 'YA QUEDO', 'ARRIBA', 'FUNCIONAL', 'ONLINE']
    if any(word in txt for word in jergasolucion): return False
    # CORRECCIÓN: Dejamos estrictamente términos de equipo caído
    keywordsfalla = ['OFFLINE', 'OFF LINE', 'LOS RED', 'PON ROJO', 'LOS EN ROJO', 'EQUIPO OFFLINE', 'ONU OFFLINE', 'ONT OFFLINE']
    return any(word in txt for word in keywordsfalla)

def depurar_archivos_en_crudo(fileactividades, filedispositivos):
    try:
        xlact = pd.ExcelFile(fileactividades, engine='openpyxl')
        sheetp = 'Prueba' if 'Prueba' in xlact.sheet_names else xlact.sheet_names[0]
        dfpraw = pd.read_excel(xlact, sheet_name=sheetp)
        sheethnom = 'HistoricoNoInstaladas' if 'HistoricoNoInstaladas' in xlact.sheet_names else None
        dfhraw = pd.read_excel(xlact, sheet_name=sheethnom) if sheethnom else pd.DataFrame()
        if filedispositivos.name.lower().endswith('.csv'):
            dfdispfull = pd.read_csv(filedispositivos, sep=None, engine='python')
        else:
            dfdispfull = pd.read_excel(filedispositivos, engine='openpyxl')
        dfdispref = pd.DataFrame()
        coltec = [c for c in dfdispfull.columns if any(x in str(c).upper() for x in['TECNICO', 'USER', 'OPERADOR'])]
        colmx = [c for c in dfdispfull.columns if any(x in str(c).upper() for x in['MX', 'VEHICULO', 'PLACA'])]
        dfdispref['TECREF'] = dfdispfull[coltec[0]].astype(str).str.strip().str.upper() if coltec else "N/D"
        dfdispref['MXREF'] = dfdispfull[colmx[0]].astype(str).str.strip() if colmx else "N/D"
        dfp = procesar_dataframe_base(dfpraw)
        dfp['TECKEY'] = dfp['TECNICO'].astype(str).str.strip().str.upper()
        dffinal = dfp.merge(dfdispref.drop_duplicates('TECREF'), left_on='TECKEY', right_on='TECREF', how='left')
        if 'MXREF' in dffinal.columns:
            dffinal['MX'] = dffinal['MXREF'].combine_first(dffinal.get('MX', pd.Series(dtype=str)))
        return dffinal.drop(columns=['TECKEY', 'TECREF', 'MXREF'], errors='ignore'), procesar_dataframe_base(dfhraw)
    except Exception as e:
        raise Exception(f"Error en cruce: {str(e)}")

def procesar_dataframe_base(df):
    df.columns = df.columns.astype(str).str.strip()
    mapeocolumnas = {}
    for nombreinterno, listaopciones in COLUMNS_MAPPING.items():
        for opcion in listaopciones:
            if opcion.upper() in [str(c).upper() for c in df.columns]:
                realname = next(c for c in df.columns if str(c).upper() == opcion.upper())
                mapeocolumnas[realname] = nombreinterno
                break
    df = df.rename(columns=mapeocolumnas)

    # =========================================================================
    # 🏢 FILTRO POR EMPRESA: Solo descartar registros de OTRA empresa explícita
    # (ej. Cable Color). No se descartan filas con EMPRESA vacía/nula, ya que
    # eso es común en órdenes recién creadas o aún sin técnico asignado, y la
    # API (usuario xMonitorIsca) ya viene pre-filtrada a ISCA-MAXCOM del lado
    # del servidor de Cepheus.
    # =========================================================================
    if 'EMPRESA' in df.columns:
        empresa_upper = df['EMPRESA'].astype(str).str.strip().str.upper()
        mask_otra_empresa = (empresa_upper != '') & (empresa_upper != 'NAN') & (empresa_upper != 'NONE') & (~empresa_upper.str.contains('ISCA', na=False))
        df = df[~mask_otra_empresa].copy()

    for colv in COLUMNAS_VITALES_SISTEMA:
        if colv not in df.columns: df[colv] = "N/D"
        
    for cstr in ['ESTADO', 'ACTIVIDAD', 'COMENTARIO', 'CLIENTE', 'TECNICO']:
        df[cstr] = df[cstr].astype(str).replace(['nan', 'None'], 'N/D')
        
    # =========================================================================
    # 🧹 NUEVO: DEPURACIÓN RADICAL DESDE LA LECTURA DEL EXCEL
    # Elimina ACTUALIZARDATOSTECNICOS y otras basuras antes de que toquen la nube
    # =========================================================================
    if 'ACTIVIDAD' in df.columns:
        actividades_basura = [
            'ACTUALIZARDATOSTECNICOS', 
            'ACTUALIZACIONDATOS', 
            'ACTUALIZACIOFW', 
            'ACTUALIZAINFOTECNICA', 
            'ACTUALIZARSENSOR'
        ]
        # Filtramos buscando coincidencias exactas o parciales para mayor seguridad
        mask_basura = df['ACTIVIDAD'].astype(str).str.strip().str.upper().isin(actividades_basura)
        df = df[~mask_basura].copy()
        
    return df

def generar_tablas_gerenciales(df_crudo):
    df = df_crudo.copy()
    df['HORA_INI'] = df['HORA_INI'].apply(parse_date_ultra_safe)
    df['HORA_LIQ'] = df['HORA_LIQ'].apply(parse_date_ultra_safe)
    df = df.dropna(subset=['HORA_INI', 'HORA_LIQ'])
    df['FECHA'] = df['HORA_LIQ'].dt.date
    totales_tec = df.groupby('TECNICO').size().reset_index(name='Total_Tecnico')
    conteo_act = df.groupby(['TECNICO', 'ACTIVIDAD']).size().reset_index(name='Cantidad')
    tabla_produccion = pd.merge(conteo_act, totales_tec, on='TECNICO')
    tabla_produccion['Participacion_%'] = (tabla_produccion['Cantidad'] / tabla_produccion['Total_Tecnico'] * 100).round(1)

    df['MINUTOS'] = (df['HORA_LIQ'] - df['HORA_INI']).dt.total_seconds() / 60
    df.loc[df['MINUTOS'] <= 0, 'MINUTOS'] = None 
    tabla_eficiencia = df.groupby(['TECNICO', 'ACTIVIDAD'])['MINUTOS'].mean().reset_index()
    tabla_eficiencia.columns = ['TECNICO', 'ACTIVIDAD', 'Promedio_Minutos']
    tabla_eficiencia['Promedio_Minutos'] = tabla_eficiencia['Promedio_Minutos'].round(1)

    jornada = df.groupby(['TECNICO', 'FECHA']).agg(Hora_Apertura=('HORA_INI', 'min'), Hora_Cierre=('HORA_LIQ', 'max'), Total_Ordenes=('NUM', 'count')).reset_index()
    jornada['Horas_En_Calle'] = (jornada['Hora_Cierre'] - jornada['Hora_Apertura']).dt.total_seconds() / 3600
    jornada.loc[jornada['Horas_En_Calle'] <= 0, 'Horas_En_Calle'] = None

    resumen_jornada = jornada.groupby('TECNICO').agg(Promedio_Horas_Dia=('Horas_En_Calle', 'mean'), Dias_Laborados=('FECHA', 'nunique'), Max_Horas_Dia=('Horas_En_Calle', 'max')).reset_index()
    resumen_jornada['Promedio_Horas_Dia'] = resumen_jornada['Promedio_Horas_Dia'].round(2)
    resumen_jornada['Max_Horas_Dia'] = resumen_jornada['Max_Horas_Dia'].round(2)

    return tabla_produccion, tabla_eficiencia, resumen_jornada

def cargar_y_limpiar_crudos_diamante_monitor(file_activ, file_dispos):
    try:
        if isinstance(file_dispos, bytes):
            file_dispos_obj = io.BytesIO(file_dispos)
            file_dispos_obj.name = "FttxActiveDevice_cached.xlsx"
        elif hasattr(file_dispos, 'read'): file_dispos.seek(0); file_dispos_obj = file_dispos
        else: file_dispos_obj = file_dispos

        if hasattr(file_activ, 'read'): file_activ.seek(0)
        df_act, df_hst = depurar_archivos_en_crudo(file_activ, file_dispos_obj)
        df_act = procesar_fechas_seguro(df_act, ['HORA_INI', 'HORA_LIQ', 'FECHA_APE'])
        ahora_momento_ts = pd.Timestamp(get_honduras_time())
        fecha_limite_7d_ventana = ahora_momento_ts - timedelta(days=7) 
        mask_vivas_loc = df_act['ESTADO'].astype(str).str.contains(PATRON_ASIGNADAS_VIVA_STR, na=False, case=False)
        df_act = df_act[(df_act['HORA_LIQ'] >= fecha_limite_7d_ventana) | (df_act['FECHA_APE'] >= fecha_limite_7d_ventana) | (df_act['HORA_LIQ'].isna()) | mask_vivas_loc].copy()
        
        df_act['DIAS_RETRASO'] = (ahora_momento_ts.normalize() - df_act['FECHA_APE'].dt.normalize()).dt.days.fillna(0).astype(int)
        df_act.loc[df_act['TECNICO'].str.strip().str.upper() == 'JOSUE MIGUEL SAUCEDA', 'DIAS_RETRASO'] = 0

        act_upper = df_act['ACTIVIDAD'].fillna('').astype(str).str.upper()
        est_upper = df_act['ESTADO'].fillna('').astype(str).str.upper().str.strip()
        tec_upper = df_act['TECNICO'].fillna('').astype(str).str.upper().str.strip()
        com_upper = df_act['COMENTARIO'].fillna('').astype(str).str.upper()
        cli_upper = df_act['CLIENTE'].fillna('').astype(str).str.upper()
        
        mins_diff = (ahora_momento_ts - df_act['HORA_INI']).dt.total_seconds() / 60
        mask_sop = act_upper.str.contains('SOPFIBRA', regex=True)
        mask_falsos = act_upper.str.contains('PLEXISCA|PEXTERNO|SPLITTEROPT|PLEX|INS|NUEVA|ADIC|CAMBIO|RECU|TVADICIONAL|MIGRACI', regex=True)

        df_act['ALERTA_TIEMPO'] = (
            (df_act['HORA_INI'].notnull()) & (df_act['HORA_LIQ'].isnull()) & 
            (mins_diff > 120) & (est_upper != 'CERRADA') & mask_sop & ~mask_falsos
        )
        
        mask_tec_valido = tec_upper != 'JOSUE MIGUEL SAUCEDA'
        mask_est_abierto = est_upper != 'CERRADA'
        mask_com_off = com_upper.str.contains("ONU OFFLINE|OFF LINE|OFFLINE|LOS EN ROJO|PON ROJO", regex=True)
        mask_precisa = com_upper.apply(es_offline_preciso) 
        
        df_act['ES_OFFLINE'] = (mask_tec_valido & mask_est_abierto & mask_sop & ~mask_falsos & (mask_com_off | mask_precisa))
        df_act['MINUTOS_CALC'] = (df_act['HORA_LIQ'] - df_act['HORA_INI']).dt.total_seconds() / 60
        
        texto_seg = act_upper + " " + cli_upper + " " + com_upper
        df_act['SEGMENTO'] = np.where(texto_seg.str.contains('PLEX|PEXTERNO|SPLITTEROPT', regex=True), 'PLEX', 'RESIDENCIAL')
        
        diff_temp = df_act['HORA_LIQ'] - df_act['HORA_INI']
        df_act['TIEMPO_REAL'] = np.where(
            df_act['HORA_INI'].isnull() | df_act['HORA_LIQ'].isnull(),
            "---",
            (diff_temp.dt.total_seconds() // 3600).fillna(0).astype(int).astype(str) + "h " +
            ((diff_temp.dt.total_seconds() % 3600) // 60).fillna(0).astype(int).astype(str) + "m"
        )
        
        return df_act, df_hst
    except Exception as e:
        if st: st.error(f"❌ Error fatal en el motor de depuración: {e}")
        return None, None

# ==============================================================================
# 6. UTILIDADES Y PROCESAMIENTO: AUDITORIA DE VEHICULOS
# ==============================================================================
def forzar_columnas_unicas(df):
    if df is None or df.empty: return df
    df.columns = df.columns.astype(str).str.strip()
    cols = pd.Series(df.columns)
    for dup in cols[cols.duplicated()].unique():
        dup_indices = cols[cols == dup].index.tolist()
        for i, idx in enumerate(dup_indices):
            if i != 0:
                cols.iat[idx] = f"{dup}_{i}"
    df.columns = cols
    return df

def read_file_robust(uploaded_file):
    filename = uploaded_file.name.lower()
    content = uploaded_file.getvalue()
    df = None
    
    if content.startswith(b'\xd0\xcf\x11\xe0'):
        try:
            uploaded_file.seek(0)
            df = pd.read_excel(uploaded_file, engine='xlrd')
        except ImportError:
            if st: st.error("Falta librería xlrd para Excel antiguo.")
    elif b'<table' in content.lower() or b'<html' in content.lower():
        try:
            dfs = pd.read_html(io.StringIO(content.decode('utf-8', errors='ignore')))
            df = max(dfs, key=len)
        except Exception:
            dfs = pd.read_html(io.StringIO(content.decode('latin1', errors='ignore')))
            df = max(dfs, key=len)
    else:
        uploaded_file.seek(0)
        if filename.endswith('.xlsx'): 
            df = pd.read_excel(uploaded_file, engine='openpyxl')
        else:
            try: 
                df = pd.read_csv(uploaded_file, encoding='utf-8', on_bad_lines='skip')
            except UnicodeDecodeError:
                uploaded_file.seek(0)
                df = pd.read_csv(uploaded_file, encoding='latin1', on_bad_lines='skip')

    return forzar_columnas_unicas(df)

def procesar_auditoria_vehiculos(df_input):
    try:
        df = df_input.copy()
        col_placa = next((c for c in df.columns if re.search(r'(?i)PLACA|ALIAS|VEHICULO', str(c))), None)
        if not col_placa:
            for i in range(min(15, len(df))):
                row_str = " ".join([str(x) for x in df.iloc[i].values]).upper()
                if 'PLACA' in row_str or 'VEHICULO' in row_str or 'ALIAS' in row_str:
                    df.columns = [str(x).strip() for x in df.iloc[i].values]
                    df = df.iloc[i+1:].reset_index(drop=True)
                    df = forzar_columnas_unicas(df)
                    break
                    
        col_placa = next((c for c in df.columns if re.search(r'(?i)PLACA|ALIAS|VEHICULO', str(c))), None)
        col_ingreso = next((c for c in df.columns if re.search(r'(?i)HORA.*INGRESO|HORA.*ENTRADA', str(c))), None)
        if not col_ingreso:
            col_ingreso = next((c for c in df.columns if re.search(r'(?i)INGRESO|ENTRADA', str(c)) and not re.search(r'(?i)LAT|LON', str(c))), None)
            
        col_salida = next((c for c in df.columns if re.search(r'(?i)HORA.*SALIDA', str(c))), None)
        if not col_salida:
            col_salida = next((c for c in df.columns if re.search(r'(?i)SALIDA', str(c)) and not re.search(r'(?i)LAT|LON', str(c))), None)
        
        if not (col_placa and col_ingreso and col_salida): 
            return None, "Columnas de Hora o Placa no detectadas correctamente."
            
        df = df.rename(columns={col_placa: '_P', col_ingreso: '_I', col_salida: '_S'})
        df['_P'] = df['_P'].astype(str).str.strip()
        df = df[~df['_P'].isin(['nan', '--', 'None', '', 'Columna'])]
        
        df['_I'] = df['_I'].astype(str).str.replace(r'a\.?\s*m\.?', 'AM', flags=re.I).str.replace(r'p\.?\s*m\.?', 'PM', flags=re.I)
        df['_S'] = df['_S'].astype(str).str.replace(r'a\.?\s*m\.?', 'AM', flags=re.I).str.replace(r'p\.?\s*m\.?', 'PM', flags=re.I)
        
        df['_I'] = pd.to_datetime(df['_I'], dayfirst=True, errors='coerce').fillna(pd.to_datetime(df['_I'], dayfirst=False, errors='coerce'))
        df['_S'] = pd.to_datetime(df['_S'], dayfirst=True, errors='coerce').fillna(pd.to_datetime(df['_S'], dayfirst=False, errors='coerce'))
        
        resumen = df.groupby('_P').agg(P_S=('_S', 'min'), U_E=('_I', 'max')).reset_index()
        
        def calc_tiempo(row):
            ps = row['P_S']
            ue = row['U_E']
            if pd.isnull(ps): return "Sin Salida"
            if pd.isnull(ue): return "Sin Ingreso"
            
            limite_inf = ps.replace(hour=6, minute=30, second=0, microsecond=0)
            limite_sup = ps.replace(hour=23, minute=59, second=59, microsecond=0)
            
            if ps < limite_inf: ps = limite_inf
            if ue > limite_sup: ue = limite_sup
            
            if ue >= ps:
                diff_secs = (ue - ps).total_seconds()
                if diff_secs > 3600: diff_secs -= 3600
                else: diff_secs = 0
                h, r = divmod(int(diff_secs), 3600); m, s = divmod(r, 60)
                return f"{h:02d}:{m:02d}:{s:02d}"
            return "Revisar"
                
        resumen['Tiempo Real en Calle'] = resumen.apply(calc_tiempo, axis=1)
        resumen['Primera Salida'] = resumen['P_S'].dt.strftime('%I:%M %p').fillna("---")
        resumen['Última Entrada'] = resumen['U_E'].dt.strftime('%I:%M %p').fillna("---")
        
        resumen = resumen.rename(columns={'_P': 'Vehículo / Placa'})
        final_df = resumen[['Vehículo / Placa', 'Primera Salida', 'Última Entrada', 'Tiempo Real en Calle']].copy()
        
        return forzar_columnas_unicas(final_df), "OK"
    except Exception as e: return None, str(e)

def procesar_auditoria_semanal(df_input):
    try:
        df = df_input.copy()
        
        # 1. Búsqueda inteligente de columnas
        col_placa = next((c for c in df.columns if re.search(r'(?i)PLACA|ALIAS|VEHICULO', str(c))), None)
        if not col_placa:
            for i in range(min(15, len(df))):
                row_str = " ".join([str(x) for x in df.iloc[i].values]).upper()
                if 'PLACA' in row_str or 'VEHICULO' in row_str or 'ALIAS' in row_str:
                    df.columns = [str(x).strip() for x in df.iloc[i].values]
                    df = df.iloc[i+1:].reset_index(drop=True)
                    df = forzar_columnas_unicas(df)
                    break
                    
        col_placa = next((c for c in df.columns if re.search(r'(?i)PLACA|ALIAS|VEHICULO', str(c))), None)
        col_ingreso = next((c for c in df.columns if re.search(r'(?i)HORA.*INGRESO|HORA.*ENTRADA', str(c))), None)
        if not col_ingreso:
            col_ingreso = next((c for c in df.columns if re.search(r'(?i)INGRESO|ENTRADA', str(c)) and not re.search(r'(?i)LAT|LON', str(c))), None)
            
        col_salida = next((c for c in df.columns if re.search(r'(?i)HORA.*SALIDA', str(c))), None)
        if not col_salida:
            col_salida = next((c for c in df.columns if re.search(r'(?i)SALIDA', str(c)) and not re.search(r'(?i)LAT|LON', str(c))), None)
        
        col_fecha = next((c for c in df.columns if re.search(r'(?i)^FECHA', str(c).strip())), None)
        
        if not (col_placa and col_ingreso and col_salida): 
            return None, None, "Columnas no detectadas.", None, None
            
        df = df.rename(columns={col_placa: '_P', col_ingreso: '_I', col_salida: '_S'})
        df['_P'] = df['_P'].astype(str).str.strip()
        df = df[~df['_P'].isin(['nan', '--', 'None', '', 'Columna'])]
        
        # 2. Limpieza de strings AM/PM
        raw_I = df['_I'].astype(str).str.replace(r'a\.?\s*m\.?', 'AM', flags=re.I).str.replace(r'p\.?\s*m\.?', 'PM', flags=re.I).str.strip()
        raw_S = df['_S'].astype(str).str.replace(r'a\.?\s*m\.?', 'AM', flags=re.I).str.replace(r'p\.?\s*m\.?', 'PM', flags=re.I).str.strip()
        
        # 3. PARSEO INTELIGENTE DE FECHAS (La cura al bug de los meses/días invertidos)
        dt_I_eu = pd.to_datetime(raw_I, format='mixed', dayfirst=True, errors='coerce')
        dt_I_us = pd.to_datetime(raw_I, format='mixed', dayfirst=False, errors='coerce')
        
        # Si el parseo Europeo (DD/MM) distribuye las fechas en meses distintos (> 20 días)
        # pero el US (MM/DD) las mantiene juntitas en la misma semana, entonces el GPS usó formato US.
        if pd.notnull(dt_I_eu.max()) and pd.notnull(dt_I_us.max()):
            if (dt_I_eu.max() - dt_I_eu.min()).days > 20 and (dt_I_us.max() - dt_I_us.min()).days <= 20:
                df['_I'] = dt_I_us
                df['_S'] = pd.to_datetime(raw_S, format='mixed', dayfirst=False, errors='coerce')
            else:
                df['_I'] = dt_I_eu
                df['_S'] = pd.to_datetime(raw_S, format='mixed', dayfirst=True, errors='coerce')
        else:
            df['_I'] = dt_I_eu
            df['_S'] = pd.to_datetime(raw_S, format='mixed', dayfirst=True, errors='coerce')

        # Si existe una columna FECHA explícita, la usamos para sobrescribir y proteger
        if col_fecha:
            dt_F_eu = pd.to_datetime(df[col_fecha], format='mixed', dayfirst=True, errors='coerce')
            dt_F_us = pd.to_datetime(df[col_fecha], format='mixed', dayfirst=False, errors='coerce')
            if pd.notnull(dt_F_eu.max()) and pd.notnull(dt_F_us.max()):
                if (dt_F_eu.max() - dt_F_eu.min()).days > 20 and (dt_F_us.max() - dt_F_us.min()).days <= 20:
                    df['_F_real'] = dt_F_us.dt.date
                else:
                    df['_F_real'] = dt_F_eu.dt.date
            else:
                df['_F_real'] = dt_F_eu.dt.date

            def fusionar_fecha_hora(f_real, dt_time_parsed):
                if pd.isnull(dt_time_parsed): return pd.NaT
                if pd.notnull(f_real): return pd.Timestamp.combine(f_real, dt_time_parsed.time())
                return dt_time_parsed

            df['_I'] = df.apply(lambda row: fusionar_fecha_hora(row.get('_F_real'), row['_I']), axis=1)
            df['_S'] = df.apply(lambda row: fusionar_fecha_hora(row.get('_F_real'), row['_S']), axis=1)
            df['Fecha'] = df.get('_F_real', df['_I'].dt.date).fillna(df['_I'].dt.date).fillna(df['_S'].dt.date)
        else:
            df['Fecha'] = df['_I'].dt.date.fillna(df['_S'].dt.date)
        
        df = df.dropna(subset=['Fecha'])
        if df.empty: return None, None, "No hay fechas válidas en el archivo.", None, None
        
        # --- ELIMINADO EL FILTRO ESTRICTO DE 7 DÍAS QUE BORRABA LA DATA ---
        f_inicio = df['Fecha'].min()
        f_fin = df['Fecha'].max()

        # 4. Agrupación por Vehículo Y Fecha
        diario = df.groupby(['_P', 'Fecha']).agg(P_S=('_S', 'min'), U_E=('_I', 'max')).reset_index()
        
        def calc_segs(row):
            ps = row['P_S']
            ue = row['U_E']
            if pd.isnull(ps) or pd.isnull(ue): return 0
            
            fecha_base = row['Fecha']
            try:
                ps_full = datetime.combine(fecha_base, ps.time())
                ue_full = datetime.combine(fecha_base, ue.time())
            except:
                return 0
            
            limite_inf = ps_full.replace(hour=6, minute=30, second=0, microsecond=0)
            limite_sup = ps_full.replace(hour=23, minute=59, second=59, microsecond=0)
            
            if ps_full < limite_inf: ps_full = limite_inf
            if ue_full > limite_sup: ue_full = limite_sup
            
            if ue_full > ps_full:
                diff = (ue_full - ps_full).total_seconds()
                if diff > 3600: return int(diff - 3600)
                return int(diff)
            return 0

        diario['segundos'] = diario.apply(calc_segs, axis=1)
        
        semanal = diario.groupby('_P').agg(
            Dias_Laborados=('Fecha', 'nunique'),
            Total_Segundos=('segundos', 'sum')
        ).reset_index()

        semanal['Total_Segundos'] = semanal['Total_Segundos'].fillna(0).astype(int)

        dias_reales = diario[diario['segundos'] > 0].groupby('_P').size().reset_index(name='Dias_Efectivos')
        semanal = pd.merge(semanal, dias_reales, on='_P', how='left')
        semanal['Dias_Efectivos'] = semanal['Dias_Efectivos'].fillna(semanal['Dias_Laborados'])
        
        semanal['Prom_Segundos'] = 0
        mask_efectivos = semanal['Dias_Efectivos'] > 0
        
        semanal.loc[mask_efectivos, 'Prom_Segundos'] = (semanal.loc[mask_efectivos, 'Total_Segundos'] / semanal.loc[mask_efectivos, 'Dias_Efectivos']).astype(int)
        semanal['Prom_Segundos'] = semanal['Prom_Segundos'].fillna(0).astype(int)

        def format_segs(secs):
            if pd.isnull(secs) or secs <= 0: return "00:00:00"
            h, r = divmod(int(secs), 3600); m, s = divmod(r, 60)
            return f"{h:02d}:{m:02d}:{s:02d}"

        diario['Primera Salida'] = diario['P_S'].dt.strftime('%I:%M %p').fillna("---")
        diario['Última Entrada'] = diario['U_E'].dt.strftime('%I:%M %p').fillna("---")
        diario['Tiempo Diario'] = diario['segundos'].apply(format_segs)
        diario = diario.rename(columns={'_P': 'Vehículo / Placa'})
        final_diario = diario[['Vehículo / Placa', 'Fecha', 'Primera Salida', 'Última Entrada', 'Tiempo Diario']].copy()

        semanal['Tiempo Total Semana'] = semanal['Total_Segundos'].apply(format_segs)
        semanal['Promedio Diario'] = semanal['Prom_Segundos'].apply(format_segs)
        semanal = semanal.rename(columns={'_P': 'Vehículo / Placa', 'Dias_Laborados': 'Días Trabajados'})
        final_semanal = semanal[['Vehículo / Placa', 'Días Trabajados', 'Tiempo Total Semana', 'Promedio Diario']].copy()
        
        return forzar_columnas_unicas(final_diario), forzar_columnas_unicas(final_semanal), "OK", f_inicio, f_fin
    except Exception as e: return None, None, str(e), None, None

def procesar_matriz_telemetria(df_raw):
    try:
        header_idx = None
        for i in range(min(20, len(df_raw))):
            if any(k in str(df_raw.iloc[i, 0]).upper() for k in ['PLACA', 'ALIAS', 'VEHICULO']):
                header_idx = i; break
        if header_idx is None: return None, "No se encontró encabezado en Estadístico."

        df = df_raw.iloc[header_idx + 1:].copy()
        raw_columns = df_raw.iloc[header_idx].astype(str).str.strip().tolist()
        
        clean_columns = []
        for i, col in enumerate(raw_columns):
            col_str = str(col).strip()
            if col_str.lower() in ['nan', '', 'none']:
                clean_columns.append(f"Info_{i}")
            elif i == 0:
                clean_columns.append(col_str if col_str else "Placa")
            elif i == 1:
                clean_columns.append(col_str if col_str else "Opcion")
            elif 'TOTAL' in col_str.upper():
                clean_columns.append(col_str)
            else:
                try:
                    fecha_obj = pd.to_datetime(col_str, errors='coerce')
                    if pd.notna(fecha_obj):
                        dias_semana = ["Lun", "Mar", "Mié", "Jue", "Vie", "Sáb", "Dom"]
                        nombre_dia = dias_semana[fecha_obj.weekday()]
                        clean_columns.append(f"{nombre_dia} {fecha_obj.strftime('%d/%m')}")
                    else:
                        clean_columns.append(col_str if col_str else f"Dia_{i-1}")
                except:
                    clean_columns.append(col_str if col_str else f"Dia_{i-1}")
        
        df.columns = clean_columns
        df = forzar_columnas_unicas(df)
        
        col_placa = df.columns[0]
        col_opcion = df.columns[1] if len(df.columns) > 1 else None
        
        df = df.dropna(subset=[col_placa])
        df = df[~df[col_placa].astype(str).str.contains('La versión de este equipo', case=False, na=False)]
        
        if col_opcion:
            df = df[~df[col_opcion].astype(str).str.contains('Tiempo', case=False, na=False)]
            
        df = df[df[col_placa].astype(str).str.strip() != ''].fillna(0)

        col_total = next((c for c in df.columns if 'TOTAL' in str(c).upper()), None)
        if col_total:
            df[col_total] = pd.to_numeric(df[col_total], errors='coerce').fillna(0)
            df = df[df[col_total] > 0].copy()

        return df, "OK"
    except Exception as e: return None, str(e)

def generar_pdf_auditoria_tiempos(df_resumen):
    pdf = ReporteGenerencialPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(84, 98, 143)
    pdf.cell(0, 10, safestr(f" Auditoria de Tiempos Diario - {get_hn_time().strftime('%d/%m/%Y %I:%M %p')}"), border=1, ln=True, fill=True)
    pdf.ln(5)
    pdf.seccion_titulo("Consolidado Diario de Tiempos Reales")
    
    if not df_resumen.empty:
        pdf.set_fill_color(225, 225, 225)
        pdf.set_text_color(50, 50, 50)
        pdf.set_font("Helvetica", "B", 7)
        anchos = [85, 30, 30, 45]
        for i, col in enumerate(df_resumen.columns): 
            pdf.cell(anchos[i], 6, safestr(str(col).upper()), border=1, align="C", fill=True)
        pdf.ln()
        pdf.set_font("Helvetica", "", 7)
        for _, fila in df_resumen.iterrows():
            for i, item in enumerate(fila):
                pdf.set_fill_color(255, 255, 255)
                pdf.set_text_color(0, 0, 0)
                if "Sin Salida" in str(item) or "Sin Ingreso" in str(item): 
                    pdf.set_fill_color(253, 230, 230)
                    pdf.set_text_color(180, 0, 0)
                pdf.cell(anchos[i], 5, safestr(str(item)[:45]), border=1, align="C" if i > 0 else "L", fill=True)
            pdf.ln()
    return finalizar_pdf(pdf)

def generar_pdf_semanal_tiempos(df_diario, df_semanal, f_inicio, f_fin):
    pdf = ReporteGenerencialPDF(orientation='L', unit='mm', format='A4')
    pdf.alias_nb_pages()
    pdf.add_page()
    
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(84, 98, 143)
    
    inicio_str = f_inicio.strftime('%d/%m/%Y') if hasattr(f_inicio, 'strftime') else str(f_inicio)
    fin_str = f_fin.strftime('%d/%m/%Y') if hasattr(f_fin, 'strftime') else str(f_fin)
    
    titulo = f" Auditoria Semanal Consolidada ({inicio_str} al {fin_str})"
    pdf.cell(0, 10, safestr(titulo), border=1, ln=True, fill=True, align="C")
    pdf.ln(5)
    
    if df_diario is not None and not df_diario.empty and df_semanal is not None and not df_semanal.empty:
        df_full = pd.merge(df_diario, df_semanal, on='Vehículo / Placa', how='left')
        w = [75, 25, 25, 25, 30, 25, 35, 35] 
        
        pdf.set_fill_color(210, 210, 215)
        pdf.set_text_color(50, 50, 50)
        pdf.set_font("Helvetica", "B", 8)
        
        headers = ['VEHICULO / PLACA', 'FECHA', '1RA SALIDA', 'ULT ENTRADA', 'TIEMPO DIARIO', 'DIAS TRAB.', 'TIEMPO SEMANAL', 'PROMEDIO DIARIO']
        for i, h in enumerate(headers):
            pdf.cell(w[i], 8, safestr(h), border=1, align="C", fill=True)
        pdf.ln()
        
        pdf.set_font("Helvetica", "", 8)
        last_tec = None
        
        for idx, row in df_full.iterrows():
            tec = row['Vehículo / Placa']
            fecha_str = row['Fecha'].strftime('%d/%m/%Y') if hasattr(row['Fecha'], 'strftime') else str(row['Fecha'])
            
            if tec != last_tec:
                tec_display = safestr(tec)[:40]
                dias = str(row['Días Trabajados'])
                t_sem = safestr(row['Tiempo Total Semana'])
                p_dia = safestr(row['Promedio Diario'])
                pdf.set_fill_color(240, 248, 255) 
                fill = True
                last_tec = tec
            else:
                tec_display = "" 
                dias = ""
                t_sem = ""
                p_dia = ""
                pdf.set_fill_color(255, 255, 255)
                fill = False
                
            pdf.set_text_color(0, 0, 0)
            
            if tec_display != "": pdf.set_font("Helvetica", "B", 8)
            pdf.cell(w[0], 6, tec_display, border=1, align="L", fill=fill)
            pdf.set_font("Helvetica", "", 8)
            
            pdf.cell(w[1], 6, fecha_str, border=1, align="C", fill=fill)
            pdf.cell(w[2], 6, safestr(row['Primera Salida']), border=1, align="C", fill=fill)
            pdf.cell(w[3], 6, safestr(row['Última Entrada']), border=1, align="C", fill=fill)
            
            if row['Tiempo Diario'] == "00:00:00": pdf.set_text_color(180, 180, 180)
            pdf.cell(w[4], 6, safestr(row['Tiempo Diario']), border=1, align="C", fill=fill)
            pdf.set_text_color(0, 0, 0)
            
            if tec_display != "": pdf.set_font("Helvetica", "B", 8)
            pdf.cell(w[5], 6, dias, border=1, align="C", fill=fill)
            pdf.cell(w[6], 6, t_sem, border=1, align="C", fill=fill)
            
            if tec_display != "": pdf.set_text_color(0, 100, 0) 
            pdf.cell(w[7], 6, p_dia, border=1, align="C", fill=fill)
            
            pdf.set_font("Helvetica", "", 8)
            pdf.set_text_color(0, 0, 0)
            pdf.ln()
            
    else:
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 10, "Sin datos disponibles.", border=0, ln=True)
        
    return finalizar_pdf(pdf)

def generar_pdf_telemetria_matriz(df_matriz, limite_vel):
    pdf = ReporteGenerencialPDF(orientation='L', unit='mm', format='A4') 
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(84, 98, 143)
    pdf.set_fill_color(252, 252, 252)
    pdf.cell(0, 10, safestr(f" Matriz de Infracciones y Velocidad Promedio (> {limite_vel} km/h) - {get_hn_time().strftime('%d/%m/%Y %I:%M %p')}"), border=1, ln=True, fill=True)
    pdf.ln(5)
    
    if not df_matriz.empty:
        pdf.seccion_titulo("Vehiculos con Excesos Confirmados")
        
        has_prom = 'Promedio Vel. (km/h)' in df_matriz.columns
        col_total = next((c for c in df_matriz.columns if 'TOTAL' in str(c).upper()), None)
        
        w_placa = 95  
        w_opcion = 20 
        w_prom = 25 if has_prom else 0  
        w_total = 12 if col_total else 0
        
        espacio_restante = 275 - w_placa - w_opcion - w_prom - w_total
        cols_dias = len(df_matriz.columns) - 2 - (1 if has_prom else 0) - (1 if col_total else 0)
        w_dia = espacio_restante / cols_dias if cols_dias > 0 else 10
        
        font_size = 5.5 if cols_dias <= 15 else 4.5 
        pdf.set_font("Helvetica", "B", font_size)
        pdf.set_fill_color(225, 225, 225)
        pdf.set_text_color(50, 50, 50)
        
        for i, col in enumerate(df_matriz.columns):
            if i == 0: w = w_placa
            elif i == 1: w = w_opcion
            elif col == 'Promedio Vel. (km/h)': w = w_prom
            elif str(col).upper() == 'TOTAL': w = w_total
            else: w = w_dia
            pdf.cell(w, 6, safestr(str(col)[:20]), border=1, align="C", fill=True)
        pdf.ln()
        
        pdf.set_font("Helvetica", "", font_size)
        for _, fila in df_matriz.iterrows():
            for i, (col_name, item) in enumerate(fila.items()):
                if i == 0: w = w_placa
                elif i == 1: w = w_opcion
                elif col_name == 'Promedio Vel. (km/h)': w = w_prom
                elif str(col_name).upper() == 'TOTAL': w = w_total
                else: w = w_dia
                
                valstr = str(item).replace('.0', '').strip()
                pdf.set_fill_color(255, 255, 255)
                pdf.set_text_color(0, 0, 0)
                
                if col_name == 'Promedio Vel. (km/h)':
                    if valstr != "-" and valstr != "":
                        pdf.set_fill_color(230, 240, 255)
                        pdf.set_text_color(0, 50, 150)
                        valstr = f"{valstr} km/h"
                    else:
                        valstr = "-"
                elif i > 1 and str(col_name).upper() != 'TOTAL': 
                    try:
                        num = float(valstr)
                        if num > 0:
                            pdf.set_fill_color(253, 230, 230)
                            pdf.set_text_color(180, 0, 0)
                            valstr = str(int(num))
                        else: valstr = "-" 
                    except:
                        if valstr == '0': valstr = "-"
                
                max_chars = 80 if i == 0 else (20 if i == 1 else 15)
                pdf.cell(w, 5, safestr(valstr[:max_chars]), border=1, align="C" if i > 0 else "L", fill=True)
            pdf.ln()
    else:
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(0, 100, 0)
        pdf.cell(0, 6, f"Operacion Segura: Nadie supero los {limite_vel} km/h.", ln=True)
        
    return finalizar_pdf(pdf)

# ==============================================================================
# REGISTRO MANUAL DE HORA DE ALMUERZO POR TÉCNICO (variable día a día)
# Se guarda en un caché local (JSON), igual patrón que cache_fttx.tmp: no
# depende de Google Sheets, vive mientras el contenedor/app siga corriendo.
# ==============================================================================
_CACHE_ALMUERZOS_PATH = "cache_almuerzos.json"
_CACHE_ORDENES_MANUALES_PATH = "cache_ordenes_manuales.json"

def guardar_almuerzo(conn, tecnico: str, fecha: str, hora_inicio: str, hora_fin: str, registrado_por: str = "") -> bool:
    """
    Guarda o actualiza el horario de almuerzo de un técnico para una fecha
    específica en un caché local (cache_almuerzos.json). Si ya existía un
    registro para ese mismo técnico+fecha, lo reemplaza (no se duplica).
    El parámetro 'conn' se mantiene por compatibilidad pero no se usa.
    """
    try:
        tec_norm = str(tecnico).strip().upper()
        fecha_str = str(fecha)

        registros = []
        if os.path.exists(_CACHE_ALMUERZOS_PATH):
            try:
                with open(_CACHE_ALMUERZOS_PATH, "r", encoding="utf-8") as f:
                    registros = json.load(f)
            except Exception:
                registros = []

        registros = [r for r in registros if not (str(r.get("TECNICO", "")).upper() == tec_norm and str(r.get("FECHA", "")) == fecha_str)]
        registros.append({
            "TECNICO": tec_norm,
            "FECHA": fecha_str,
            "HORA_INICIO": str(hora_inicio),
            "HORA_FIN": str(hora_fin),
            "REGISTRADO_POR": str(registrado_por),
        })

        with open(_CACHE_ALMUERZOS_PATH, "w", encoding="utf-8") as f:
            json.dump(registros, f, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Error en guardar_almuerzo: {e}")
        return False


def cargar_almuerzos(conn, fecha: str = None) -> pd.DataFrame:
    """
    Carga los registros de almuerzo desde el caché local (cache_almuerzos.json).
    Si se especifica 'fecha' (formato 'YYYY-MM-DD'), filtra solo esa fecha.
    Devuelve un DataFrame vacío si el caché no existe todavía o está vacío.
    El parámetro 'conn' se mantiene por compatibilidad pero no se usa.
    """
    try:
        if not os.path.exists(_CACHE_ALMUERZOS_PATH):
            return pd.DataFrame()
        with open(_CACHE_ALMUERZOS_PATH, "r", encoding="utf-8") as f:
            registros = json.load(f)
        if not registros:
            return pd.DataFrame()
        df = pd.DataFrame(registros)
        if fecha:
            df = df[df['FECHA'].astype(str) == str(fecha)].copy()
        return df
    except Exception as e:
        print(f"Error en cargar_almuerzos: {e}")
        return pd.DataFrame()


# ==============================================================================
# REGISTRO MANUAL DE ÓRDENES (para cuando la API de Cepheus falla y una orden
# real no se refleja en el sistema). Mismo patrón de caché local que el
# almuerzo: no depende de Google Sheets, vive mientras el contenedor/app
# siga corriendo.
# ==============================================================================
def guardar_orden_manual(num_orden: str, actividad: str, tecnico: str, fecha: str, hora_inicio: str, hora_liq: str = "", registrado_por: str = "") -> bool:
    """
    Guarda (o actualiza, si ya existe ese NUM) una orden ingresada manualmente.
    Se guarda en un caché local (cache_ordenes_manuales.json), igual que
    guardar_almuerzo, para no depender de la conexión a Google Sheets en el
    momento del ingreso.
    """
    try:
        num_norm = str(num_orden).strip()
        tec_norm = str(tecnico).strip().upper()

        registros = []
        if os.path.exists(_CACHE_ORDENES_MANUALES_PATH):
            try:
                with open(_CACHE_ORDENES_MANUALES_PATH, "r", encoding="utf-8") as f:
                    registros = json.load(f)
            except Exception:
                registros = []

        # Si ya existe una orden manual con el mismo NUM, se reemplaza (no se duplica).
        registros = [r for r in registros if str(r.get("NUM", "")).strip() != num_norm]
        registros.append({
            "NUM": num_norm,
            "ACTIVIDAD": str(actividad).strip().upper(),
            "TECNICO": tec_norm,
            "FECHA": str(fecha),
            "HORA_INICIO": str(hora_inicio),
            "HORA_LIQ": str(hora_liq) if hora_liq else "",
            "REGISTRADO_POR": str(registrado_por),
        })

        with open(_CACHE_ORDENES_MANUALES_PATH, "w", encoding="utf-8") as f:
            json.dump(registros, f, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Error en guardar_orden_manual: {e}")
        return False


def cargar_ordenes_manuales(fecha: str = None) -> pd.DataFrame:
    """
    Carga las órdenes manuales desde el caché local y las devuelve ya
    formateadas con las mismas columnas que usa el resto del sistema
    (NUM, ACTIVIDAD, TECNICO, HORA_INI, HORA_LIQ, ESTADO, FECHA_APE, etc.)
    para poder concatenarlas directamente a df_base y que aparezcan en el
    Panel Operativo y en el Gantt como cualquier otra orden.
    Si se especifica 'fecha' (formato 'YYYY-MM-DD'), filtra solo esa fecha.
    """
    try:
        if not os.path.exists(_CACHE_ORDENES_MANUALES_PATH):
            return pd.DataFrame()
        with open(_CACHE_ORDENES_MANUALES_PATH, "r", encoding="utf-8") as f:
            registros = json.load(f)
        if not registros:
            return pd.DataFrame()

        df = pd.DataFrame(registros)
        if fecha:
            df = df[df['FECHA'].astype(str) == str(fecha)].copy()
        if df.empty:
            return df

        df['HORA_INI'] = pd.to_datetime(df['FECHA'] + ' ' + df['HORA_INICIO'], errors='coerce')
        df['HORA_LIQ'] = pd.to_datetime(
            df.apply(lambda r: (r['FECHA'] + ' ' + r['HORA_LIQ']) if str(r.get('HORA_LIQ', '')).strip() else '', axis=1),
            errors='coerce'
        )
        df['FECHA_APE'] = df['HORA_INI']
        df['ESTADO'] = df['HORA_LIQ'].apply(lambda x: 'CERRADA' if pd.notnull(x) else 'PENDIENTE')
        df['DIAS_RETRASO'] = 0
        df['ORIGEN'] = 'MANUAL'

        diff_manual = df['HORA_LIQ'] - df['HORA_INI']
        df['TIEMPO_REAL'] = np.where(
            df['HORA_INI'].isnull() | df['HORA_LIQ'].isnull(),
            "En curso (Abierta)",
            (diff_manual.dt.total_seconds() // 3600).fillna(0).astype(int).astype(str) + "h " +
            ((diff_manual.dt.total_seconds() % 3600) // 60).fillna(0).astype(int).astype(str) + "m"
        )

        for col_extra in ['MOTIVO', 'CLIENTE', 'NOMBRE', 'COMENTARIO', 'EMPRESA', 'COLONIA', 'TIPO_CLIENTE', 'GPS']:
            if col_extra not in df.columns:
                df[col_extra] = ''
        df.loc[df['COMENTARIO'] == '', 'COMENTARIO'] = 'Orden ingresada manualmente por ' + df['REGISTRADO_POR'].astype(str)

        return df[['NUM', 'ACTIVIDAD', 'TECNICO', 'HORA_INI', 'HORA_LIQ', 'FECHA_APE', 'ESTADO', 'DIAS_RETRASO', 'TIEMPO_REAL',
                   'MOTIVO', 'CLIENTE', 'NOMBRE', 'COMENTARIO', 'EMPRESA', 'COLONIA', 'TIPO_CLIENTE', 'GPS', 'ORIGEN']].copy()
    except Exception as e:
        print(f"Error en cargar_ordenes_manuales: {e}")
        return pd.DataFrame()


def cargar_catalogo_tecnicos():
    """Lee y clasifica el archivo personal_tecnico.txt según reglas de MaxCom."""
    path = "personal_tecnico.txt"
    datos = []
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            for linea in f:
                linea = linea.strip()
                if not linea: continue
                
                # Separar por coma
                partes = linea.split(',')
                
                # El nombre está en la primera parte, reemplazamos tabs por espacios y limpiamos dobles espacios
                nombre_bruto = partes[0].replace('\t', ' ')
                nombre = ' '.join(nombre_bruto.split()).upper()
                
                if len(partes) > 1:
                    cargo_area = partes[1].strip().upper()
                else:
                    cargo_area = "N/D"
                    
                estatus = "ACTIVO"
                if len(partes) > 2 and "VACACIONES" in partes[2].upper():
                    estatus = "VACACIONES"
                
                # Regla del usuario: Clasificación de Técnico Principal
                # Ahora acepta "TECNICO PRINCIPAL", "TECNICO" y los formatos estándar de áreas
                if cargo_area in ['PLEX', 'HFC', 'FTTH', 'TECNICO PRINCIPAL', 'TECNICO', 'TECNICO_PRINCIPAL']:
                    clasificacion = "TÉCNICO PRINCIPAL"
                elif 'AYUDANTE' in cargo_area:
                    clasificacion = "AYUDANTE"
                elif 'SUPERVISOR' in cargo_area:
                    clasificacion = "SUPERVISOR"
                else:
                    clasificacion = "OTRAS ÁREAS / SOPORTE"
                    
                datos.append({
                    'Nombre': nombre,
                    'Cargo/Área': cargo_area,
                    'Clasificación': clasificacion,
                    'Estatus': estatus
                })
    return pd.DataFrame(datos)

def procesar_asistencia_vs_catalogo(df_biometrico, df_catalogo):
    """Cruza marcaciones biométricas contra el catálogo maestro de personal agrupando por áreas."""
    import unicodedata

    def limpiar_para_cruce(texto):
        """Quita tildes, dobles espacios y pasa a mayúsculas para un cruce perfecto."""
        if pd.isnull(texto): return ""
        t = str(texto).upper()
        t = ''.join(c for c in unicodedata.normalize('NFD', t) if unicodedata.category(c) != 'Mn')
        return ' '.join(t.split())

    if df_catalogo.empty:
        return pd.DataFrame()
        
    df_cat = df_catalogo.copy()
    
    # Clasificador automático de Áreas (Plex, Residencial, Otras)
    def categorizar_grupo(area):
        a = str(area).upper()
        if 'PLEX' in a: return 'PLEX'
        elif 'FTTH' in a or 'HFC' in a: return 'RESIDENCIAL'
        else: return 'OTRAS ÁREAS'
        
    df_cat['Grupo_Tabla'] = df_cat['Cargo/Área'].apply(categorizar_grupo)
    
    # Si no hay biométrico cargado aún
    if df_biometrico.empty:
        df_cat['Asistencia'] = df_cat['Estatus'].apply(lambda x: '🌴 VACACIONES' if x == 'VACACIONES' else '❌ NO MARCÓ')
        df_cat['Entrada'] = "---"
        return df_cat[['Nombre', 'Clasificación', 'Cargo/Área', 'Asistencia', 'Entrada', 'Grupo_Tabla']]

    df_bio = df_biometrico.copy()
    
    df_cat['KEY_CRUCE'] = df_cat['Nombre'].apply(limpiar_para_cruce)
    df_bio['KEY_CRUCE'] = df_bio['Empleado'].apply(limpiar_para_cruce)
    
    resultado = pd.merge(df_cat, df_bio, on='KEY_CRUCE', how='left')
    
    def determinar_asistencia(row):
        if row['Estatus'] == 'VACACIONES': return '🌴 VACACIONES'
        if pd.notnull(row.get('Entrada')) and row['Entrada'] != "---" and str(row['Entrada']).strip() != "": 
            return '✅ MARCÓ'
        return '❌ NO MARCÓ'
        
    resultado['Asistencia'] = resultado.apply(determinar_asistencia, axis=1)
    resultado['Entrada'] = resultado['Entrada'].fillna("---")
    
    # Retornamos solo Entrada, Clasificación y el Grupo_Tabla invisible
    return resultado[['Nombre', 'Clasificación', 'Cargo/Área', 'Asistencia', 'Entrada', 'Grupo_Tabla']]

def generar_pdf_ordenes_totales(df_base, fecha_corte):
    """Genera un PDF con el listado de todas las órdenes PENDIENTES, ordenadas por retraso y con columna Colonia."""
    
    # --- 1. ORDENAMIENTO ---
    df_base['DIAS_RETRASO'] = pd.to_numeric(df_base['DIAS_RETRASO'], errors='coerce').fillna(0)
    df_base = df_base.sort_values(by='DIAS_RETRASO', ascending=False)
    
    pdf = ReporteGenerencialPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(40, 50, 100)
    pdf.cell(0, 10, safestr("REPORTE DE ORDENES TOTALES PENDIENTES"), border=0, ln=True, align="C")
    
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 6, safestr(f"Corte Operativo del Día: {fecha_corte.strftime('%d/%m/%Y')}"), ln=True, align="C")
    pdf.ln(5)
    
    pdf.seccion_titulo(f"Listado Total ({len(df_base)} Órdenes en Ruta / Sin Asignar)")
    
    pdf.set_fill_color(240, 240, 240)
    pdf.set_text_color(50, 50, 50)
    pdf.set_font("Helvetica", "B", 7)
    
    # Anchos ajustados: Días(12), Orden(18), Cliente(20), Tecnico(40), Actividad(50), Colonia(50) = 190mm
    w = [12, 18, 20, 40, 50, 50] 
    
    # Encabezado
    pdf.cell(w[0], 6, "Días", border=1, align="C", fill=True)
    pdf.cell(w[1], 6, "Orden", border=1, align="C", fill=True)
    pdf.cell(w[2], 6, "Cliente", border=1, align="C", fill=True)
    pdf.cell(w[3], 6, "Tecnico", border=1, align="C", fill=True)
    pdf.cell(w[4], 6, "Actividad", border=1, align="C", fill=True)
    pdf.cell(w[5], 6, "Colonia", border=1, align="C", fill=True)
    pdf.ln()
    
    pdf.set_font("Helvetica", "", 6)
    
    for _, row in df_base.iterrows():
        if pdf.get_y() > 270:
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 7)
            pdf.set_fill_color(240, 240, 240)
            pdf.set_text_color(50, 50, 50)
            pdf.cell(w[0], 6, "Días", border=1, align="C", fill=True)
            pdf.cell(w[1], 6, "Orden", border=1, align="C", fill=True)
            pdf.cell(w[2], 6, "Cliente", border=1, align="C", fill=True)
            pdf.cell(w[3], 6, "Tecnico", border=1, align="C", fill=True)
            pdf.cell(w[4], 6, "Actividad", border=1, align="C", fill=True)
            pdf.cell(w[5], 6, "Colonia", border=1, align="C", fill=True)
            pdf.ln()
            pdf.set_font("Helvetica", "", 6)

        # 1. Semáforo de Días
        dias_val = int(row.get('DIAS_RETRASO', 0))
        if dias_val >= 7:
            pdf.set_fill_color(211, 47, 47)   # Rojo
            pdf.set_text_color(255, 255, 255)
        elif 4 <= dias_val <= 6:
            pdf.set_fill_color(245, 124, 0)   # Naranja
            pdf.set_text_color(255, 255, 255)
        elif 1 <= dias_val <= 3:
            pdf.set_fill_color(251, 192, 45)  # Amarillo
            pdf.set_text_color(0, 0, 0)
        else:
            pdf.set_fill_color(56, 142, 60)   # Verde
            pdf.set_text_color(255, 255, 255)

        pdf.set_font("Helvetica", "B", 6)
        pdf.cell(w[0], 5, str(dias_val), border=1, align="C", fill=True)

        # 2. Reset para el resto de la fila
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 6)
        
        # Datos
        num = safestr(str(row.get('NUM', 'N/D')))
        cliente = safestr(str(row.get('CLIENTE', 'N/D')))
        
        tec_raw = str(row.get('TECNICO', ''))
        tec = "SIN ASIGNAR" if pd.isna(tec_raw) or tec_raw.strip().upper() in ['NONE', 'NAN', 'N/D', 'NULL', ''] else safestr(tec_raw)[:25]
        
        act = safestr(str(row.get('ACTIVIDAD', 'N/D')))[:32]
        colonia = safestr(str(row.get('COLONIA', 'N/D')))[:35] # Obtenemos Colonia
        
        pdf.cell(w[1], 5, num, border=1, align="C")
        pdf.cell(w[2], 5, cliente, border=1, align="C")
        pdf.cell(w[3], 5, tec, border=1, align="L")
        pdf.cell(w[4], 5, act, border=1, align="L")
        pdf.cell(w[5], 5, colonia, border=1, align="L") # Imprimimos Colonia
        pdf.ln()
        
    return finalizar_pdf(pdf)


# ==============================================================================
# 7. MÓDULO DE PERSISTENCIA EN GOOGLE CLOUD STORAGE (NUEVO)
# ==============================================================================

def obtener_cliente_gcs_nativo():
    """Inicializa el cliente de GCS reciclando las credenciales de gsheets de Streamlit."""
    import streamlit as st
    creds_dict = st.secrets["connections"]["gsheets"]
    creds = service_account.Credentials.from_service_account_info(creds_dict)
    return storage.Client(credentials=creds, project=creds.project_id)

def sobrescribir_archivo_gcs(dataframe_o_bytes, nombre_bucket, nombre_archivo_destino):
    """
    Sube un DataFrame (como CSV) o un archivo binario directo a GCS.
    Si el archivo ya existe, GCS lo borra/sobrescribe automáticamente en el acto.
    """
    try:
        cliente = obtener_cliente_gcs_nativo()
        bucket = cliente.bucket(nombre_bucket)
        blob = bucket.blob(nombre_archivo_destino)
        
        # Detectar si es un DataFrame de Pandas (para el Historial Maestro)
        if isinstance(dataframe_o_bytes, pd.DataFrame):
            csv_en_ram = dataframe_o_bytes.to_csv(index=False)
            blob.upload_from_string(csv_en_ram, content_type="text/csv")
        # Detectar si son bytes puros (para guardar el archivo FTTX tal cual)
        elif isinstance(dataframe_o_bytes, bytes):
            blob.upload_from_string(dataframe_o_bytes, content_type="application/octet-stream")
        else:
            # Si es un objeto de archivo subido de Streamlit (UploadedFile)
            dataframe_o_bytes.seek(0)
            blob.upload_from_file(dataframe_o_bytes)
            
        return True
    except Exception as e:
        print(f"Error de persistencia en GCS: {e}")
        return False
        

def leer_espejo_gcs(nombre_bucket, nombre_archivo_destino):
    """
    Descarga el archivo desde GCS en memoria RAM y lo devuelve como un DataFrame de Pandas.
    """
    import io
    try:
        cliente = obtener_cliente_gcs_nativo()
        bucket = cliente.bucket(nombre_bucket)
        blob = bucket.blob(nombre_archivo_destino)
        
        if blob.exists():
            contenido = blob.download_as_bytes()
            return pd.read_csv(io.BytesIO(contenido))
        else:
            print(f"El archivo {nombre_archivo_destino} no existe en GCS.")
            return None
    except Exception as e:
        print(f"Error al leer desde GCS: {e}")
        return None


# ==============================================================================
# GESTIÓN DE ARCHIVOS BINARIOS EN GCS (repositorio documental de expedientes)
# ==============================================================================
# leer_espejo_gcs() solo sabe leer CSV (hace pd.read_csv), así que sirve para el
# índice de documentos pero NO para recuperar un PDF, Word o Excel. Estas cuatro
# funciones cubren ese hueco: subir, descargar como bytes, eliminar y listar.

def generar_pdf_rendimiento_integral_360(df_m, df_tipo_ord, df_exp_det):
    pdf = ReporteIntegral360PDF(orientation='L', unit='mm', format='A4')
    pdf.add_page()
    
    # --- MÓDULO 1: RESUMEN GERENCIAL (KPIs) ---
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(16, 185, 129) # Verde
    pdf.cell(0, 8, safestr("1. RESUMEN GERENCIAL (Indicadores Clave)"), ln=True)
    
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(50, 50, 50)
    total_tecs = len(df_m)
    total_ords = int(df_m['ÓRDENES CANTIDAD'].sum()) if 'ÓRDENES CANTIDAD' in df_m else 0
    prom_gral = round(df_m['TIEMPO PROM. EN ORDEN (Min)'].mean(), 1) if not df_m.empty else 0
    tot_inc = int(df_m['DÍAS FALTADOS'].sum() + df_m['LLAMADOS ATENCIÓN'].sum() + df_m['DÍAS NO PRESENTADO'].sum())
    
    kpi_text = (f"Técnicos Analizados: {total_tecs}   |   "
                f"Total Órdenes Ejecutadas: {total_ords}   |   "
                f"Tiempo Promedio Global: {prom_gral} min   |   "
                f"Total Incidencias y Faltas: {tot_inc}")
    pdf.cell(0, 6, safestr(kpi_text), ln=True)
    pdf.ln(5)
    
    # --- MÓDULO 2: PRODUCTIVIDAD EN CAMPO (TABLA DESGLOSADA EN PLEX Y RESIDENCIAL) ---
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(59, 130, 246) # Azul
    pdf.cell(0, 8, safestr("2. EJECUCIÓN OPERATIVA (Productividad y Tiempos de Cierre)"), ln=True)
    
    pdf.set_fill_color(240, 245, 250)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "B", 8)
    
    # Se ajustaron los anchos y las columnas para dar espacio al promedio segmentado
    w1 = [55, 18, 18, 18, 28, 28, 28, 28] 
    h1 = ["TÉCNICO", "TOTAL", "PLEX", "RESID.", "PROM PLEX(Min)", "PROM RESID(Min)", "1ra ORDEN", "ÚLT. ORDEN"]
    
    for i, h in enumerate(h1):
        pdf.cell(w1[i], 7, safestr(h), border=1, fill=True, align="C")
    pdf.ln()
    
    pdf.set_font("Helvetica", "", 8)
    for _, row in df_m.iterrows():
        pdf.cell(w1[0], 6, safestr(row.get('TÉCNICO', ''))[:32], border=1)
        pdf.cell(w1[1], 6, safestr(row.get('ÓRDENES CANTIDAD', 0)), border=1, align="C")
        pdf.cell(w1[2], 6, safestr(row.get('ÓRDENES PLEX', 0)), border=1, align="C")
        pdf.cell(w1[3], 6, safestr(row.get('ÓRDENES RESIDENCIAL', 0)), border=1, align="C")
        
        # Tiempo Promedio PLEX
        t_plex = row.get('TIEMPO PROM. PLEX (Min)', 0.0)
        t_plex_str = f"{t_plex:.1f}" if t_plex > 0 else "---"
        pdf.set_text_color(220, 38, 38) if t_plex > 120 else pdf.set_text_color(0, 0, 0)
        pdf.cell(w1[4], 6, safestr(t_plex_str), border=1, align="C")
        pdf.set_text_color(0, 0, 0)
        
        # Tiempo Promedio RESIDENCIAL
        t_resi = row.get('TIEMPO PROM. RESIDENCIAL (Min)', 0.0)
        t_resi_str = f"{t_resi:.1f}" if t_resi > 0 else "---"
        pdf.set_text_color(220, 38, 38) if t_resi > 90 else pdf.set_text_color(0, 0, 0)
        pdf.cell(w1[5], 6, safestr(t_resi_str), border=1, align="C")
        pdf.set_text_color(0, 0, 0)
        
        pdf.cell(w1[6], 6, safestr(row.get('HORA 1ra ORDEN', '--')), border=1, align="C")
        pdf.cell(w1[7], 6, safestr(row.get('HORA ÚLT. ORDEN', '--')), border=1, align="C")
        pdf.ln()

    # --- MÓDULO 3: LOGÍSTICA Y RRHH (GPS y Disciplina) ---
    pdf.ln(8)
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(245, 158, 11) # Naranja/Ambar
    pdf.cell(0, 8, safestr("3. LOGÍSTICA Y RRHH (Control de Flotilla y Asistencia)"), ln=True)
    
    pdf.set_fill_color(254, 243, 199)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "B", 8)
    
    w2 = [60, 35, 35, 35, 40]
    h2 = ["TÉCNICO", "SALIDA GPS", "ENTRADA GPS", "DÍAS FALTADOS", "LLAMADOS ATENCIÓN"]
    
    for i, h in enumerate(h2):
        pdf.cell(w2[i], 7, safestr(h), border=1, fill=True, align="C")
    pdf.ln()
    
    pdf.set_font("Helvetica", "", 8)
    for _, row in df_m.iterrows():
        faltas_tot = row.get('DÍAS FALTADOS', 0) + row.get('DÍAS NO PRESENTADO', 0)
        llamados = row.get('LLAMADOS ATENCIÓN', 0)
        gps_s = row.get('SALIDA PLANTEL (GPS)', '--')
        gps_e = row.get('ENTRADA PLANTEL (GPS)', '--')
        
        pdf.cell(w2[0], 6, safestr(row.get('TÉCNICO', ''))[:35], border=1)
        pdf.cell(w2[1], 6, safestr(gps_s), border=1, align="C")
        pdf.cell(w2[2], 6, safestr(gps_e), border=1, align="C")
        
        pdf.set_text_color(220, 38, 38) if faltas_tot > 0 else pdf.set_text_color(0, 0, 0)
        pdf.cell(w2[3], 6, safestr(faltas_tot), border=1, align="C")
        
        pdf.set_text_color(217, 119, 6) if llamados > 0 else pdf.set_text_color(0, 0, 0)
        pdf.cell(w2[4], 6, safestr(llamados), border=1, align="C")
        pdf.set_text_color(0, 0, 0)
        pdf.ln()

    # --- MÓDULO 4: DESGLOSE POR ACTIVIDAD ---
    if df_tipo_ord is not None and not df_tipo_ord.empty:
        pdf.add_page() # Nueva página
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(139, 92, 246) # Morado
        pdf.cell(0, 8, safestr("4. MATRIZ DE RENDIMIENTO POR TIPO DE ACTIVIDAD"), ln=True)
        pdf.set_font("Helvetica", "I", 9)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 5, safestr("Detalle de volumen, tiempo de ejecucion y porcentaje de eficiencia por categoria tecnica."), ln=True)
        pdf.ln(3)

        pdf.set_fill_color(243, 232, 255)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "B", 8)
        
        w3 = [80, 75, 25, 25, 30]
        h3 = ["TÉCNICO", "TIPO DE ACTIVIDAD", "CANTIDAD", "PROM(Min)", "REND(%)"]
        
        for i, h in enumerate(h3):
            pdf.cell(w3[i], 7, safestr(h), border=1, fill=True, align="C")
        pdf.ln()
        
        pdf.set_font("Helvetica", "", 8)
        df_listado = df_tipo_ord.sort_values(['TECNICO', 'Ordenes'], ascending=[True, False])
        
        activ_evaluables = [
            'INSEQUIPO', 'INSFIBRA', 'INSFIBRACORP', 'INSFIBRACOPR', 
            'SOP', 'SOPFIBRA', 'SOPFIBRACORP', 
            'TRASLADOEXTFIBRA', 'TRASLADOINTERNOFIBRA', 'TVADICIONAL'
        ]
        
        tec_actual = ""
        for _, row_t in df_listado.iterrows():
            tec_print = safestr(row_t['TECNICO'])[:45] if row_t['TECNICO'] != tec_actual else ""
            tec_actual = row_t['TECNICO']
            
            pdf.cell(w3[0], 6, tec_print, border='L' if tec_print == "" else 1)
            pdf.cell(w3[1], 6, safestr(row_t['TipoOrden'])[:40], border=1)
            pdf.cell(w3[2], 6, safestr(row_t['Ordenes']), border=1, align="C")
            
            t_min = round(row_t['MinProm'], 1)
            pdf.set_text_color(220, 38, 38) if t_min > 90 else (pdf.set_text_color(16, 185, 129) if t_min < 45 else pdf.set_text_color(0, 0, 0))
            pdf.cell(w3[3], 6, safestr(t_min), border=1, align="C")
            pdf.set_text_color(0, 0, 0)
            
            # Evaluación de Rendimiento %
            tipo_act = str(row_t.get('TipoOrden', '')).upper().strip()
            rend_val = row_t.get('Rendimiento_Prom') if 'Rendimiento_Prom' in df_tipo_ord.columns else None
            
            if tipo_act in activ_evaluables and pd.notna(rend_val):
                rend_str = f"{int(round(rend_val))}%"
                if rend_val >= 100:
                    pdf.set_text_color(16, 185, 129) # Verde
                elif rend_val >= 80:
                    pdf.set_text_color(217, 119, 6)   # Naranja
                else:
                    pdf.set_text_color(220, 38, 38)  # Rojo
            else:
                rend_str = "---"
                
            pdf.cell(w3[4], 6, safestr(rend_str), border=1, align="C")
            pdf.set_text_color(0, 0, 0)
            pdf.ln()

    return finalizar_pdf(pdf) # Asegúrate de que esta función exista en tools.py
    
# ==============================================================================
# BLINDAJE CORE - USO EXCLUSIVO PARA APP.PY (CEREBRO PRINCIPAL)
# ==============================================================================
def generar_pdf_gastos_vehiculo(df_gastos, vehiculo, rango_fechas, total):
    pdf = ReporteGastosPDF(orientation='P', unit='mm', format='A4')
    pdf.add_page()
    
    # Info de cabecera
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(16, 185, 129) # Verde
    pdf.cell(0, 8, safestr(f"1. RESUMEN DE GASTOS - UNIDAD: {vehiculo}"), ln=True)
    
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(50, 50, 50)
    
    if isinstance(rango_fechas, (list, tuple)) and len(rango_fechas) == 2:
        fecha_txt = f"{rango_fechas[0].strftime('%d/%m/%Y')} al {rango_fechas[1].strftime('%d/%m/%Y')}"
    elif isinstance(rango_fechas, (list, tuple)) and len(rango_fechas) == 1:
        fecha_txt = f"{rango_fechas[0].strftime('%d/%m/%Y')}"
    else:
        fecha_txt = "Todas las fechas"

    kpi_text = f"Periodo: {fecha_txt}   |   Total Facturas: {len(df_gastos)}   |   Monto Invertido: L. {total:,.2f}"
    pdf.cell(0, 6, safestr(kpi_text), ln=True)
    pdf.ln(5)
    
    # Tabla de detalle
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(59, 130, 246) # Azul
    pdf.cell(0, 8, safestr("2. HISTORIAL DETALLADO DE FACTURAS Y SERVICIOS"), ln=True)
    
    pdf.set_fill_color(240, 245, 250)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "B", 9)
    
    w = [25, 45, 90, 30] 
    h = ["FECHA", "CATEGORÍA", "DESCRIPCIÓN / FACTURA", "MONTO (L)"]
    
    for i, head in enumerate(h):
        pdf.cell(w[i], 7, safestr(head), border=1, fill=True, align="C")
    pdf.ln()
    
    pdf.set_font("Helvetica", "", 9)
    for _, row in df_gastos.iterrows():
        desc = safestr(row.get('DESCRIPCION', ''))[:55]
        cat = safestr(row.get('TIPO_GASTO', ''))[:20]
        
        pdf.cell(w[0], 6, safestr(row.get('FECHA', '')), border=1, align="C")
        pdf.cell(w[1], 6, cat, border=1, align="C")
        pdf.cell(w[2], 6, desc, border=1)
        pdf.cell(w[3], 6, safestr(f"{float(row.get('MONTO', 0)):,.2f}"), border=1, align="R")
        pdf.ln()
        
    # Totales finales en la tabla
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_fill_color(220, 230, 240)
    pdf.cell(w[0] + w[1] + w[2], 7, safestr("TOTAL GASTOS DEL PERIODO:"), border=1, align="R", fill=True)
    pdf.cell(w[3], 7, safestr(f"L. {total:,.2f}"), border=1, align="R", fill=True)
    
    fd, path = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    pdf.output(path)
    with open(path, "rb") as f: data = f.read()
    os.remove(path)
    return data

def generar_pdf_reporte_general_gastos(df_gastos):
    """
    Genera un PDF gerencial de gastos y flota, agrupado por vehículo,
    mostrando fecha cronológica y descripciones detalladas de forma robusta.
    """
    from fpdf import FPDF
    import pandas as pd
    
    # Configuramos el documento (Letter / Carta)
    pdf = FPDF(orientation='P', unit='mm', format='Letter')
    pdf.add_page()
    
    # =========================================================
    # ENCABEZADO PRINCIPAL DEL DOCUMENTO
    # =========================================================
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(15, 23, 42)
    pdf.cell(0, 10, safestr("REPORTE GERENCIAL DE GASTOS Y FLOTA"), ln=True, align="C")
    
    pdf.set_font("Helvetica", "", 10)
    fecha_gen = get_honduras_time().strftime("%Y-%m-%d %H:%M")
    pdf.cell(0, 6, safestr(f"Fecha de Generación: {fecha_gen}"), ln=True, align="C")
    pdf.ln(5)
    
    # Si la base viene vacía, retornamos un PDF con el aviso
    if df_gastos is None or df_gastos.empty:
        pdf.set_font("Helvetica", "I", 10)
        pdf.cell(0, 10, "No hay gastos registrados en el periodo seleccionado.", ln=True, align="C")
        try: return pdf.output(dest='S').encode('latin1')
        except: return bytes(pdf.output())

    # =========================================================
    # DETECCIÓN ROBUSTA DE COLUMNAS (Independiente de mayúsculas/minúsculas o espacios)
    # =========================================================
    # Creamos un mapa de columnas normalizadas (sin espacios y en minúsculas)
    df_cols = {col.strip().lower(): col for col in df_gastos.columns}
    
    # 1. Buscar columna del vehículo
    col_vehiculo = None
    for opt in ['vehiculo', 'vehículo', 'placa', 'vehicle', 'unidad']:
        if opt in df_cols:
            col_vehiculo = df_cols[opt]
            break
            
    # 2. Buscar columna de la fecha
    col_fecha = None
    for opt in ['fecha', 'date', 'fec', 'fecha_gasto']:
        if opt in df_cols:
            col_fecha = df_cols[opt]
            break

    # 3. Buscar columna de categoría / tipo de gasto
    col_tipo = None
    for opt in ['tipo_gasto', 'tipo de gasto', 'tipo_de_gasto', 'categoria', 'categoría', 'tipo', 'category']:
        if opt in df_cols:
            col_tipo = df_cols[opt]
            break

    # 4. Buscar columna de descripción
    col_desc = None
    for opt in ['descripcion', 'descripción', 'detalle', 'description', 'desc']:
        if opt in df_cols:
            col_desc = df_cols[opt]
            break

    # 5. Buscar columna del monto total
    col_total = None
    for opt in ['total', 'costo', 'monto', 'total (l.)', 'total(l.)', 'monto_gasto', 'costo_total', 'precio', 'importe', 'valor']:
        if opt in df_cols:
            col_total = df_cols[opt]
            break

    # Aseguramos la existencia de la columna vehículo para agrupar
    vehiculo_col_real = col_vehiculo if col_vehiculo else 'VEHICULO'
    if vehiculo_col_real not in df_gastos.columns:
        df_gastos[vehiculo_col_real] = "FLOTA GENERAL"
        
    vehiculos_unicos = sorted(df_gastos[vehiculo_col_real].dropna().unique().tolist())
    gran_total = 0.0
    
    # =========================================================
    # RECORRIDO Y AGRUPACIÓN POR CADA VEHÍCULO
    # =========================================================
    for vehiculo in vehiculos_unicos:
        df_vehiculo = df_gastos[df_gastos[vehiculo_col_real] == vehiculo]
        if df_vehiculo.empty:
            continue
            
        # --- TÍTULO DE LA SECCIÓN DEL VEHÍCULO ---
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_fill_color(30, 58, 138)  # Azul oscuro institucional
        pdf.set_text_color(255, 255, 255) # Letras blancas
        pdf.cell(190, 8, safestr(f" VEHÍCULO: {vehiculo}"), border=1, ln=True, fill=True, align="L")
        
        # --- ENCABEZADOS DE LA TABLA (190mm ancho total) ---
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_fill_color(200, 200, 200) # Fondo gris claro para los títulos de columnas
        pdf.set_text_color(15, 23, 42)    # Letra oscura
        
        pdf.cell(25, 6, "FECHA", border=1, fill=True, align="C")
        pdf.cell(40, 6, "TIPO DE GASTO", border=1, fill=True, align="C")
        pdf.cell(95, 6, "DESCRIPCION / DETALLE", border=1, fill=True, align="C")
        pdf.cell(30, 6, "TOTAL (L.)", border=1, fill=True, align="C")
        pdf.ln()

        subtotal_vehiculo = 0.0
        pdf.set_font("Helvetica", "", 8) # Letra estándar para el contenido
        
        # --- IMPRESIÓN DEL DETALLE DE GASTOS ---
        for _, row in df_vehiculo.iterrows():
            # 1. Fecha limpia
            fecha_val = row.get(col_fecha) if col_fecha else None
            fecha_str = safestr(fecha_val)[:10] if pd.notna(fecha_val) else 'N/D'
            
            # 2. Categoría
            cat_val = row.get(col_tipo) if col_tipo else None
            cat = safestr(cat_val)[:25] if pd.notna(cat_val) else 'N/D'
            
            # 3. Descripción detallada
            desc_val = row.get(col_desc) if col_desc else None
            desc = safestr(desc_val)[:70] if pd.notna(desc_val) else 'Sin detalle'
            
            # 4. Obtención y parsing del valor numérico
            raw_val = row.get(col_total) if col_total else 0.0
            
            if pd.isna(raw_val) or raw_val == '':
                tot = 0.0
            else:
                try:
                    # Si el valor ya es numérico, lo convertimos directamente
                    if isinstance(raw_val, (int, float)):
                        tot = float(raw_val)
                    else:
                        # Si es de tipo string, removemos símbolos de moneda y separadores
                        val_clean = str(raw_val).upper().replace('L.', '').replace('L', '').replace('$', '').replace(',', '').strip()
                        tot = float(val_clean)
                except Exception:
                    tot = 0.0
                
            subtotal_vehiculo += tot
            
            # Escribir la fila en el PDF
            pdf.cell(25, 6, fecha_str, border=1, align="C")
            pdf.cell(40, 6, cat, border=1, align="L")
            pdf.cell(95, 6, desc, border=1, align="L")
            pdf.cell(30, 6, safestr(f"{tot:,.2f}"), border=1, align="R")
            pdf.ln()
        
        # --- FILA DE SUBTOTAL DEL VEHÍCULO ---
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(230, 240, 250) # Azul muy clarito
        pdf.cell(160, 6, "SUBTOTAL VEHICULO:", border=1, align="R", fill=True)
        pdf.cell(30, 6, safestr(f"L. {subtotal_vehiculo:,.2f}"), border=1, align="R", fill=True)
        pdf.ln(8) # Espacio de separación antes de imprimir el siguiente vehículo
        
        # Acumulamos para el total general
        gran_total += subtotal_vehiculo

    # =========================================================
    # GRAN TOTAL INVERTIDO AL FINAL DEL DOCUMENTO
    # =========================================================
    pdf.ln(5)
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_fill_color(30, 58, 138)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(140, 10, "GRAN TOTAL INVERTIDO EN FLOTA:", border=1, align="R", fill=True)
    pdf.cell(50, 10, safestr(f"L. {gran_total:,.2f}"), border=1, align="C", fill=True)
    
    # =========================================================
    # RETORNO SEGURO DE BYTES
    # =========================================================
    try:
        return pdf.output(dest='S').encode('latin1')
    except Exception:
        return bytes(pdf.output())



# ==============================================================================
# MOTOR DOCX: REPORTE DE FALTAS INDIVIDUAL (DISEÑO AJUSTADO A UNA SOLA HOJA)
# ==============================================================================
def generar_docx_reporte_faltas_individual(df):
    import io
    import os
    import re
    from datetime import datetime, timedelta
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml, OxmlElement
        from docx.oxml.ns import nsdecls, qn
    except ImportError:
        return b""
        
    doc = Document()
    
    # Ajustar márgenes a 0.5 pulgadas para maximizar el espacio útil de la hoja
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    # Configuración de fuente base Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(8.5)
    
    # Helpers XML de docx para sombreados, bordes y márgenes
    def set_cell_shading(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        
    def set_cell_borders(cell, color="555555", sz="4", val="single"):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'''
            <w:tcBorders {nsdecls("w")}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            </w:tcBorders>
        ''')
        tcPr.append(tcBorders)

    def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def format_cell_p(cell, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(1.5)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.0
        return p

    def add_micro_spacing(doc, space_pt):
        # Genera un espaciador ultra-preciso sin crear párrafos vacíos gigantescos
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_pt)
        p.paragraph_format.line_spacing = Pt(1)
        r = p.add_run()
        r.font.size = Pt(1)
        
    NAVY_HEX = "1A1D54"  # Azul marino idéntico al corporativo
    BORDER_COLOR = "000000"  # Bordes negros limpios

    # 1. Tabla de Encabezado Principal (1 fila, 3 columnas)
    header_table = doc.add_table(rows=1, cols=3)
    header_table.style = 'Table Grid'
    header_table.autofit = False
    
    col_widths = [Inches(1.8), Inches(3.7), Inches(2.0)]
    cells = header_table.rows[0].cells
    for i, width in enumerate(col_widths):
        cells[i].width = width
        set_cell_borders(cells[i], color=BORDER_COLOR, sz="4")
        set_cell_margins(cells[i], top=40, bottom=40, left=100, right=100)
        
    # Celda 1: Área del logo de Maxcom
    p1 = format_cell_p(cells[0], WD_ALIGN_PARAGRAPH.CENTER)
    _logo_docx = obtener_logo_reporte()
    if _logo_docx:
        try:
            p1.add_run().add_picture(_logo_docx, height=Inches(0.42))  # por alto: cualquier proporción cabe
        except:
            r = p1.add_run("maxcom")
            r.font.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(30, 58, 138)
    else:
        r = p1.add_run("maxcom")
        r.font.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(30, 58, 138)
        
    # Celda 2: Nombre del Formulario
    p2 = format_cell_p(cells[1], WD_ALIGN_PARAGRAPH.CENTER)
    p2.paragraph_format.space_before = Pt(8)
    r2 = p2.add_run("REPORTE DE FALTAS")
    r2.font.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(80, 80, 80)
    
    # Celda 3: Datos de Control de Calidad / nested table
    nested_meta = cells[2].add_table(rows=4, cols=2)
    nested_meta.autofit = False
    nested_meta.columns[0].width = Inches(0.9)
    nested_meta.columns[1].width = Inches(1.1)
    
    metadata = [
        ("CÓDIGO", "HN-GG-RH-FR-09"),
        ("VERSIÓN", "1.0"),
        ("FECHA", "12/05/2026"),
        ("CLASIFICACIÓN", "INTERNO")
    ]
    for row_idx, (key, val) in enumerate(metadata):
        n_row = nested_meta.rows[row_idx]
        n_row.cells[0].width = Inches(0.9)
        n_row.cells[1].width = Inches(1.1)
        for nc in n_row.cells:
            set_cell_borders(nc, color=BORDER_COLOR, sz="4")
            set_cell_margins(nc, top=15, bottom=15, left=30, right=30)
            
        p_k = format_cell_p(n_row.cells[0], WD_ALIGN_PARAGRAPH.CENTER)
        r_k = p_k.add_run(key)
        r_k.font.bold = True
        r_k.font.size = Pt(7)
        
        p_v = format_cell_p(n_row.cells[1], WD_ALIGN_PARAGRAPH.CENTER)
        r_v = p_v.add_run(val)
        r_v.font.size = Pt(7)
        
    add_micro_spacing(doc, 4)
    
    # Título central "REPORTE DE FALTAS"
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("REPORTE DE FALTAS")
    r_title.font.bold = True
    r_title.font.size = Pt(12)
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(6)
    
    # 2. Tabla: Datos del Jefe del Departamento Solicitante
    table_jefe = doc.add_table(rows=6, cols=2)
    table_jefe.style = 'Table Grid'
    table_jefe.autofit = False
    table_jefe.columns[0].width = Inches(2.8)
    table_jefe.columns[1].width = Inches(4.7)
    
    hdr_cell = table_jefe.rows[0].cells[0]
    hdr_cell.merge(table_jefe.rows[0].cells[1])
    set_cell_shading(hdr_cell, NAVY_HEX)
    set_cell_margins(hdr_cell, top=45, bottom=40, left=100, right=100)
    p_hdr = format_cell_p(hdr_cell, WD_ALIGN_PARAGRAPH.CENTER)
    r_hdr = p_hdr.add_run("Datos del Jefe del Departamento Solicitante")
    r_hdr.font.bold = True
    r_hdr.font.size = Pt(9)
    r_hdr.font.color.rgb = RGBColor(255, 255, 255)
    
    jefe_labels = [
        "Nombre del Jefe del Departamento",
        "Puesto",
        "Departamento",
        "Ciudad",
        "Fecha de Ingreso"
    ]
    
    for idx, label in enumerate(jefe_labels, start=1):
        row = table_jefe.rows[idx]
        row.cells[0].width = Inches(2.8)
        row.cells[1].width = Inches(4.7)
        
        for c in row.cells:
            set_cell_borders(c, color=BORDER_COLOR, sz="4")
            set_cell_margins(c, top=30, bottom=30, left=80, right=80)
            
        p_lbl = format_cell_p(row.cells[0])
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(8.5)
        
    add_micro_spacing(doc, 4)
    
    # 3. Tabla: Datos de Empleado Reportado
    table_emp = doc.add_table(rows=7, cols=2)
    table_emp.style = 'Table Grid'
    table_emp.autofit = False
    table_emp.columns[0].width = Inches(2.8)
    table_emp.columns[1].width = Inches(4.7)
    
    hdr_cell_emp = table_emp.rows[0].cells[0]
    hdr_cell_emp.merge(table_emp.rows[0].cells[1])
    set_cell_shading(hdr_cell_emp, NAVY_HEX)
    set_cell_margins(hdr_cell_emp, top=45, bottom=40, left=100, right=100)
    p_hdr_emp = format_cell_p(hdr_cell_emp, WD_ALIGN_PARAGRAPH.CENTER)
    r_hdr_emp = p_hdr_emp.add_run("Datos de Empleado Reportado")
    r_hdr_emp.font.bold = True
    r_hdr_emp.font.size = Pt(9)
    r_hdr_emp.font.color.rgb = RGBColor(255, 255, 255)
    
    emp_name = "N/D"
    if not df.empty and 'TECNICO' in df.columns:
        emp_name = str(df['TECNICO'].iloc[0]).upper().strip()
        
    fechas_list = []
    if not df.empty and 'FECHA_INCIDENCIA' in df.columns:
        raw_dates = df['FECHA_INCIDENCIA'].dropna().astype(str).unique()
        def date_key(d_str):
            for fmt in ('%d/%m/%Y', '%Y-%m-%d', '%d-%m-%Y'):
                try: return datetime.strptime(d_str.strip(), fmt)
                except: continue
            return datetime.min
        fechas_list = sorted(list(raw_dates), key=date_key)
        
    fecha_falta_str = " - ".join(fechas_list) if fechas_list else "N/D"
    
    horas_list = []
    if not df.empty and 'FECHA_REGISTRO' in df.columns:
        for val in df['FECHA_REGISTRO'].dropna():
            match = re.search(r'(\d{2}):(\d{2})', str(val))
            if match:
                horas_list.append(f"{match.group(1)}:{match.group(2)}")
        horas_list = sorted(list(set(horas_list)))
        
    hora_falta_str = " - ".join(horas_list) if horas_list else "N/D"
    
    emp_labels = [
        ("Nombre del Empleado", emp_name),
        ("Código de Empleado", ""),
        ("Puesto", ""),
        ("Horario de Trabajo", ""),
        ("Fecha de la Falta Ocurrida", fecha_falta_str),
        ("Hora de la Falta Ocurrida", hora_falta_str)
    ]
    
    for idx, (label, val_text) in enumerate(emp_labels, start=1):
        row = table_emp.rows[idx]
        row.cells[0].width = Inches(2.8)
        row.cells[1].width = Inches(4.7)
        
        for c in row.cells:
            set_cell_borders(c, color=BORDER_COLOR, sz="4")
            set_cell_margins(c, top=30, bottom=30, left=80, right=80)
            
        p_lbl = format_cell_p(row.cells[0])
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(8.5)
        
        p_val = format_cell_p(row.cells[1])
        r_val = p_val.add_run(val_text)
        r_val.font.size = Pt(8.5)
        
    add_micro_spacing(doc, 4)
    
    # 4. Tabla: Detalle de la Falta (DISEÑO RESPONSIVO DINÁMICO PARA INTEGRACIÓN EN 1 HOJA)
    incidents = []
    if not df.empty:
        df_sorted = df.copy()
        def parse_row_date(r):
            for fmt in ('%d/%m/%Y', '%Y-%m-%d', '%d-%m-%Y'):
                try: return datetime.strptime(str(r.get('FECHA_INCIDENCIA', '')).strip(), fmt)
                except: continue
            return datetime.min
        df_sorted['_sort_date'] = df_sorted.apply(parse_row_date, axis=1)
        df_sorted = df_sorted.sort_values(by='_sort_date')
        
        for _, row in df_sorted.iterrows():
            f_inc = str(row.get('FECHA_INCIDENCIA', ''))
            motivo = str(row.get('TIPO_FALTA', '')).upper()
            coment = str(row.get('COMENTARIO', ''))
            bullet_text = f"• [{f_inc}] {motivo}: {coment}"
            incidents.append(bullet_text)

    # El número total de filas vacías se ajusta dinámicamente para prevenir desbordes de página
    num_incidents = len(incidents)
    empty_rows_needed = max(2, 8 - num_incidents)  # Mínimo 2 filas vacías, máximo adaptativo
    TOTAL_ROWS = num_incidents + empty_rows_needed
    
    table_detail = doc.add_table(rows=1 + TOTAL_ROWS, cols=1)
    table_detail.style = 'Table Grid'
    table_detail.autofit = False
    table_detail.columns[0].width = Inches(7.5)
    
    hdr_cell_det = table_detail.rows[0].cells[0]
    set_cell_shading(hdr_cell_det, NAVY_HEX)
    set_cell_margins(hdr_cell_det, top=45, bottom=40, left=100, right=100)
    p_hdr_det = format_cell_p(hdr_cell_det, WD_ALIGN_PARAGRAPH.CENTER)
    r_hdr_det = p_hdr_det.add_run("Detalle de la Falta")
    r_hdr_det.font.bold = True
    r_hdr_det.font.size = Pt(9)
    r_hdr_det.font.color.rgb = RGBColor(255, 255, 255)
    
    for idx in range(1, TOTAL_ROWS + 1):
        row_cell = table_detail.rows[idx].cells[0]
        row_cell.width = Inches(7.5)
        set_cell_borders(row_cell, color=BORDER_COLOR, sz="4")
        set_cell_margins(row_cell, top=25, bottom=25, left=80, right=80)
        
        p_det = format_cell_p(row_cell)
        p_det.paragraph_format.space_before = Pt(0.5)
        p_det.paragraph_format.space_after = Pt(0.5)
        
        if idx - 1 < len(incidents):
            r_det = p_det.add_run(incidents[idx - 1])
            r_det.font.size = Pt(8)
        else:
            r_det = p_det.add_run("")
            r_det.font.size = Pt(8)
            
    p_sign_note = doc.add_paragraph()
    p_sign_note.paragraph_format.space_before = Pt(6)
    p_sign_note.paragraph_format.space_after = Pt(2)
    p_sign_note.paragraph_format.line_spacing = Pt(1)
    r_sign_note = p_sign_note.add_run("Firmo en señal de solicitud/autorización/aprobación lo de arriba detallado.")
    r_sign_note.font.italic = True
    r_sign_note.font.size = Pt(8)
    
    # 5. Tabla Footer (Estructurada en Rejilla Unificada compacta para calce de página)
    footer_table = doc.add_table(rows=3, cols=3)
    footer_table.style = 'Table Grid'
    footer_table.autofit = False
    
    col_widths_foot = [Inches(1.9), Inches(1.7), Inches(3.9)]
    for row in footer_table.rows:
        for i, w in enumerate(col_widths_foot):
            row.cells[i].width = w
            
    # Fecha/Hora del sistema en tiempo real
    hn_now = datetime.utcnow() - timedelta(hours=6)
    system_date_str = hn_now.strftime("%d/%m/%Y")
    system_time_str = hn_now.strftime("%I:%M%p").lower()
    
    # Configurar Fila 0 (Encabezados Navy compactos)
    cell_navy_left = footer_table.rows[0].cells[0].merge(footer_table.rows[0].cells[1])
    set_cell_shading(cell_navy_left, NAVY_HEX)
    set_cell_borders(cell_navy_left, color=BORDER_COLOR, sz="4")
    set_cell_margins(cell_navy_left, top=30, bottom=30, left=60, right=60)
    format_cell_p(cell_navy_left)
    
    cell_navy_right = footer_table.rows[0].cells[2]
    set_cell_shading(cell_navy_right, NAVY_HEX)
    set_cell_borders(cell_navy_right, color=BORDER_COLOR, sz="4")
    set_cell_margins(cell_navy_right, top=30, bottom=30, left=60, right=60)
    p_rhdr = format_cell_p(cell_navy_right, WD_ALIGN_PARAGRAPH.CENTER)
    r_rhdr = p_rhdr.add_run("Firma de Jefe de\nDepartamento Solicitante")
    r_rhdr.font.bold = True
    r_rhdr.font.size = Pt(8.5)
    r_rhdr.font.color.rgb = RGBColor(255, 255, 255)
    
    # Configurar Fila 1 (Fecha de Elaboración)
    cell_f_lbl = footer_table.rows[1].cells[0]
    set_cell_borders(cell_f_lbl, color=BORDER_COLOR, sz="4")
    set_cell_margins(cell_f_lbl, top=30, bottom=30, left=80, right=80)
    p_f_lbl = format_cell_p(cell_f_lbl)
    r_f_lbl = p_f_lbl.add_run("Fecha de elaboración\nde reporte")
    r_f_lbl.font.bold = True
    r_f_lbl.font.size = Pt(8.5)
    
    cell_f_val = footer_table.rows[1].cells[1]
    set_cell_borders(cell_f_val, color=BORDER_COLOR, sz="4")
    set_cell_margins(cell_f_val, top=30, bottom=30, left=80, right=80)
    p_f_val = format_cell_p(cell_f_val)
    p_f_val.paragraph_format.space_before = Pt(4)
    r_f_val = p_f_val.add_run(system_date_str)
    r_f_val.font.size = Pt(8.5)
    
    # Configurar Fila 2 (Hora de Elaboración)
    cell_h_lbl = footer_table.rows[2].cells[0]
    set_cell_borders(cell_h_lbl, color=BORDER_COLOR, sz="4")
    set_cell_margins(cell_h_lbl, top=30, bottom=30, left=80, right=80)
    p_h_lbl = format_cell_p(cell_h_lbl)
    r_h_lbl = p_h_lbl.add_run("Hora de elaboración\nde reporte")
    r_h_lbl.font.bold = True
    r_h_lbl.font.size = Pt(8.5)
    
    cell_h_val = footer_table.rows[2].cells[1]
    set_cell_borders(cell_h_val, color=BORDER_COLOR, sz="4")
    set_cell_margins(cell_h_val, top=30, bottom=30, left=80, right=80)
    p_h_val = format_cell_p(cell_h_val)
    p_h_val.paragraph_format.space_before = Pt(4)
    r_h_val = p_h_val.add_run(system_time_str)
    r_h_val.font.size = Pt(8.5)
    
    # Unificar espacio para firma física (snug fit para evitar saltos de página)
    cell_signature_box = footer_table.rows[1].cells[2].merge(footer_table.rows[2].cells[2])
    set_cell_borders(cell_signature_box, color=BORDER_COLOR, sz="4")
    set_cell_margins(cell_signature_box, top=120, bottom=120, left=60, right=60)
    p_sig = format_cell_p(cell_signature_box)
    p_sig.add_run("")
    
    b_io = io.BytesIO()
    doc.save(b_io)
    return b_io.getvalue()


# ==============================================================================
# MOTOR DE INVENTARIO: CONTROL DE EQUIPOS Y DROP ACOMETIDAS
# ==============================================================================
def clasificar_materiales(row) -> str:
    """
    Analiza la actividad y el comentario de cierre para categorizar de forma robusta
    si la transacción fue un cambio de equipo terminal o un reemplazo de cable Drop.
    """
    act = str(row.get('ACTIVIDAD', '')).upper()
    com = str(row.get('COMENTARIO', '')).upper()
    texto = act + " " + com
    
    # 1. Detección de cableado y drop de acometidas
    if any(k in texto for k in ['DROP', 'ACOMETIDA', 'REEMPLAZO CABLE', 'REEMPLAZO DE CABLE', 'TENDIDO CABLE', 'CABLE DROP']):
        return 'CAMBIO_ACOMETIDA'
        
    # 2. Detección de equipos terminales (ONT, ONU, CPE o Módems)
    if any(k in texto for k in ['ONT', 'ONU', 'CPE', 'CPE ', 'MODEM', 'EQUIPO', 'CAMBIO EQUIPO', 'ROUTER', 'CAMBIO DE EQUIPO']):
        return 'CAMBIO_EQUIPO'
        
    return 'OTROS'


def generar_pdf_materiales_mensual(df_equipos, df_acometidas, tech_summary, mes_sel, anio_sel) -> bytes:
    """
    Dibuja y genera el reporte PDF de materiales e inventarios de forma ejecutiva.
    """
    pdf = ReporteGenerencialPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(40, 50, 100)
    pdf.cell(0, 10, safestr("REPORTE MENSUAL DE CONTROL DE MATERIALES"), ln=True, align="C")
    
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 6, safestr(f"Periodo: {mes_sel} {anio_sel}"), ln=True, align="C")
    pdf.ln(10)
    
    # --- SECCIÓN 1: RESUMEN DE INTERVENCIONES POR COLABORADOR ---
    pdf.seccion_titulo("1. RESUMEN DE INTERVENCIONES POR COLABORADOR")
    if not tech_summary.empty:
        pdf.set_fill_color(230, 235, 245)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "B", 8)
        
        headers = ["COLABORADOR", "EQUIPOS CAMBIADOS", "ACOMETIDAS CAMBIADAS", "TOTAL INTERVENCIONES"]
        w = [80, 40, 40, 30]
        for idx_h, h in enumerate(headers):
            pdf.cell(w[idx_h], 7, safestr(h), border=1, fill=True, align="C")
        pdf.ln()
        
        pdf.set_font("Helvetica", "", 8)
        for _, row in tech_summary.iterrows():
            pdf.cell(w[0], 6, safestr(row.get('TECNICO', '')), border=1)
            pdf.cell(w[1], 6, str(int(row.get('Equipos Cambiados', 0))), border=1, align="C")
            pdf.cell(w[2], 6, str(int(row.get('Acometidas Cambiadas', 0))), border=1, align="C")
            pdf.cell(w[3], 6, str(int(row.get('Total Intervenciones', 0))), border=1, align="C")
            pdf.ln()
    else:
        pdf.cell(0, 6, "Sin transacciones de materiales registradas en el periodo.", ln=True)
        
    pdf.ln(10)
    
    # --- SECCIÓN 2: DETALLE DE CAMBIOS DE EQUIPOS ---
    pdf.seccion_titulo("2. DETALLE DE CAMBIOS DE EQUIPO TERMINAL (Primeras 20 transacciones)")
    if not df_equipos.empty:
        pdf.set_fill_color(240, 240, 240)
        pdf.set_font("Helvetica", "B", 7)
        headers_eq = ["ORDEN", "CLIENTE", "TECNICO", "COMENTARIO DE CIERRE"]
        w_eq = [20, 25, 45, 100]
        for idx_h, h in enumerate(headers_eq):
            pdf.cell(w_eq[idx_h], 6, safestr(h), border=1, fill=True, align="C")
        pdf.ln()
        
        pdf.set_font("Helvetica", "", 7)
        for _, row in df_equipos.head(20).iterrows():
            p_com = safestr(row.get('COMENTARIO', ''))[:65]
            pdf.cell(w_eq[0], 5, safestr(row.get('NUM', '')), border=1, align="C")
            pdf.cell(w_eq[1], 5, safestr(row.get('CLIENTE', '')), border=1, align="C")
            pdf.cell(w_eq[2], 5, safestr(row.get('TECNICO', ''))[:28], border=1)
            pdf.cell(w_eq[3], 5, p_com, border=1)
            pdf.ln()
    else:
        pdf.cell(0, 6, "No se registraron cambios de equipos en el mes.", ln=True)
        
    pdf.ln(10)
    
    # --- SECCIÓN 3: DETALLE DE ACOMETIDAS ---
    pdf.seccion_titulo("3. DETALLE DE REEMPLAZOS DE ACOMETIDA - DROP (Primeras 20 transacciones)")
    if not df_acometidas.empty:
        pdf.set_fill_color(240, 240, 240)
        pdf.set_font("Helvetica", "B", 7)
        headers_ac = ["ORDEN", "CLIENTE", "TECNICO", "COMENTARIO DE CIERRE"]
        w_ac = [20, 25, 45, 100]
        for idx_h, h in enumerate(headers_ac):
            pdf.cell(w_ac[idx_h], 6, safestr(h), border=1, fill=True, align="C")
        pdf.ln()
        
        pdf.set_font("Helvetica", "", 7)
        for _, row in df_acometidas.head(20).iterrows():
            p_com = safestr(row.get('COMENTARIO', ''))[:65]
            pdf.cell(w_ac[0], 5, safestr(row.get('NUM', '')), border=1, align="C")
            pdf.cell(w_ac[1], 5, safestr(row.get('CLIENTE', '')), border=1, align="C")
            pdf.cell(w_ac[2], 5, safestr(row.get('TECNICO', ''))[:28], border=1)
            pdf.cell(w_ac[3], 5, p_com, border=1)
            pdf.ln()
    else:
        pdf.cell(0, 6, "No se registraron reemplazos de acometidas en el mes.", ln=True)
        
    return finalizar_pdf(pdf)


# ==============================================================================
# MOTOR DE INVENTARIO Y CONTROL DE CALIDAD (ccalidad.py) - WATI, PDF Y ELIMINACIÓN
# ==============================================================================

def guardar_registro_calidad(conn, data_dict: dict) -> bool:
    """
    Guarda una evaluación de calidad tanto en Google Sheets como en un histórico CSV en GCS.
    Detecta automáticamente si la pestaña 'Calidad' no existe en Google Sheets y la crea.
    """
    # Definir localmente el nombre del bucket de GCS de forma segura para evitar NameError
    try:
        df_new = pd.DataFrame([data_dict])
        
        # 1. Guardar copia de seguridad en Google Cloud Storage (GCS)
        df_gcs = leer_espejo_gcs(NOMBRE_BUCKET_SISTEMA, "calidad_maestro.csv")
        if df_gcs is not None and not df_gcs.empty:
            df_combined = pd.concat([df_gcs, df_new], ignore_index=True)
        else:
            df_combined = df_new
        sobrescribir_archivo_gcs(df_combined, NOMBRE_BUCKET_SISTEMA, "calidad_maestro.csv")
        
        # 2. Guardar en Google Sheets (Si la conexión está disponible)
        if conn is not None:
            try:
                # Intentamos leer la pestaña "Calidad"
                df_sheets = conn.read(spreadsheet=st.secrets["url_base_datos"], worksheet="Calidad", ttl=0)
                if df_sheets is not None and not df_sheets.empty:
                    df_sheets.columns = df_sheets.columns.astype(str).str.upper().str.strip()
                    df_new.columns = df_new.columns.astype(str).str.upper().str.strip()
                    df_sheets_combined = pd.concat([df_sheets, df_new], ignore_index=True)
                    conn.update(spreadsheet=st.secrets["url_base_datos"], worksheet="Calidad", data=df_sheets_combined)
                else:
                    # Si la pestaña está completamente vacía, guardamos los datos directamente
                    df_new.columns = df_new.columns.astype(str).str.upper().str.strip()
                    conn.update(spreadsheet=st.secrets["url_base_datos"], worksheet="Calidad", data=df_new)
            except Exception:
                # Si la pestaña "Calidad" no existe, forzamos a Google Sheets a CREAR la pestaña
                try:
                    df_new.columns = df_new.columns.astype(str).str.upper().str.strip()
                    conn.update(spreadsheet=st.secrets["url_base_datos"], worksheet="Calidad", data=df_new)
                except Exception as e_create:
                    print(f"Error al crear y escribir la pestaña Calidad en Google Sheets: {e_create}")
                    return False
        return True
    except Exception as e:
        print(f"Error general en guardar_registro_calidad: {e}")
        return False

def disparar_encuesta_wati(data_dict: dict) -> bool:
    """
    Envía una plantilla de mensaje oficial de WhatsApp directamente a través de la API REST de WATI.
    """
    wati_config = st.secrets.get("wati", {})
    api_url = wati_config.get("api_url", None)
    access_token = wati_config.get("access_token", None)
    template_name = wati_config.get("template_name", None)
    
    if not api_url or not access_token or not template_name:
        print("WATI API_URL, ACCESS_TOKEN o TEMPLATE_NAME no configurado en st.secrets['wati']")
        return False
        
    try:
        tel_clean = re.sub(r'\D', '', str(data_dict.get("TELEFONO", "")))
        if len(tel_clean) == 8:
            tel_clean = "504" + tel_clean
            
        endpoint = f"{api_url.rstrip('/')}/api/v1/sendTemplateMessageByPhone?whatsappNumber={tel_clean}"
        
        payload = {
            "template_name": template_name,
            "broadcast_name": f"QA_ORD_{data_dict.get('NUM_ORDEN')}",
            "parameters": [
                {"name": "cliente", "value": data_dict.get("NOMBRE_CLIENTE")},
                {"name": "tecnico", "value": data_dict.get("TECNICO")},
                {"name": "orden", "value": f"ORD-{data_dict.get('NUM_ORDEN')}"}
            ]
        }
        
        headers = {
            "Authorization": access_token if access_token.startswith("Bearer ") else f"Bearer {access_token}",
            "Content-Type": "application/json"
        }
        
        response = requests.post(endpoint, json=payload, headers=headers, timeout=15)
        return response.status_code in [200, 201, 202]
    except Exception as e:
        print(f"Error al disparar mensaje de WATI: {e}")
        return False

def eliminar_registro_calidad(conn, ticket: str, fecha_gestion: str) -> bool:
    """
    Elimina un registro específico del histórico de calidad (GCS y Google Sheets)
    basado en el Ticket y la Fecha de Gestión exacta.
    """
    try:
        # 1. Leer GCS y filtrar fila a eliminar
        df_gcs = leer_espejo_gcs(NOMBRE_BUCKET_SISTEMA, "calidad_maestro.csv")
        if df_gcs is not None and not df_gcs.empty:
            mask_eliminar_gcs = (df_gcs['TICKET'].astype(str) == str(ticket)) & (df_gcs['FECHA_GESTION'].astype(str) == str(fecha_gestion))
            df_gcs_filtered = df_gcs[~mask_eliminar_gcs].copy()
            sobrescribir_archivo_gcs(df_gcs_filtered, NOMBRE_BUCKET_SISTEMA, "calidad_maestro.csv")
            
        # 2. Leer e intentar actualizar Google Sheets
        if conn is not None:
            try:
                df_sheets = conn.read(spreadsheet=st.secrets["url_base_datos"], worksheet="Calidad", ttl=0)
                if df_sheets is not None:
                    df_sheets.columns = df_sheets.columns.astype(str).str.upper().str.strip()
                    mask_eliminar_sheets = (df_sheets['TICKET'].astype(str) == str(ticket)) & (df_sheets['FECHA_GESTION'].astype(str) == str(fecha_gestion))
                    df_sheets_filtered = df_sheets[~mask_eliminar_sheets].copy()
                    
                    # Si quedó vacío tras eliminar, guardamos un DataFrame con las mismas columnas
                    if df_sheets_filtered.empty:
                        df_sheets_filtered = pd.DataFrame(columns=df_sheets.columns)
                        
                    conn.update(spreadsheet=st.secrets["url_base_datos"], worksheet="Calidad", data=df_sheets_filtered)
            except Exception as e_sheet:
                print(f"Error al eliminar en Google Sheets: {e_sheet}")
        return True
    except Exception as e:
        print(f"Error general al eliminar registro de calidad: {e}")
        return False

def generar_pdf_reporte_calidad(df_filtered, f_inicio, f_fin) -> bytes:
    """
    Genera un reporte PDF gerencial detallado de las auditorías de calidad filtradas.
    Evita KeyErrors utilizando búsquedas seguras y fallbacks automáticos.
    """
    pdf = ReporteGenerencialPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(40, 50, 100)
    pdf.cell(0, 10, safestr("REPORTE GERENCIAL: AUDITORÍAS DE CALIDAD"), ln=True, align="C")
    
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 100, 100)
    
    inicio_str = f_inicio.strftime('%d/%m/%Y') if hasattr(f_inicio, 'strftime') else str(f_inicio)
    fin_str = f_fin.strftime('%d/%m/%Y') if hasattr(f_fin, 'strftime') else str(f_fin)
    pdf.cell(0, 6, safestr(f"Rango de Fechas: {inicio_str} al {fin_str}"), ln=True, align="C")
    pdf.ln(5)
    
    if df_filtered.empty:
        pdf.set_font("Helvetica", "I", 10)
        pdf.cell(0, 10, "No se encontraron registros de calidad en el rango seleccionado.", ln=True, align="C")
        return finalizar_pdf(pdf)
        
    # Calcular métricas básicas con tolerancia a columnas ausentes
    total_evaluaciones = len(df_filtered)
    
    # Mapeo inteligente de CSAT / Satisfacción General
    if 'CSAT' in df_filtered.columns:
        df_filtered['CSAT_NUM'] = pd.to_numeric(df_filtered['CSAT'], errors='coerce')
    elif 'SATISFACCION_GENERAL' in df_filtered.columns:
        # Traducimos las respuestas de texto de la nueva encuesta a escala numérica para calcular promedio
        satisfaccion_map = {
            "MUY SATISFECHO": 5.0,
            "SATISFECHO": 4.0,
            "POCO SATISFECHO": 2.5,
            "INSATISFECHO": 1.0,
            "N/A": np.nan,
            "PENDIENTE": np.nan
        }
        df_filtered['CSAT_NUM'] = df_filtered['SATISFACCION_GENERAL'].astype(str).str.upper().str.strip().map(satisfaccion_map)
    else:
        df_filtered['CSAT_NUM'] = np.nan
        
    promedio_csat = df_filtered['CSAT_NUM'].mean()
    promedio_csat_str = f"{promedio_csat:.2f} / 5.0" if pd.notnull(promedio_csat) and not pd.isna(promedio_csat) else "N/A"
    
    # Conteo de estados de aprobación
    aprobadas = 0
    observadas = 0
    rechazadas = 0
    
    if 'APROBACION_INTERNA' in df_filtered.columns:
        conteo_aprobacion = df_filtered['APROBACION_INTERNA'].value_counts()
        aprobadas = conteo_aprobacion.get("Servicio aprobado", 0)
        observadas = conteo_aprobacion.get("Servicio con observaciones", 0)
        rechazadas = conteo_aprobacion.get("Servicio no aprobado – requiere seguimiento", 0)
    
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(40, 50, 100)
    pdf.cell(0, 6, safestr("1. RESUMEN DE INDICADORES CLAVE"), ln=True)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(50, 50, 50)
    pdf.cell(0, 6, safestr(f"Total Auditorías Realizadas: {total_evaluaciones}"), ln=True)
    pdf.cell(0, 6, safestr(f"Índice de Satisfacción Promedio (CSAT Equivalente): {promedio_csat_str}"), ln=True)
    pdf.cell(0, 6, safestr(f"Estatus de Aprobación: Aprobadas: {aprobadas} | Con Observaciones: {observadas} | Reclamos/Rechazadas: {rechazadas}"), ln=True)
    pdf.ln(5)
    
    # Tabla detallada (190mm de ancho total)
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(40, 50, 100)
    pdf.cell(0, 6, safestr("2. LISTADO DETALLADO DE AUDITORÍAS"), ln=True)
    
    pdf.set_fill_color(230, 235, 245)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "B", 7)
    
    w = [25, 45, 50, 15, 30, 25]
    headers = ["TICKET", "TÉCNICO", "CLIENTE", "SATISF.", "ESTÉTICA", "APROBACIÓN"]
    
    for i in range(len(headers)):
        pdf.cell(w[i], 7, safestr(headers[i]), border=1, fill=True, align="C")
    pdf.ln()
    
    pdf.set_font("Helvetica", "", 7)
    for _, row in df_filtered.iterrows():
        if pdf.get_y() > 270:
            pdf.add_page()
            pdf.set_fill_color(230, 235, 245)
            pdf.set_font("Helvetica", "B", 7)
            for i in range(len(headers)):
                pdf.cell(w[i], 7, safestr(headers[i]), border=1, fill=True, align="C")
            pdf.ln()
            pdf.set_font("Helvetica", "", 7)
            
        ticket = str(row.get('TICKET', 'N/D'))
        tec = str(row.get('TECNICO', 'N/D'))[:22]
        cli = str(row.get('NOMBRE_CLIENTE', 'N/D'))[:24]
        
        # Obtener nivel de satisfacción de forma segura
        satisf_val = str(row.get('CSAT', ''))
        if not satisf_val or satisf_val.upper() in ["NAN", "NONE", "N/D", "", "PENDIENTE"]:
            satisf_val = str(row.get('SATISFACCION_GENERAL', 'N/D'))
            # Abreviar los textos largos del nuevo CSAT para que quepan en la columna de ancho 15
            if "muy satisfecho" in satisf_val.lower(): satisf_val = "Muy Sat."
            elif "insatisfecho" in satisf_val.lower(): satisf_val = "Insat."
            elif "poco satisfecho" in satisf_val.lower(): satisf_val = "Poco Sat."
            elif "satisfecho" in satisf_val.lower(): satisf_val = "Sat."
            
        estetica = str(row.get('ESTETICA', 'N/D'))[:15]
        aprob = str(row.get('APROBACION_INTERNA', 'N/D'))
        
        if "aprobado" in aprob.lower() and "no" not in aprob.lower() and "con" not in aprob.lower():
            aprob_str = "Aprobado"
        elif "observaciones" in aprob.lower():
            aprob_str = "Con Obs."
        else:
            aprob_str = "Rechazado"
            
        pdf.cell(w[0], 6, safestr(ticket), border=1, align="C")
        pdf.cell(w[1], 6, safestr(tec), border=1, align="L")
        pdf.cell(w[2], 6, safestr(cli), border=1, align="L")
        pdf.cell(w[3], 6, safestr(satisf_val), border=1, align="C")
        pdf.cell(w[4], 6, safestr(estetica), border=1, align="C")
        pdf.cell(w[5], 6, safestr(aprob_str), border=1, align="C")
        pdf.ln()
        
    return finalizar_pdf(pdf)

# ==============================================================================
# AUDITORÍA DE CAMPO: GESTIÓN DE OPERACIONES E INSTALACIONES (ccalidad.py)
# ==============================================================================

def guardar_auditoria_campo(conn, data_dict: dict, tipo: str) -> bool:
    """
    Guarda registros de auditoría de campo (operaciones o instalaciones)
    en Google Sheets y realiza una copia de respaldo en GCS de forma segura.
    """
    worksheet_name = "Operaciones" if tipo == "operaciones" else "Instalaciones"
    filename_gcs = "operaciones_maestro.csv" if tipo == "operaciones" else "instalaciones_maestro.csv"
    
    try:
        df_new = pd.DataFrame([data_dict])
        
        # 1. Respaldo en GCS
        df_gcs = leer_espejo_gcs(NOMBRE_BUCKET_SISTEMA, filename_gcs)
        if df_gcs is not None and not df_gcs.empty:
            df_combined = pd.concat([df_gcs, df_new], ignore_index=True)
        else:
            df_combined = df_new
        sobrescribir_archivo_gcs(df_combined, NOMBRE_BUCKET_SISTEMA, filename_gcs)
        
        # 2. Google Sheets
        if conn is not None:
            try:
                df_sheets = conn.read(spreadsheet=st.secrets["url_base_datos"], worksheet=worksheet_name, ttl=0)
                if df_sheets is not None and not df_sheets.empty:
                    df_sheets.columns = df_sheets.columns.astype(str).str.upper().str.strip()
                    df_new.columns = df_new.columns.astype(str).str.upper().str.strip()
                    df_sheets_combined = pd.concat([df_sheets, df_new], ignore_index=True)
                    conn.update(spreadsheet=st.secrets["url_base_datos"], worksheet=worksheet_name, data=df_sheets_combined)
                else:
                    df_new.columns = df_new.columns.astype(str).str.upper().str.strip()
                    conn.update(spreadsheet=st.secrets["url_base_datos"], worksheet=worksheet_name, data=df_new)
            except Exception:
                try:
                    df_new.columns = df_new.columns.astype(str).str.upper().str.strip()
                    conn.update(spreadsheet=st.secrets["url_base_datos"], worksheet=worksheet_name, data=df_new)
                except Exception as e_create:
                    print(f"Error al crear pestaña {worksheet_name} en Google Sheets: {e_create}")
                    return False
        return True
    except Exception as e:
        print(f"Error en guardar_auditoria_campo ({tipo}): {e}")
        return False

def eliminar_registro_campo(conn, ticket: str, fecha_gestion: str, tipo: str) -> bool:
    """
    Elimina un registro específico de auditoría de campo (operaciones o instalaciones)
    en Google Sheets y realiza la copia de respaldo en GCS.
    """
    worksheet_name = "Operaciones" if tipo == "operaciones" else "Instalaciones"
    filename_gcs = "operaciones_maestro.csv" if tipo == "operaciones" else "instalaciones_maestro.csv"
    
    try:
        # 1. Leer GCS y filtrar fila a eliminar
        df_gcs = leer_espejo_gcs(NOMBRE_BUCKET_SISTEMA, filename_gcs)
        if df_gcs is not None and not df_gcs.empty:
            mask_eliminar_gcs = (df_gcs['ORDEN_NUM'].astype(str) == str(ticket)) & (df_gcs['FECHA_AUDITORIA'].astype(str) == str(fecha_gestion))
            df_gcs_filtered = df_gcs[~mask_eliminar_gcs].copy()
            sobrescribir_archivo_gcs(df_gcs_filtered, NOMBRE_BUCKET_SISTEMA, filename_gcs)
            
        # 2. Leer e intentar actualizar Google Sheets
        if conn is not None:
            try:
                df_sheets = conn.read(spreadsheet=st.secrets["url_base_datos"], worksheet=worksheet_name, ttl=0)
                if df_sheets is not None:
                    df_sheets.columns = df_sheets.columns.astype(str).str.upper().str.strip()
                    mask_eliminar_sheets = (df_sheets['ORDEN_NUM'].astype(str) == str(ticket)) & (df_sheets['FECHA_AUDITORIA'].astype(str) == str(fecha_gestion))
                    df_sheets_filtered = df_sheets[~mask_eliminar_sheets].copy()
                    
                    if df_sheets_filtered.empty:
                        df_sheets_filtered = pd.DataFrame(columns=df_sheets.columns)
                        
                    conn.update(spreadsheet=st.secrets["url_base_datos"], worksheet=worksheet_name, data=df_sheets_filtered)
            except Exception as e_sheet:
                print(f"Error al eliminar en Google Sheets: {e_sheet}")
        return True
    except Exception as e:
        print(f"Error al eliminar registro de campo: {e}")
        return False

def generar_pdf_reporte_campo(df_filtered, f_inicio, f_fin, tipo: str) -> bytes:
    """
    Genera un reporte PDF gerencial detallado de las auditorías de campo filtradas (Operaciones o Instalaciones).
    """
    pdf = ReporteGenerencialPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(40, 50, 100)
    titulo_rep = "REPORTE GERENCIAL: AUDITORÍA DE OPERACIONES" if tipo == "operaciones" else "REPORTE GERENCIAL: AUDITORÍA DE INSTALACIONES (INSFIBRA)"
    pdf.cell(0, 10, safestr(titulo_rep), ln=True, align="C")
    
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 100, 100)
    inicio_str = f_inicio.strftime('%d/%m/%Y') if hasattr(f_inicio, 'strftime') else str(f_inicio)
    fin_str = f_fin.strftime('%d/%m/%Y') if hasattr(f_fin, 'strftime') else str(f_fin)
    pdf.cell(0, 6, safestr(f"Rango de Fechas: {inicio_str} al {fin_str}"), ln=True, align="C")
    pdf.ln(5)
    
    if df_filtered.empty:
        pdf.set_font("Helvetica", "I", 10)
        pdf.cell(0, 10, "No se encontraron registros de auditoría en el rango seleccionado.", ln=True, align="C")
        return finalizar_pdf(pdf)
        
    pdf.seccion_titulo("1. RESUMEN DE COMPROBACIONES TÉCNICAS")
    total_auditorias = len(df_filtered)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(50, 50, 50)
    pdf.cell(0, 6, safestr(f"Total Auditorías de Campo Realizadas: {total_auditorias}"), ln=True)
    pdf.ln(5)
    
    pdf.seccion_titulo("2. LISTADO DETALLADO DE COMPROBACIONES")
    pdf.set_fill_color(230, 235, 245)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "B", 7)
    
    if tipo == "operaciones":
        w = [25, 25, 45, 25, 25, 45]
        headers = ["ORDEN", "CLIENTE", "SERVICIO", "VIÑETA", "MUFA", "ESTÉTICA"]
        for i in range(len(headers)):
            pdf.cell(w[i], 7, safestr(headers[i]), border=1, fill=True, align="C")
        pdf.ln()
        
        pdf.set_font("Helvetica", "", 7)
        for _, row in df_filtered.iterrows():
            if pdf.get_y() > 270:
                pdf.add_page()
                pdf.set_fill_color(230, 235, 245)
                pdf.set_font("Helvetica", "B", 7)
                for i in range(len(headers)):
                    pdf.cell(w[i], 7, safestr(headers[i]), border=1, fill=True, align="C")
                pdf.ln()
                pdf.set_font("Helvetica", "", 7)
                
            pdf.cell(w[0], 6, safestr(str(row.get('ORDEN_NUM', ''))), border=1, align="C")
            pdf.cell(w[1], 6, safestr(str(row.get('CODIGO_CLIENTE', ''))), border=1, align="C")
            pdf.cell(w[2], 6, safestr(str(row.get('CODIGO_SERVICIO', '')))[:24], border=1, align="L")
            pdf.cell(w[3], 6, safestr(str(row.get('VINETA', ''))), border=1, align="C")
            pdf.cell(w[4], 6, safestr(str(row.get('MUFA', ''))), border=1, align="C")
            pdf.cell(w[5], 6, safestr(str(row.get('ESTETICA', ''))), border=1, align="C")
            pdf.ln()
    else:
        w = [20, 20, 45, 30, 25, 25, 25]
        headers = ["ORDEN", "CLIENTE", "TECNICO", "TIPO FO", "METROS", "VIÑETA", "MUFA"]
        for i in range(len(headers)):
            pdf.cell(w[i], 7, safestr(headers[i]), border=1, fill=True, align="C")
        pdf.ln()
        
        pdf.set_font("Helvetica", "", 7)
        for _, row in df_filtered.iterrows():
            if pdf.get_y() > 270:
                pdf.add_page()
                pdf.set_fill_color(230, 235, 245)
                pdf.set_font("Helvetica", "B", 7)
                for i in range(len(headers)):
                    pdf.cell(w[i], 7, safestr(headers[i]), border=1, fill=True, align="C")
                pdf.ln()
                pdf.set_font("Helvetica", "", 7)
                
            pdf.cell(w[0], 6, safestr(str(row.get('ORDEN_NUM', ''))), border=1, align="C")
            pdf.cell(w[1], 6, safestr(str(row.get('CODIGO_CLIENTE', ''))), border=1, align="C")
            pdf.cell(w[2], 6, safestr(str(row.get('TECNICO', '')))[:24], border=1, align="L")
            pdf.cell(w[3], 6, safestr(str(row.get('TIPO_FO', ''))), border=1, align="C")
            pdf.cell(w[4], 6, safestr(str(row.get('METROS_FO', ''))), border=1, align="C")
            pdf.cell(w[5], 6, safestr(str(row.get('VINETA', ''))), border=1, align="C")
            pdf.cell(w[6], 6, safestr(str(row.get('MUFA', ''))), border=1, align="C")
            pdf.ln()
            
    return finalizar_pdf(pdf)
